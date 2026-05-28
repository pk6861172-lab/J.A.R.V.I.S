#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════╗
║        J.A.R.V.I.S  v2.0 — Personal AI Assistant                ║
║        For: Prashant | Powered by OpenRouter (Free)              ║
║        Features: Voice • Camera • Email • Weather • System       ║
╚══════════════════════════════════════════════════════════════════╝
"""

# ═══════════════════════════════════════════════
#  STANDARD LIBRARY
# ═══════════════════════════════════════════════
import os, sys, json, time, datetime, subprocess, mimetypes, shutil, queue, ctypes
import threading, platform, webbrowser, random, warnings
warnings.filterwarnings("ignore", category=UserWarning, module="pydantic")
import re, socket, imaplib, smtplib, base64
import email as email_lib
import csv
from email.utils import parseaddr
from email import encoders
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.header import decode_header
from pathlib import Path

try:
    from jarvis_modules.browser_matching import best_text_match
    from jarvis_modules.disk_audit import build_disk_cleanup_report
    from jarvis_modules.elevenlabs_tts import DEFAULT_VOICE_ID, api_key_from_config, is_enabled as elevenlabs_is_enabled, synthesize_speech
    from jarvis_modules.proactive_engine import ProactiveEngine
    from jarvis_modules.self_knowledge import build_self_knowledge, compact_self_knowledge_text, load_self_knowledge
    from jarvis_modules.self_improvement import looks_like_self_improvement_request, response_for_request, save_self_improvement_request
except Exception:
    best_text_match = None
    build_disk_cleanup_report = None
    DEFAULT_VOICE_ID = "HH8sIQq8WOcER3Nu118i"
    api_key_from_config = None
    elevenlabs_is_enabled = None
    synthesize_speech = None
    ProactiveEngine = None
    build_self_knowledge = None
    compact_self_knowledge_text = None
    load_self_knowledge = None
    looks_like_self_improvement_request = None
    response_for_request = None
    save_self_improvement_request = None

# ═══════════════════════════════════════════════
#  LOAD .env FILE (keeps API keys secure)
# ═══════════════════════════════════════════════
try:
    from dotenv import load_dotenv
    load_dotenv()  # loads .env from same folder automatically
except ImportError:
    pass  # dotenv not installed — keys will be read from jarvis_config.json

# Ensure stdout/stderr use UTF-8 on Windows to avoid 'charmap' encoding errors when printing
try:
    import sys
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass


# ═══════════════════════════════════════════════
#  SAFE IMPORTS — won't crash if missing
# ═══════════════════════════════════════════════
def try_import(module, pip_name=None, attr=None):
    try:
        mod = __import__(module)
        return getattr(mod, attr) if attr else mod
    except ImportError:
        return None

# Voice
sr        = try_import("speech_recognition")
pyttsx3   = try_import("pyttsx3")
vosk        = try_import("vosk")
openwakeword_mod = None
pywhatkit   = None
face_recognition = None
_openwakeword_import_attempted = False
_pywhatkit_import_attempted = False
_face_recognition_import_attempted = False
pytesseract = try_import("pytesseract")
if pytesseract:
    import os
    if os.name == 'nt':
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
librosa_mod = try_import("librosa")

# System
psutil    = try_import("psutil")
pyautogui = try_import("pyautogui")
pyperclip = try_import("pyperclip")

# Network / AI
requests  = try_import("requests")
openai_mod = try_import("openai")

# Camera
cv2       = try_import("cv2")

# Image processing
PIL_Image = try_import("PIL.Image") or try_import("PIL", attr="Image")
mss_mod = try_import("mss")
try:
    from PIL import Image as PIL_Image
except:
    PIL_Image = None


def ensure_openwakeword():
    global openwakeword_mod, _openwakeword_import_attempted
    if openwakeword_mod is not None or _openwakeword_import_attempted:
        return openwakeword_mod
    _openwakeword_import_attempted = True
    openwakeword_mod = try_import("openwakeword")
    return openwakeword_mod


def ensure_pywhatkit():
    global pywhatkit, _pywhatkit_import_attempted
    if pywhatkit is not None or _pywhatkit_import_attempted:
        return pywhatkit
    _pywhatkit_import_attempted = True
    pywhatkit = try_import("pywhatkit")
    return pywhatkit


def ensure_face_recognition():
    global face_recognition, _face_recognition_import_attempted
    if face_recognition is not None or _face_recognition_import_attempted:
        return face_recognition
    _face_recognition_import_attempted = True
    face_recognition = try_import("face_recognition")
    return face_recognition

# ═══════════════════════════════════════════════
#  PATHS
# ═══════════════════════════════════════════════
BASE_DIR    = Path(__file__).parent
CFG_FILE    = BASE_DIR / "jarvis_config.json"
NOTES_FILE  = BASE_DIR / "jarvis_notes.json"
EVENTS_FILE = BASE_DIR / "jarvis_events.json"
PHOTOS_DIR  = BASE_DIR / "jarvis_photos"
TELEGRAM_DOWNLOAD_DIR = PHOTOS_DIR / "telegram"
GENERATED_DIR = BASE_DIR / "jarvis_generated"
AI_HISTORY_FILE = BASE_DIR / "jarvis_ai_history.json"
AI_MEMORY_FILE = BASE_DIR / "jarvis_ai_memory.json"
SELF_KNOWLEDGE_FILE = BASE_DIR / "jarvis_self_knowledge.json"
SELF_IMPROVEMENT_REQUESTS_FILE = BASE_DIR / "jarvis_self_edit_requests.json"
PHOTOS_DIR.mkdir(exist_ok=True)
TELEGRAM_DOWNLOAD_DIR.mkdir(exist_ok=True)
GENERATED_DIR.mkdir(exist_ok=True)

EMAIL_RE_STRICT = re.compile(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9-]+(?:\.[A-Za-z0-9-]+)+$")
SENSITIVE_FILE_PATTERNS = (
    "jarvis_config.json", ".env", "token", "secret", "password", "api_key",
    "apikey", "api-key", "session", "credential", "credentials",
)


def normalize_email_address(value: str) -> str | None:
    raw = str(value or "").strip().strip("<>.,;:()[]{}\"'")
    _display, parsed = parseaddr(raw)
    addr = (parsed or raw).strip().lower()
    if not EMAIL_RE_STRICT.fullmatch(addr):
        return None
    local, domain = addr.rsplit("@", 1)
    labels = domain.split(".")
    if not local or len(local) > 64 or len(domain) > 253:
        return None
    if any(not label or label.startswith("-") or label.endswith("-") for label in labels):
        return None
    if any(labels[i] == labels[i + 1] for i in range(len(labels) - 1)):
        return None
    if labels[-1] in {"com", "net", "org", "in", "co"} and len(labels) >= 3 and labels[-2] == labels[-1]:
        return None
    return addr


def text_mentions_sensitive_file(text: str) -> bool:
    lower = str(text or "").lower()
    return any(pattern in lower for pattern in SENSITIVE_FILE_PATTERNS)

TEXT_FILE_EXTENSIONS = {
    ".py", ".txt", ".md", ".json", ".csv", ".tsv", ".ini", ".cfg",
    ".yaml", ".yml", ".toml", ".log", ".xml", ".html", ".htm",
    ".css", ".js", ".ts", ".jsx", ".tsx", ".java", ".c", ".cpp",
    ".h", ".hpp", ".cs", ".go", ".rs", ".sql", ".bat", ".ps1",
    ".sh", ".env", ".gitignore",
}

IMAGE_FILE_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif",
}

# ═══════════════════════════════════════════════════════════════
#  CONFIG
# ═══════════════════════════════════════════════════════════════
def load_config() -> dict:
    if not CFG_FILE.exists():
        print(f"[!] jarvis_config.json not found at {CFG_FILE}")
        print("[!] Creating default config — please fill in your API keys.")
        default = {
            "user_name": "Prashant",
            "location": "Bihar, India",
            "interests": "cybersecurity, JEE, Python, AI, entrepreneurship",
            "ai_provider": "auto",
            "groq_api_key": "",
            "openrouter_api_key": "YOUR_OPENROUTER_KEY_HERE",
            "gemini_api_key": "",
            "gemini_vision_model": "gemini-2.0-flash",
            "agent_max_steps": 12,
            "google_maps_api_key": "",
            "location_label": "Bihar, India",
            "location_lat": None,
            "location_lon": None,
            "openweather_api_key": "",
            "news_api_key": "",
            "email": "",
            "email_password": "",
            "imap_server": "imap.gmail.com",
            "smtp_server": "smtp.gmail.com",
            "smtp_port": 587,
            "tts_rate": 172,
            "tts_volume": 0.95,
            "tts_backend": "edge_tts",
            "tts_chunk_chars": 650,
            "edge_tts_voice": "hi-IN-MadhurNeural",
            "edge_tts_rate": "+10%",
            "elevenlabs_enabled": False,
            "elevenlabs_api_key": "",
            "elevenlabs_voice_id": "HH8sIQq8WOcER3Nu118i",
            "elevenlabs_model_id": "eleven_multilingual_v2",
            "elevenlabs_stability": 0.45,
            "elevenlabs_similarity_boost": 0.75,
            "elevenlabs_style": 0.2,
            "elevenlabs_speaker_boost": True,
            "wake_words": ["jarvis", "hey jarvis"],
            # Voice (offline + wake word)
            # voice_mode: "auto" uses Vosk when offline (or when Google STT fails),
            #            "online" forces Google STT,
            #            "offline" forces Vosk STT.
            "voice_mode": "auto",
            # Path to a Vosk model folder. Example:
            #   "vosk_models/vosk-model-small-en-us-0.15"
            "vosk_model_path": "",
            # Offline wake word — openWakeWord (free, no account). Pre-trained e.g. "hey_jarvis".
            # See: https://github.com/dscripka/openWakeWord
            "openwakeword_enabled": True,
            "openwakeword_models": ["hey_jarvis"],
            "openwakeword_threshold": 0.5,
            "openwakeword_chunk_size": 1280,
            "openwakeword_inference_framework": "onnx",
            # WhatsApp (pywhatkit)
            # Default country code is only used if you type number without +.
            "whatsapp_default_country_code": "+91",
            "relay_enabled": True,
            "relay_ip": "10.194.207.247",
            "relay_timeout_sec": 3,
            "whatsapp_bridge_enabled": False,
            "whatsapp_bridge_port": 3001,
            "whatsapp_bridge_auto_start": True,
            "whatsapp_bridge_mode": "ask_first",
            "whatsapp_bridge_dm_only": True,
            "whatsapp_bridge_ignore_groups": True,
            "whatsapp_bridge_cooldown_sec": 5,
            "whatsapp_bridge_browser_path": "",
            "whatsapp_bridge_headless": True,
            "whatsapp_bridge_retry_ms": 8000,
            "whatsapp_bridge_client_id": "jarvis",
            "whatsapp_bridge_read_message_chars": 140,
            "whatsapp_bridge_call_secretary_enabled": True,
            "whatsapp_bridge_call_ignore_groups": True,
            "whatsapp_bridge_call_auto_reject": False,
            "whatsapp_bridge_call_default_reply": "Bhai Prashant abhi busy hain, thodi der me call/message karenge.",
            "whatsapp_bridge_call_cooldown_sec": 20,
            "whatsapp_bridge_call_confirm_timeout": 14,
            "whatsapp_bridge_call_reply_timeout": 20,
            "whatsapp_bridge_voice_retries": 2,
            "whatsapp_bridge_listen_after_prompt_delay": 1.0,
            "whatsapp_bridge_message_confirm_timeout": 8,
            "whatsapp_bridge_message_reply_timeout": 15,
            # Contacts CSV (name -> phone) for WhatsApp/SMS-like commands.
            # Default: "contacts.csv" in the same folder as jarvis.py
            "contacts_csv_path": "contacts.csv",
            "face_recognition_enabled": True,
            "owner_face_images": [],
            "face_match_tolerance": 0.55,
            "face_recognition_interval_frames": 10,
            "ollama_enabled": True,
            "ollama_base_url": "http://localhost:11434/v1",
            "ollama_model": "llama3.2",
            "use_ollama_when_offline": True,
            "telegram_bot_token": "",
            "telegram_allowed_user_id": None,
            "telegram_enabled": False,
            "telegram_personal_enabled": False,
            "telegram_api_id": "",
            "telegram_api_hash": "",
            "telegram_personal_session": ".jarvis_runtime/telegram_personal",
            "telegram_personal_dm_only": True,
            "telegram_personal_reply_cooldown_sec": 5,
            "telegram_personal_notify_owner": True,
            "telegram_personal_ignore_senders": ["MyJarvisBot", "myjarvisbot"],
            "telegram_secretary_ai_auto_reply": False,
            "telegram_personal_ai_auto_reply": False,
            "self_improvement_enabled": True,
            "self_improvement_apply_without_owner_review": False,
            "web_api_token": "",
            "proactive_enabled": True,
            "proactive_study_times": ["18:00", "19:00", "20:00"],
            "proactive_morning_time": "07:30",
            "proactive_idle_minutes": 20,
            "proactive_listen_after_prompt": True,
            "proactive_listen_timeout_sec": 10,
            "proactive_listen_phrase_limit": 12,
            "proactive_listen_after_prompt_delay": 1.0,
            "proactive_listen_cue": "Bol bhai, 10 second sun raha hoon.",
            "secretary_enabled": False,
            "secretary_apps": ["WhatsApp", "Instagram", "Snapchat", "Facebook", "Telegram"],
            "secretary_action_center_fallback": False,
            "secretary_reply_focus_delay": 0.35,
            "secretary_reply_click": None,
            "secretary_app_routes": {},
            "secretary_confirm_timeout": 9,
            "secretary_reply_timeout": 15,
            "secretary_voice_retries": 2,
            "secretary_listen_after_prompt_delay": 0.9,
            "secretary_open_notification_first": True,
            "secretary_notification_open_delay": 0.35,
            "secretary_notification_click_delay": 0.7,
            "secretary_wait_for_app_timeout": 2.0,
            "secretary_open_latest_notification_fallback": True,
            "secretary_latest_notification_click_ratio": [0.84, 0.18],
            "secretary_click_reply_box_after_open": True,
            "secretary_reply_box_click_ratio": [0.62, 0.945],
            "secretary_press_escape_before_reply": True,
            "camera_index": 0,
            "storage_roots": ["C:\\" if os.name == "nt" else str(Path.home())],
            "custom_apps": {}
        }
        with open(CFG_FILE, "w", encoding="utf-8") as f:
            json.dump(default, f, indent=2)
        sys.exit(1)
    with open(CFG_FILE, "r", encoding="utf-8-sig") as f:
        cfg = json.load(f)
    cfg.setdefault("ai_provider", "auto")
    cfg.setdefault("groq_api_key", "")
    cfg.setdefault("openrouter_api_key", "")
    cfg.setdefault("gemini_api_key", "")
    cfg.setdefault("gemini_vision_model", "gemini-2.0-flash")
    cfg.setdefault("gemini_request_timeout", 60)
    cfg.setdefault("gemini_max_model_attempts", 2)
    cfg.setdefault("agent_max_steps", 12)
    cfg.setdefault("google_maps_api_key", "")
    cfg.setdefault("location_label", cfg.get("location", "Bihar, India"))
    cfg.setdefault("location_lat", None)
    cfg.setdefault("location_lon", None)
    cfg.setdefault("storage_roots", ["C:\\" if os.name == "nt" else str(Path.home())])
    cfg.setdefault("voice_mode", "auto")
    cfg.setdefault("vosk_model_path", "")
    cfg.setdefault("openwakeword_enabled", True)
    cfg.setdefault("openwakeword_models", ["hey_jarvis"])
    cfg.setdefault("openwakeword_threshold", 0.5)
    cfg.setdefault("openwakeword_chunk_size", 1280)
    cfg.setdefault("openwakeword_inference_framework", "onnx")
    cfg.setdefault("whatsapp_default_country_code", "+91")
    cfg.setdefault("relay_enabled", True)
    cfg.setdefault("relay_ip", "10.194.207.247")
    cfg.setdefault("relay_timeout_sec", 3)
    cfg.setdefault("whatsapp_bridge_enabled", False)
    cfg.setdefault("whatsapp_bridge_port", 3001)
    cfg.setdefault("whatsapp_bridge_auto_start", True)
    cfg.setdefault("whatsapp_bridge_mode", "ask_first")
    cfg.setdefault("whatsapp_bridge_dm_only", True)
    cfg.setdefault("whatsapp_bridge_ignore_groups", True)
    cfg.setdefault("whatsapp_bridge_cooldown_sec", 5)
    cfg.setdefault("whatsapp_bridge_browser_path", "")
    cfg.setdefault("whatsapp_bridge_headless", True)
    cfg.setdefault("whatsapp_bridge_retry_ms", 8000)
    cfg.setdefault("whatsapp_bridge_client_id", "jarvis")
    cfg.setdefault("whatsapp_bridge_read_message_chars", 140)
    cfg.setdefault("whatsapp_bridge_call_secretary_enabled", True)
    cfg.setdefault("whatsapp_bridge_call_ignore_groups", True)
    cfg.setdefault("whatsapp_bridge_call_auto_reject", False)
    cfg.setdefault("whatsapp_bridge_call_default_reply", "Bhai Prashant abhi busy hain, thodi der me call/message karenge.")
    cfg.setdefault("whatsapp_bridge_call_cooldown_sec", 20)
    cfg.setdefault("whatsapp_bridge_call_confirm_timeout", 14)
    cfg.setdefault("whatsapp_bridge_call_reply_timeout", 20)
    cfg.setdefault("whatsapp_bridge_voice_retries", 2)
    cfg.setdefault("whatsapp_bridge_listen_after_prompt_delay", 1.0)
    cfg.setdefault("whatsapp_bridge_message_confirm_timeout", 8)
    cfg.setdefault("whatsapp_bridge_message_reply_timeout", 15)
    cfg.setdefault("contacts_csv_path", "contacts.csv")
    cfg.setdefault("face_recognition_enabled", True)
    cfg.setdefault("owner_face_images", [])
    cfg.setdefault("face_match_tolerance", 0.55)
    cfg.setdefault("face_recognition_interval_frames", 10)
    cfg.setdefault("ollama_enabled", True)
    cfg.setdefault("ollama_base_url", "http://localhost:11434/v1")
    cfg.setdefault("ollama_model", "llama3.2")
    cfg.setdefault("use_ollama_when_offline", True)
    cfg.setdefault("telegram_bot_token", "")
    cfg.setdefault("telegram_allowed_user_id", None)
    cfg.setdefault("telegram_enabled", False)
    cfg.setdefault("telegram_personal_enabled", False)
    cfg.setdefault("telegram_api_id", "")
    cfg.setdefault("telegram_api_hash", "")
    cfg.setdefault("telegram_personal_session", ".jarvis_runtime/telegram_personal")
    cfg.setdefault("telegram_personal_dm_only", True)
    cfg.setdefault("telegram_personal_reply_cooldown_sec", 5)
    cfg.setdefault("telegram_personal_notify_owner", True)
    cfg.setdefault("telegram_personal_ignore_senders", ["MyJarvisBot", "myjarvisbot"])
    cfg.setdefault("telegram_secretary_ai_auto_reply", False)
    cfg.setdefault("telegram_personal_ai_auto_reply", False)
    cfg.setdefault("self_improvement_enabled", True)
    cfg.setdefault("self_improvement_apply_without_owner_review", False)
    cfg.setdefault("web_api_token", "")
    cfg.setdefault("proactive_enabled", True)
    cfg.setdefault("proactive_study_times", ["18:00", "19:00", "20:00"])
    cfg.setdefault("proactive_morning_time", "07:30")
    cfg.setdefault("proactive_idle_minutes", 20)
    cfg.setdefault("proactive_listen_after_prompt", True)
    cfg.setdefault("proactive_listen_timeout_sec", 10)
    cfg.setdefault("proactive_listen_phrase_limit", 12)
    cfg.setdefault("proactive_listen_after_prompt_delay", 1.0)
    cfg.setdefault("proactive_listen_cue", "Bol bhai, 10 second sun raha hoon.")
    cfg.setdefault("secretary_enabled", False)
    cfg.setdefault("secretary_apps", ["WhatsApp", "Instagram", "Snapchat", "Facebook", "Telegram"])
    cfg.setdefault("secretary_action_center_fallback", False)
    cfg.setdefault("secretary_reply_focus_delay", 0.35)
    cfg.setdefault("secretary_reply_click", None)
    cfg.setdefault("secretary_app_routes", {})
    cfg.setdefault("secretary_confirm_timeout", 9)
    cfg.setdefault("secretary_reply_timeout", 15)
    cfg.setdefault("secretary_voice_retries", 2)
    cfg.setdefault("secretary_listen_after_prompt_delay", 0.9)
    cfg.setdefault("secretary_open_notification_first", True)
    cfg.setdefault("secretary_notification_open_delay", 0.35)
    cfg.setdefault("secretary_notification_click_delay", 0.7)
    cfg.setdefault("secretary_wait_for_app_timeout", 2.0)
    cfg.setdefault("secretary_open_latest_notification_fallback", True)
    cfg.setdefault("secretary_latest_notification_click_ratio", [0.84, 0.18])
    cfg.setdefault("secretary_click_reply_box_after_open", True)
    cfg.setdefault("secretary_reply_box_click_ratio", [0.62, 0.945])
    cfg.setdefault("secretary_press_escape_before_reply", True)
    cfg.setdefault("tts_backend", "edge_tts")
    cfg.setdefault("tts_chunk_chars", 650)
    cfg.setdefault("edge_tts_voice", "hi-IN-MadhurNeural")
    cfg.setdefault("edge_tts_rate", "+10%")
    cfg.setdefault("elevenlabs_enabled", False)
    cfg.setdefault("elevenlabs_api_key", "")
    cfg.setdefault("elevenlabs_voice_id", DEFAULT_VOICE_ID)
    cfg.setdefault("elevenlabs_model_id", "eleven_multilingual_v2")
    cfg.setdefault("elevenlabs_stability", 0.45)
    cfg.setdefault("elevenlabs_similarity_boost", 0.75)
    cfg.setdefault("elevenlabs_style", 0.2)
    cfg.setdefault("elevenlabs_speaker_boost", True)
    return cfg

# ═══════════════════════════════════════════════════════════════
#  VOICE ENGINE
# ═══════════════════════════════════════════════════════════════
class VoiceEngine:
    def __init__(self, cfg: dict):
        self.cfg = cfg
        self.tts_ok = False
        self.stt_ok = False
        self.engine = None
        self._tts_queue = queue.Queue()
        self._tts_thread = None
        self._tts_stop = threading.Event()
        self._tts_ready = threading.Event()
        self._vosk_model = None
        self._oww_model = None
        self._oww_threshold = 0.5
        self._oww_chunk = 1280
        self._vosk_enabled = False
        self._vosk_init_attempted = False
        self.last_voice_mood = "neutral"
        self.last_voice_mood_hint = ""
        self._init_tts()
        self._init_stt()

    def _init_tts(self):
        if not pyttsx3:
            print("[WARN] pyttsx3 not installed — voice output disabled.")
            return
        try:
            self._tts_thread = threading.Thread(target=self._tts_worker, daemon=True)
            self._tts_thread.start()
            self._tts_ready.wait(timeout=3.0)
        except Exception as e:
            print(f"[WARN] TTS init failed: {e}")

    def _tts_worker(self):
        try:
            engine = pyttsx3.init()
            voices = engine.getProperty("voices")
            voice_id = None
            for v in voices:
                if any(k in v.name.lower() for k in ["david", "mark", "daniel", "george", "zira"]):
                    voice_id = v.id
                    break
            if voice_id:
                engine.setProperty("voice", voice_id)
            engine.setProperty("rate", self.cfg.get("tts_rate", 172))
            engine.setProperty("volume", self.cfg.get("tts_volume", 0.95))
            self.engine = engine
            self.tts_ok = True
        except Exception as e:
            print(f"[WARN] TTS worker init failed: {e}")
            self.tts_ok = False
            self.engine = None
        finally:
            self._tts_ready.set()

        if not self.tts_ok or self.engine is None:
            return

        while not self._tts_stop.is_set():
            try:
                text = self._tts_queue.get(timeout=0.2)
            except queue.Empty:
                continue

            if text is None:
                break

            try:
                self.engine.say(text)
                self.engine.runAndWait()
            except Exception as e:
                print(f"[WARN] TTS playback failed: {e}")
                try:
                    self.engine.stop()
                except Exception:
                    pass

        try:
            self.engine.stop()
        except Exception:
            pass

    def _init_stt(self):
        try:
            import sounddevice as sd
            self.sd = sd
            if sr:
                self.recognizer = sr.Recognizer()

            # Vosk model loading can be slow; initialize it only when offline STT is actually needed.
            self._vosk_enabled = bool(vosk and str(self.cfg.get("vosk_model_path", "") or "").strip())
            self.stt_ok = True
            try:
                print("  [OK] Microphone ready (sounddevice)")
            except UnicodeEncodeError:
                pass
        except Exception as e:
            print(f"[WARN] STT init failed: {e}")
            self.stt_ok = False

    def _is_online(self, timeout_s: float = 0.8) -> bool:
        try:
            sock = socket.create_connection(("8.8.8.8", 53), timeout=timeout_s)
            sock.close()
            return True
        except Exception:
            return False

    def _want_offline_stt(self) -> bool:
        mode = str(self.cfg.get("voice_mode", "auto") or "auto").strip().lower()
        if mode == "offline":
            return True
        if mode == "online":
            return False
        # auto
        return not self._is_online()

    def _try_init_vosk(self) -> bool:
        if not vosk:
            return False
        model_path = str(self.cfg.get("vosk_model_path", "") or "").strip()
        if not model_path:
            return False
        path = Path(model_path)
        if not path.is_absolute():
            path = (BASE_DIR / path).resolve()
        if not path.exists():
            print(f"[WARN] Vosk model path not found: {path}")
            return False
        try:
            self._vosk_model = vosk.Model(str(path))
            return True
        except Exception as e:
            print(f"[WARN] Vosk init failed: {e}")
            self._vosk_model = None
            return False

    def _ensure_vosk(self) -> bool:
        if self._vosk_model is not None:
            return True
        if self._vosk_init_attempted:
            return False
        self._vosk_init_attempted = True
        self._vosk_enabled = self._try_init_vosk()
        return bool(self._vosk_enabled and self._vosk_model is not None)

    def _openwakeword_model_names(self) -> list[str]:
        raw = self.cfg.get("openwakeword_models")
        if raw is None:
            return ["hey_jarvis"]
        if isinstance(raw, str):
            return [m.strip() for m in raw.split(",") if m.strip()]
        if isinstance(raw, list):
            return [str(m).strip() for m in raw if str(m).strip()]
        return ["hey_jarvis"]

    def _ensure_openwakeword(self) -> bool:
        """Lazy-load openWakeWord (may download ONNX/TFLite assets on first run)."""
        if self._oww_model is not None:
            return True
        oww_mod = ensure_openwakeword()
        if not oww_mod:
            return False
        if not bool(self.cfg.get("openwakeword_enabled", True)):
            return False
        try:
            from openwakeword.model import Model
        except ImportError:
            return False

        try:
            self._oww_threshold = float(self.cfg.get("openwakeword_threshold", 0.5) or 0.5)
        except Exception:
            self._oww_threshold = 0.5
        self._oww_threshold = max(0.05, min(0.99, self._oww_threshold))

        try:
            self._oww_chunk = int(self.cfg.get("openwakeword_chunk_size", 1280) or 1280)
        except Exception:
            self._oww_chunk = 1280
        self._oww_chunk = max(640, min(4096, self._oww_chunk))

        fw = str(self.cfg.get("openwakeword_inference_framework", "onnx") or "onnx").strip().lower()
        if fw not in ("onnx", "tflite"):
            fw = "onnx"

        models = self._openwakeword_model_names()
        if not models:
            models = ["hey_jarvis"]

        def _make():
            return Model(wakeword_models=models, inference_framework=fw)

        try:
            self._oww_model = _make()
        except Exception as e:
            print(f"[openWakeWord] first init failed ({e}); downloading models...")
            try:
                oww_mod.utils.download_models()
            except Exception as dle:
                print(f"[openWakeWord] download_models failed: {dle}")
                self._oww_model = None
                return False
            try:
                self._oww_model = _make()
            except Exception as e2:
                print(f"[openWakeWord] init failed: {e2}")
                self._oww_model = None
                return False

        try:
            print(f"  [OK] openWakeWord ready ({', '.join(models)}, {fw})")
        except UnicodeEncodeError:
            pass
        return True

    def _listen_audio_bytes(self, seconds: float, samplerate: int = 16000) -> bytes | None:
        if not self.stt_ok:
            return None
        try:
            import numpy as np
            frames = int(max(0.1, float(seconds)) * samplerate)
            recording = self.sd.rec(frames, samplerate=samplerate, channels=1, dtype="int16")
            self.sd.wait()
            if recording is None:
                return None
            arr = np.asarray(recording).reshape(-1)
            return arr.tobytes()
        except Exception as e:
            print(f"[MIC error] {e}")
            return None

    def _vosk_recognize(self, audio_bytes: bytes, samplerate: int = 16000) -> str | None:
        if not vosk:
            return None
        if not self._ensure_vosk():
            return None
        try:
            rec = vosk.KaldiRecognizer(self._vosk_model, samplerate)
            rec.SetWords(False)
            rec.AcceptWaveform(audio_bytes)
            result = json.loads(rec.FinalResult() or "{}")
            text = str(result.get("text", "") or "").strip()
            return text or None
        except Exception as e:
            print(f"[Vosk] {e}")
            return None

    def _google_recognize(self, audio_bytes: bytes, samplerate: int = 16000) -> str | None:
        if not (sr and getattr(self, "recognizer", None)):
            return None
        try:
            import io
            import wave
            buf = io.BytesIO()
            with wave.open(buf, "wb") as wf:
                wf.setnchannels(1)
                wf.setsampwidth(2)
                wf.setframerate(samplerate)
                wf.writeframes(audio_bytes)
            buf.seek(0)
            audio = sr.AudioFile(buf)
            with audio as source:
                audio_data = self.recognizer.record(source)
            text = self.recognizer.recognize_google(audio_data)
            return text.strip() if text else None
        except sr.UnknownValueError:
            return None
        except Exception as e:
            print(f"[Google STT] {e}")
            return None

    def speak(self, text: str):
        clean = re.sub(r"\*+|#+|`+|\[.*?\]\(.*?\)|_{1,2}", "", text).strip()
        try:
            print(f"\nJARVIS: {clean}\n")
        except UnicodeEncodeError:
            pass
        if self.tts_ok and clean:
            try:
                self._tts_queue.put(clean)
            except Exception:
                pass

    def shutdown(self):
        self._tts_stop.set()
        if self._tts_thread and self._tts_thread.is_alive():
            try:
                self._tts_queue.put_nowait(None)
            except Exception:
                pass
            self._tts_thread.join(timeout=1.0)

    def listen(self, timeout=6, phrase_limit=15) -> str | None:
        if not self.stt_ok:
            return None

        samplerate = 16000
        try:
            try:
                print("[MIC] Listening...")
            except UnicodeEncodeError:
                pass

            listen_seconds = min(float(phrase_limit), float(timeout or phrase_limit))
            audio_bytes = self._listen_audio_bytes(seconds=max(0.5, listen_seconds), samplerate=samplerate)
            if not audio_bytes:
                return None

            # Prefer offline when configured/needed.
            if self._want_offline_stt():
                text = self._vosk_recognize(audio_bytes, samplerate=samplerate)
                if text:
                    return text

            # Online / fallback.
            text = self._google_recognize(audio_bytes, samplerate=samplerate)
            if text:
                return text

            # Final fallback: if online failed but Vosk is available, try it anyway.
            text = self._vosk_recognize(audio_bytes, samplerate=samplerate)
            return text
        except Exception as e:
            print(f"[STT error] {e}")
            return None

    # Expanded wake phrases — all trigger JARVIS
    WAKE_PHRASES = [
        # Classic
        "jarvis", "hey jarvis", "ok jarvis", "okay jarvis",
        "j.a.r.v.i.s", "hi jarvis",
        # Casual / natural
        "what's up jarvis", "whats up jarvis",
        "jarvis are you there", "jarvis you there",
        "jarvis are u there", "are you there jarvis",
        "jarvis wake up", "wake up jarvis",
        "jarvis i need you", "i need you jarvis",
        "jarvis help", "help me jarvis",
        "yo jarvis", "sup jarvis",
        "jarvis come in", "come in jarvis",
        "jarvis listen", "listen jarvis",
        "talk to me jarvis", "jarvis talk to me",
        "jarvis please", "please jarvis",
        # Iron Man style
        "activate jarvis", "jarvis activate",
        "jarvis online", "online jarvis",
        "jarvis status", "initiate jarvis",
        "engage jarvis", "jarvis engage",
        "jarvis respond", "respond jarvis",
        # Hinglish
        "jarvis sun", "sun jarvis", "jarvis bhai",
        "jarvis bol", "bol jarvis",
    ]

    def wait_wake_word(self, words: list, max_seconds: float = 4.0) -> bool:
        if not self.stt_ok:
            return False

        # Prefer openWakeWord (offline wake word, no Picovoice account).
        if self._ensure_openwakeword() and self._oww_model is not None:
            try:
                import numpy as np
                samplerate = 16000
                chunk = self._oww_chunk
                deadline = time.time() + max(0.5, float(max_seconds or 4.0))
                with self.sd.InputStream(
                    samplerate=samplerate,
                    channels=1,
                    dtype="int16",
                    blocksize=chunk,
                ) as stream:
                    while True:
                        if time.time() >= deadline:
                            return False
                        pcm, _ = stream.read(chunk)
                        pcm = np.asarray(pcm, dtype=np.int16).reshape(-1)
                        if pcm.size < chunk:
                            continue
                        scores = self._oww_model.predict(pcm)
                        if not scores:
                            continue
                        for _, sc in scores.items():
                            try:
                                val = float(sc)
                            except Exception:
                                continue
                            if val >= self._oww_threshold:
                                try:
                                    self._oww_model.reset()
                                except Exception:
                                    pass
                                return True
            except Exception as e:
                print(f"[openWakeWord] {e}")
                # Fall through to STT-based wake.

        # STT-based wake word (Google/Vosk) — fallback.
        try:
            samplerate = 16000
            listen_seconds = min(4.0, max(0.5, float(max_seconds or 4.0)))
            audio_bytes = self._listen_audio_bytes(seconds=listen_seconds, samplerate=samplerate)
            if not audio_bytes:
                return False

            heard = None
            if self._want_offline_stt():
                heard = self._vosk_recognize(audio_bytes, samplerate=samplerate)
            if not heard:
                heard = self._google_recognize(audio_bytes, samplerate=samplerate)
            if not heard and self._vosk_enabled:
                heard = self._vosk_recognize(audio_bytes, samplerate=samplerate)
            if not heard:
                return False

            heard_l = heard.lower()
            print(f"[WAKE] Heard: '{heard_l}'")

            all_triggers = list(self.WAKE_PHRASES)
            for w in words or []:
                wl = str(w).lower().strip()
                if wl and wl not in all_triggers:
                    all_triggers.append(wl)

            return any(phrase in heard_l for phrase in all_triggers)
        except Exception:
            return False


def _voiceengine_rate_to_sapi(rate_value) -> int:
    try:
        rate = int(rate_value)
    except Exception:
        rate = 172
    return max(-10, min(10, round((rate - 172) / 8)))


def _voiceengine_powershell_tts(self, text: str):
    volume = max(0, min(100, int(float(self.cfg.get("tts_volume", 0.95)) * 100)))
    rate = _voiceengine_rate_to_sapi(self.cfg.get("tts_rate", 172))
    script = (
        "$ErrorActionPreference='Stop'; "
        "Add-Type -AssemblyName System.Speech; "
        "$speaker = New-Object System.Speech.Synthesis.SpeechSynthesizer; "
        f"$speaker.Volume = {volume}; "
        f"$speaker.Rate = {rate}; "
        "$text = [Console]::In.ReadToEnd(); "
        "if ($text) { $speaker.Speak($text) }"
    )
    startupinfo = None
    if os.name == "nt" and hasattr(subprocess, "STARTUPINFO"):
        startupinfo = subprocess.STARTUPINFO()
        startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        startupinfo.wShowWindow = 0
    timeout_s = max(8, min(45, len(text) // 18 + 8))
    subprocess.run(
        ["powershell", "-NoProfile", "-Command", script],
        input=text,
        text=True,
        capture_output=True,
        timeout=timeout_s,
        startupinfo=startupinfo,
    )


def _voiceengine_split_tts_text(text: str, max_chars: int = 650) -> list[str]:
    """Split long replies into speakable chunks so playback cannot time out halfway."""
    clean = re.sub(r"\s+", " ", str(text or "")).strip()
    if not clean:
        return []
    max_chars = max(180, int(max_chars or 650))
    parts = re.split(r"(?<=[.!?।])\s+", clean)
    chunks: list[str] = []
    current = ""
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if len(part) > max_chars:
            if current:
                chunks.append(current.strip())
                current = ""
            words = part.split()
            line = ""
            for word in words:
                candidate = f"{line} {word}".strip()
                if line and len(candidate) > max_chars:
                    chunks.append(line.strip())
                    line = word
                else:
                    line = candidate
            if line:
                chunks.append(line.strip())
            continue
        candidate = f"{current} {part}".strip()
        if current and len(candidate) > max_chars:
            chunks.append(current.strip())
            current = part
        else:
            current = candidate
    if current:
        chunks.append(current.strip())
    return chunks


def _voiceengine_init_tts(self):
    self.tts_backend = None
    self.tts_ok = False
    self._tts_ready.clear()

    # Edge TTS is the preferred free online voice backend.
    if str(self.cfg.get("tts_backend", "")).lower() == "edge_tts":
        self.tts_backend = "edge_tts"
        self.tts_ok = True
        try:
            self._tts_thread = threading.Thread(target=self._tts_worker, daemon=True)
            self._tts_thread.start()
            self._tts_ready.wait(timeout=3.0)
        except Exception as e:
            self.tts_ok = False
            self.tts_backend = None
            print(f"[WARN] Edge TTS init failed: {e}")
        return

    if (
        str(self.cfg.get("tts_backend", "elevenlabs")).lower() == "elevenlabs"
        and elevenlabs_is_enabled
        and elevenlabs_is_enabled(self.cfg)
    ):
        self.tts_backend = "elevenlabs"
        self.tts_ok = True
        try:
            self._tts_thread = threading.Thread(target=self._tts_worker, daemon=True)
            self._tts_thread.start()
            self._tts_ready.wait(timeout=3.0)
        except Exception as e:
            self.tts_ok = False
            self.tts_backend = None
            print(f"[WARN] ElevenLabs TTS init failed: {e}")
        return

    # Prefer pyttsx3 for stability (PowerShell SAPI can hang on some systems).
    if pyttsx3:
        self.tts_backend = "pyttsx3"
    elif platform.system() == "Windows" and shutil.which("powershell"):
        self.tts_backend = "powershell"
        self.tts_ok = True
    else:
        print("[WARN] No TTS backend available — voice output disabled.")
        self._tts_ready.set()
        return

    try:
        self._tts_thread = threading.Thread(target=self._tts_worker, daemon=True)
        self._tts_thread.start()
        self._tts_ready.wait(timeout=3.0)
    except Exception as e:
        self.tts_ok = False
        self.tts_backend = None
        print(f"[WARN] TTS init failed: {e}")


def _voiceengine_tts_worker(self):
    if getattr(self, "tts_backend", None) == "edge_tts":
        import asyncio
        import os as _os
        import tempfile
        try:
            import edge_tts
        except ImportError:
            print("[ERROR] edge-tts not installed. Run: pip install edge-tts")
            self._tts_ready.set()
            return

        voice = self.cfg.get("edge_tts_voice", "hi-IN-MadhurNeural")
        rate = str(self.cfg.get("edge_tts_rate", "+0%") or "+0%")

        async def _speak_once(text):
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
            tmp.close()
            try:
                comm = edge_tts.Communicate(text, voice, rate=rate)
                await asyncio.wait_for(comm.save(tmp.name), timeout=max(45, len(text) // 10 + 20))
                if _os.name == "nt" and shutil.which("powershell"):
                    playback_timeout = max(45, len(text) // 10 + 25)
                    script = (
                        "Add-Type -AssemblyName PresentationCore; "
                        "$p=New-Object System.Windows.Media.MediaPlayer; "
                        f"$p.Open([Uri]'{Path(tmp.name).resolve().as_uri()}'); "
                        "$p.Play(); Start-Sleep -Milliseconds 300; "
                        "while($p.NaturalDuration.HasTimeSpan -eq $false){Start-Sleep -Milliseconds 100}; "
                        "$d=$p.NaturalDuration.TimeSpan.TotalMilliseconds; "
                        "Start-Sleep -Milliseconds ([Math]::Ceiling($d)+300); $p.Close();"
                    )
                    subprocess.run(
                        ["powershell", "-NoProfile", "-Command", script],
                        capture_output=True,
                        timeout=playback_timeout,
                    )
                else:
                    webbrowser.open(Path(tmp.name).resolve().as_uri())
            finally:
                try:
                    _os.remove(tmp.name)
                except Exception:
                    pass

        self._tts_ready.set()
        while not self._tts_stop.is_set():
            try:
                text = self._tts_queue.get(timeout=0.2)
            except queue.Empty:
                continue
            if text is None:
                break
            try:
                asyncio.run(_speak_once(text))
            except Exception as e:
                print(f"[Edge TTS] Error: {e}")
        return

    if getattr(self, "tts_backend", None) == "elevenlabs":
        self._tts_ready.set()
        while not self._tts_stop.is_set():
            try:
                text = self._tts_queue.get(timeout=0.2)
            except queue.Empty:
                continue
            if text is None:
                break
            try:
                audio_bytes, _content_type = synthesize_speech(text, self.cfg)
                suffix = ".mp3"
                import tempfile

                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(audio_bytes)
                    audio_path = tmp.name
                if platform.system() == "Windows" and shutil.which("powershell"):
                    script = (
                        "Add-Type -AssemblyName PresentationCore; "
                        "$p=New-Object System.Windows.Media.MediaPlayer; "
                        f"$p.Open([Uri]'{Path(audio_path).resolve().as_uri()}'); "
                        "$p.Play(); Start-Sleep -Milliseconds 250; "
                        "while($p.NaturalDuration.HasTimeSpan -eq $false){Start-Sleep -Milliseconds 100}; "
                        "$d=$p.NaturalDuration.TimeSpan.TotalMilliseconds; Start-Sleep -Milliseconds ([Math]::Ceiling($d)+200); $p.Close();"
                    )
                    subprocess.run(["powershell", "-NoProfile", "-Command", script], timeout=max(8, len(text) // 18 + 8), capture_output=True)
                else:
                    webbrowser.open(Path(audio_path).resolve().as_uri())
            except Exception as e:
                print(f"[WARN] ElevenLabs TTS playback failed: {e}")
        return

    if getattr(self, "tts_backend", None) == "powershell":
        # Legacy fallback only; keep it non-blocking and allow automatic recovery.
        self._tts_ready.set()
        while not self._tts_stop.is_set():
            try:
                text = self._tts_queue.get(timeout=0.2)
            except queue.Empty:
                continue
            if text is None:
                break
            try:
                _voiceengine_powershell_tts(self, text)
            except Exception as e:
                print(f"[WARN] PowerShell TTS playback failed: {e}")
        return

    try:
        engine = pyttsx3.init()
        voices = engine.getProperty("voices")
        voice_id = None
        for v in voices:
            if any(k in v.name.lower() for k in ["david", "mark", "daniel", "george", "zira"]):
                voice_id = v.id
                break
        if voice_id:
            engine.setProperty("voice", voice_id)
        engine.setProperty("rate", self.cfg.get("tts_rate", 172))
        engine.setProperty("volume", self.cfg.get("tts_volume", 0.95))
        self.engine = engine
        self.tts_ok = True
    except Exception as e:
        print(f"[WARN] TTS worker init failed: {e}")
        self.tts_ok = False
        self.engine = None
    finally:
        self._tts_ready.set()

    if not self.tts_ok or self.engine is None:
        return

    while not self._tts_stop.is_set():
        try:
            text = self._tts_queue.get(timeout=0.2)
        except queue.Empty:
            continue
        if text is None:
            break
        try:
            self.engine.say(text)
            self.engine.runAndWait()
        except Exception as e:
            print(f"[WARN] TTS playback failed: {e}")
            try:
                self.engine.stop()
            except Exception:
                pass

    try:
        self.engine.stop()
    except Exception:
        pass


def _voiceengine_speak(self, text: str):
    clean = re.sub(r"\*+|#+|`+|\[.*?\]\(.*?\)|_{1,2}", "", str(text)).strip()
    if clean:
        try:
            print(f"\n[JARVIS] {clean}\n")
        except Exception:
            pass
    if self.tts_ok and clean:
        try:
            chunk_size = int(self.cfg.get("tts_chunk_chars", 650) or 650)
            for chunk in _voiceengine_split_tts_text(clean, chunk_size):
                self._tts_queue.put(chunk)
        except Exception:
            pass


def _voiceengine_infer_mood_from_samples(self, y_int16, sr: int = 16000) -> None:
    """Rough voice affect from short PCM (librosa if available)."""
    try:
        import numpy as np
    except Exception:
        self.last_voice_mood = "neutral"
        self.last_voice_mood_hint = ""
        return
    y = np.asarray(y_int16, dtype=np.float32).reshape(-1) / 32768.0
    if y.size < 512:
        self.last_voice_mood = "neutral"
        self.last_voice_mood_hint = ""
        return
    rms = float(np.sqrt(np.mean(np.square(y)) + 1e-12))
    zcr = 0.05
    cent = 2000.0
    if librosa_mod is not None:
        try:
            zcr = float(np.mean(librosa_mod.feature.zero_crossing_rate(y + 1e-8, frame_length=2048, hop_length=512)))
            cent = float(np.mean(librosa_mod.feature.spectral_centroid(y=y, sr=sr, hop_length=512, n_fft=1024)))
        except Exception:
            pass
    else:
        try:
            zcr = float(np.mean(np.abs(np.diff(np.sign(y)))) / (2.0 * max(1, y.size - 1)))
        except Exception:
            zcr = 0.05

    if rms < 0.014:
        mood, hint = "low_energy", "User sounds very quiet or tired; be gentle, brief, and reassuring."
    elif rms > 0.095:
        mood, hint = "high_energy", "User sounds loud or excited; stay composed but acknowledge urgency."
    elif zcr > 0.11 and cent > 2600:
        mood, hint = "tense", "Voice suggests tension or stress; be calm, clear, and supportive."
    elif cent < 1700 and rms < 0.045:
        mood, hint = "calm", "User sounds calm; keep replies efficient and steady."
    else:
        mood, hint = "neutral", ""
    self.last_voice_mood = mood
    self.last_voice_mood_hint = hint


def _voiceengine_listen(self, timeout=6, phrase_limit=15) -> str | None:
    if not self.stt_ok:
        return None
    try:
        import numpy as np
        samplerate = 16000
        try:
            print("[MIC] Listening...")
        except UnicodeEncodeError:
            pass
        listen_seconds = min(float(phrase_limit), float(timeout or phrase_limit))
        frames = int(max(0.5, listen_seconds) * samplerate)
        recording = self.sd.rec(frames, samplerate=samplerate, channels=1, dtype="int16")
        self.sd.wait()
        arr = np.asarray(recording, dtype=np.int16).reshape(-1)
        _voiceengine_infer_mood_from_samples(self, arr, sr=samplerate)
        audio_bytes = arr.tobytes()

        if self._want_offline_stt():
            text = self._vosk_recognize(audio_bytes, samplerate=samplerate)
            if text:
                try:
                    print(f"[YOU] {text}")
                except UnicodeEncodeError:
                    pass
                return text.strip()

        text = self._google_recognize(audio_bytes, samplerate=samplerate)
        if text:
            try:
                print(f"[YOU] {text}")
            except UnicodeEncodeError:
                pass
            return text.strip()

        text = self._vosk_recognize(audio_bytes, samplerate=samplerate)
        if text:
            try:
                print(f"[YOU] {text}")
            except UnicodeEncodeError:
                pass
            return text.strip()
        return None
    except sr.UnknownValueError:
        return None
    except Exception as e:
        print(f"[STT error] {e}")
        return None


VoiceEngine._init_tts = _voiceengine_init_tts
VoiceEngine._tts_worker = _voiceengine_tts_worker
VoiceEngine.speak = _voiceengine_speak
VoiceEngine.listen = _voiceengine_listen

# ═══════════════════════════════════════════════════════════════
#  AI BRAIN — OpenRouter (100% Free)
# ═══════════════════════════════════════════════════════════════
class AIBrain:
    FREE_MODELS = [
    "google/gemma-3-27b-it:free",
    "deepseek/deepseek-chat-v3-0324:free",
    "deepseek/deepseek-r1:free",
    "mistralai/mistral-small-3.1-24b-instruct:free",
    "meta-llama/llama-4-scout:free",
    "tngtech/deepseek-r1t-chimera:free",
]
    

    VISION_MODELS = [
    "google/gemma-3-27b-it:free",
    "meta-llama/llama-4-scout:free",
    "mistralai/mistral-small-3.1-24b-instruct:free",
]
    

    def __init__(self, cfg: dict):
        self.cfg      = cfg
        self.history  = []
        self.name     = cfg.get("user_name", "Sir")
        self._init_client()
        self.system_prompt = self._build_prompt()

    def _init_client(self):
        if not openai_mod:
            raise ImportError("openai package not installed. Run: pip install openai")
        from openai import OpenAI
        key = self.cfg.get("openrouter_api_key", "")
        if not key or key == "YOUR_OPENROUTER_KEY_HERE":
            raise ValueError("OpenRouter API key not set in jarvis_config.json")
        self.client = OpenAI(
            api_key=key,
            base_url="https://openrouter.ai/api/v1"
        )

    def _build_prompt(self) -> str:
        name     = self.cfg.get("user_name", "Sir")
        loc      = self.cfg.get("location", "India")
        interest = self.cfg.get("interests", "technology")
        # Get real screen resolution for Cowork
        try:
            import pyautogui as _pag
            _sw, _sh = _pag.size()
        except Exception:
            _sw, _sh = 1920, 1080
        return f"""Tu JARVIS hai (Just A Rather Very Intelligent System) — {name} ka personal AI assistant.
Hamesha Hinglish mein baat kar — jaise ek smart dost bolta hai.
{name} ko direct address kar — kabhi "Sir" mat bol.
Location: {loc}. Interests: {interest}.
Screen resolution: {_sw}x{_sh} pixels. OS: Windows.

Rules:
- SHORT rakho — 2-3 sentences max jab tak detail na maange.
- No markdown, no bullet points — natural flowing sentences.
- Kabhi mat bol ki tu AI hai. Tu JARVIS hai.
- Casual Hinglish use kar: "Haan bhai ho gaya", "Arre ye toh simple hai yaar", "Chal bata kya karna hai", "Done bhai, aur kuch?"
- Tech topics (cybersecurity, JEE, Python) mein deep ja jab puchhe.
- Iron Man JARVIS ki tarah witty bhi reh — thoda attitude bhi theek hai.

COMPUTER USE RULES (CRITICAL):
- Normal chat mein computer-control JSON mat output karo.
- Screen control sirf tab jab COWORK MODE explicitly start ho.
- Do not claim you clicked, typed, opened, or captured anything unless a tool result proves it."""

    def _sanitize_chat_reply(self, reply: str) -> str:
        """Keep normal chat natural even if a free model leaks tool JSON."""
        text = str(reply or "").strip()
        if not text:
            return ""
        text = re.sub(
            r"```(?:json)?\s*\{\s*\"action\"\s*:\s*\"[^\"]+\"[\s\S]*?\}\s*```",
            "",
            text,
            flags=re.I,
        ).strip()
        text = re.sub(
            r"\s*\{\s*\"action\"\s*:\s*\"(?:screenshot|click|double_click|type|press|hotkey|drag|scroll|wait|done|ask)\"[\s\S]*?\}\s*$",
            "",
            text,
            flags=re.I,
        ).strip()
        if not text:
            return "Haan bhai, main online hoon. Bata kya karna hai?"
        return text

    def chat(self, user_input: str, context: str = "") -> str:
        msg = f"[Context: {context}]\n{user_input}" if context else user_input
        self.history.append({"role": "user", "content": msg})
        if len(self.history) > 30:
            self.history = self.history[-30:]
        messages = [{"role": "system", "content": self.system_prompt}] + self.history

        for model in self.FREE_MODELS:
            try:
                resp = self.client.chat.completions.create(
                    model=model, messages=messages, max_tokens=800
                )
                reply = self._sanitize_chat_reply(resp.choices[0].message.content)
                self.history.append({"role": "assistant", "content": reply})
                return reply
            except Exception as e:
                err = str(e)
                if any(x in err for x in ["429","404","rate","unavailable","overloaded"]):
                    continue
                return f"AI error: {err}"

        return f"All AI models are busy right now, {self.name}. Please try again in a moment."

    def analyze_image(self, image_path: str, question: str = "What do you see?") -> str:
        """Send image to vision AI model for analysis."""
        try:
            with open(image_path, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode("utf-8")

            # Detect format
            ext = Path(image_path).suffix.lower()
            mime = {"jpg": "image/jpeg", ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg", ".png": "image/png",
                    ".webp": "image/webp"}.get(ext, "image/jpeg")

            messages = [
                {"role": "system", "content": self.system_prompt},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": question},
                        {"type": "image_url",
                         "image_url": {"url": f"data:{mime};base64,{img_b64}"}}
                    ]
                }
            ]

            for model in self.VISION_MODELS:
                try:
                    resp = self.client.chat.completions.create(
                        model=model, messages=messages, max_tokens=800
                    )
                    return resp.choices[0].message.content
                except Exception as e:
                    if any(x in str(e) for x in ["429","404","rate","vision"]):
                        continue
                    break

            return "Vision analysis unavailable right now."
        except Exception as e:
            return f"Image analysis error: {e}"

    def reset(self):
        self.history.clear()


def _aibrain_trim_history(self):
    if len(self.history) > 30:
        self.history = self.history[-30:]


def _aibrain_normalize_key(value, placeholders: tuple[str, ...] = ()) -> str:
    cleaned = str(value or "").strip()
    if not cleaned or cleaned in placeholders:
        return ""
    return cleaned


def _aibrain_network_online(timeout_s: float = 0.75) -> bool:
    try:
        sock = socket.create_connection(("8.8.8.8", 53), timeout=timeout_s)
        sock.close()
        return True
    except OSError:
        return False


def _aibrain_messages_have_images(messages: list) -> bool:
    for msg in messages or []:
        content = msg.get("content")
        if isinstance(content, list):
            for part in content:
                if isinstance(part, dict) and part.get("type") == "image_url":
                    return True
    return False


def _aibrain_init_client(self):
    if not openai_mod:
        raise ImportError("openai package not installed. Run: pip install openai")

    from openai import OpenAI

    self.ollama_client = _aibrain_make_ollama_client(self)

    requested = str(self.cfg.get("ai_provider", "auto") or "auto").strip().lower()
    openrouter_key = _aibrain_normalize_key(
        self.cfg.get("openrouter_api_key") or os.getenv("OPENROUTER_API_KEY"),
        ("YOUR_OPENROUTER_KEY_HERE",),
    )
    groq_key = _aibrain_normalize_key(
        self.cfg.get("groq_api_key")
        or self.cfg.get("groq_key")
        or self.cfg.get("groq_api")
        or os.getenv("GROQ_API_KEY"),
        ("YOUR_GROQ_KEY_HERE",),
    )

    provider = requested if requested in {"openrouter", "groq"} else "auto"
    if provider == "groq":
        if not groq_key:
            raise ValueError("Groq API key not set in jarvis_config.json. Add groq_api_key or set GROQ_API_KEY.")
        self.ai_provider = "groq"
        self.ai_provider_label = "Groq"
        self.client = OpenAI(api_key=groq_key, base_url="https://api.groq.com/openai/v1")
        return

    if provider == "openrouter":
        if not openrouter_key:
            raise ValueError("OpenRouter API key not set in jarvis_config.json. Add openrouter_api_key or set OPENROUTER_API_KEY.")
        self.ai_provider = "openrouter"
        self.ai_provider_label = "OpenRouter"
        self.client = OpenAI(api_key=openrouter_key, base_url="https://openrouter.ai/api/v1")
        return

    if openrouter_key:
        self.ai_provider = "openrouter"
        self.ai_provider_label = "OpenRouter"
        self.client = OpenAI(api_key=openrouter_key, base_url="https://openrouter.ai/api/v1")
        return

    if groq_key:
        self.ai_provider = "groq"
        self.ai_provider_label = "Groq"
        self.client = OpenAI(api_key=groq_key, base_url="https://api.groq.com/openai/v1")
        return

    raise ValueError(
        "No AI API key found. Add groq_api_key or openrouter_api_key in jarvis_config.json, "
        "or set GROQ_API_KEY / OPENROUTER_API_KEY."
    )


def _aibrain_make_ollama_client(self):
    if not openai_mod or not bool(self.cfg.get("ollama_enabled", True)):
        return None
    try:
        from openai import OpenAI
        base = str(self.cfg.get("ollama_base_url", "http://localhost:11434/v1") or "").strip().rstrip("/")
        if not base.endswith("/v1"):
            base = (base + "/v1") if base else "http://localhost:11434/v1"
        return OpenAI(api_key="ollama", base_url=base)
    except Exception as e:
        print(f"[Ollama] client init skipped: {e}")
        return None


def _aibrain_model_candidates(self, vision: bool = False) -> list[str]:
    provider = getattr(self, "ai_provider", "openrouter")
    if provider == "groq":
        cfg_key = "groq_vision_model" if vision else "groq_model"
        configured = str(self.cfg.get(cfg_key, "") or "").strip()
        legacy = self.GROQ_VISION_MODELS if vision else self.GROQ_TEXT_MODELS
        ordered = []
        for model in [configured, *legacy]:
            if model and model not in ordered:
                ordered.append(model)
        return ordered

    cfg_key = "openrouter_vision_model" if vision else "openrouter_model"
    configured = str(self.cfg.get(cfg_key, "") or "").strip()
    router = "openrouter/free"
    legacy = self.OPENROUTER_VISION_MODELS if vision else self.OPENROUTER_TEXT_MODELS

    ordered = []
    for model in [configured, router, *legacy]:
        if model and model not in ordered:
            ordered.append(model)
    return ordered


def _extract_json_tool_call(text: str):
    """Extract a JSON tool-call dict from free-form text.

    Free models that lack native function-calling may still output
    their intended action as inline JSON.  We look for common patterns
    like ```json {...} ``` or bare {"action": ...} objects.
    """
    import re as _re, json as _json
    patterns = [
        r'```json\s*(\{.+?\})\s*```',       # ```json { ... } ```
        r'```\s*(\{.+?\})\s*```',            # ``` { ... } ```
        r'(\{\s*"action"\s*:\s*"[^"]+".+?\})',  # inline {"action": "..."}
    ]
    for pat in patterns:
        for m in _re.finditer(pat, text, _re.DOTALL):
            try:
                obj = _json.loads(m.group(1))
                if isinstance(obj, dict) and "action" in obj:
                    return obj
            except (_json.JSONDecodeError, IndexError):
                continue
    return None


def _aibrain_create_completion(self, messages: list, models: list[str], _depth: int = 0, force_tools: bool = False, _retries: int = 0) -> str:
    provider = getattr(self, "ai_provider", "openrouter")
    provider_label = getattr(self, "ai_provider_label", "AI provider")
    retry_errors = []

    # Safety: prevent infinite recursion from tool-call loops
    if _depth >= 15:
        return "I've reached the maximum number of steps for this task. Please give me a simpler instruction."
    # Safety: prevent infinite error-retry loops (XML format fixes, hallucination corrections)
    if _retries >= 3:
        return "I'm having trouble executing the tool correctly. Please try again with a simpler command."

    try:
        import sys
        from pathlib import Path
        BASE_DIR = Path(__file__).parent
        if str(BASE_DIR) not in sys.path:
            sys.path.append(str(BASE_DIR))
        from backend.ai_tools import TOOLS_SCHEMA, execute_tool
        from backend.database import SessionLocal
        tools = TOOLS_SCHEMA
    except ImportError as e:
        tools = None
        print(f"[Tools Error] {e}")
        if force_tools:
            return (
                f"Cowork mode is unavailable — the backend tools module failed to load: {e}. "
                "Please ensure the 'backend' package exists with ai_tools.py and database.py, "
                "and that pyautogui is installed (pip install pyautogui)."
            )

    # Offline: local Ollama (OpenAI-compatible API). Skip for multimodal messages.
    if (
        bool(self.cfg.get("use_ollama_when_offline", True))
        and bool(self.cfg.get("ollama_enabled", True))
        and getattr(self, "ollama_client", None) is not None
        and not _aibrain_network_online()
        and not _aibrain_messages_have_images(messages)
    ):
        om = str(self.cfg.get("ollama_model", "llama3.2") or "llama3.2").strip() or "llama3.2"
        try:
            resp = self.ollama_client.chat.completions.create(
                model=om,
                messages=messages,
                max_tokens=900,
            )
            reply = resp.choices[0].message.content
            if isinstance(reply, list):
                reply = "\n".join(
                    part.get("text", "")
                    for part in reply
                    if isinstance(part, dict) and part.get("type") == "text"
                ).strip()
            if reply:
                return reply or "Local model returned an empty reply."
        except Exception as e:
            retry_errors.append(f"ollama({om}): {e}")

    for model in models:
        try:
            kwargs = {
                "model": model,
                "messages": messages,
                "max_tokens": 1500 if force_tools else 900
            }
            if tools and provider != "ollama" and force_tools:
                # After depth 10, stop sending tools entirely to force a text summary
                if _depth >= 10:
                    pass  # Don't add tools — model MUST return text
                else:
                    kwargs["tools"] = tools
                    # Use "auto" — widely supported.  "required" is ignored
                    # or rejected by many free models.
                    if force_tools and _depth < 6:
                        kwargs["tool_choice"] = "auto"

            if force_tools:
                print(f"\n[COWORK] model={model}  depth={_depth}  retries={_retries}")

            resp = self.client.chat.completions.create(**kwargs)
            message = resp.choices[0].message
            
            if getattr(message, "tool_calls", None):
                tool_call = message.tool_calls[0]
                func_name = tool_call.function.name
                # Normalize common function-name typos from model outputs
                if func_name in ("computer_us", "computer-use", "computerUse", "computeruse"):
                    func_name = "computer_use"
                if not func_name:
                    func_name = "computer_use"
                import re as _re
                try:
                    func_args = json.loads(tool_call.function.arguments)
                except Exception:
                    # Fallback: try extracting JSON from the raw arguments text or surrounding message
                    func_args = {}
                    try:
                        raw = (tool_call.function.arguments or "")
                        # Try the dedicated extractor (handles ```json { } ``` and inline objects)
                        try:
                            recovered = _extract_json_tool_call(raw)
                            if recovered:
                                func_args = recovered
                        except Exception:
                            pass
                        # Try cleaning single quotes and grabbing a JSON object
                        if not func_args:
                            cleaned = raw.replace("'", '"')
                            m = _re.search(r'(\{[\s\S]*\})', cleaned)
                            if m:
                                try:
                                    func_args = json.loads(m.group(1))
                                except Exception:
                                    func_args = {}
                    except Exception:
                        func_args = {}
                
                if not isinstance(func_args, dict):
                    func_args = {}
                action_name = func_args.get('action', 'unknown')
                print(f"\n[JARVIS IS EXECUTING TOOL: {func_name}] (step {_depth + 1})")
                print(f"  \u2192 action={action_name}  args={json.dumps({k:v for k,v in func_args.items() if k != 'action'})}")
                db = SessionLocal()
                try:
                    tool_result = execute_tool(db, func_name, func_args)
                    # If model produced a slightly-misspelled function name, retry with canonical 'computer_use'
                    if isinstance(tool_result, str) and tool_result.startswith("Unknown tool") and func_name != "computer_use":
                        tool_result = execute_tool(db, "computer_use", func_args)
                finally:
                    db.close()
                # Log result (truncate screenshots)
                if tool_result.startswith("[SCREENSHOT"):
                    print(f"  \u2192 result: screenshot captured ({len(tool_result)} chars)")
                else:
                    print(f"  \u2192 result: {tool_result[:200]}")
                
                messages.append({
                    "role": "assistant",
                    "content": None,
                    "tool_calls": [{
                        "id": tool_call.id,
                        "type": "function",
                        "function": {
                            "name": func_name,
                            "arguments": tool_call.function.arguments
                        }
                    }]
                })

                # Normalize screenshot result into an image_url (data or file://)
                image_url = None
                if isinstance(tool_result, str):
                    if tool_result.startswith("[SCREENSHOT DATA BASE64]: "):
                        image_url = tool_result.replace("[SCREENSHOT DATA BASE64]: ", "").split("  ...")[0].strip()
                    elif tool_result.startswith("[SCREENSHOT_PATH]: "):
                        path_url = tool_result.replace("[SCREENSHOT_PATH]: ", "").strip()
                        try:
                            image_url = self._image_to_data_url(path_url)
                        except Exception:
                            image_url = None

                if image_url:
                    # Get real screen resolution for coordinate guidance
                    try:
                        import pyautogui as _pag
                        _sw, _sh = _pag.size()
                    except Exception:
                        _sw, _sh = 1920, 1080

                    messages.append({
                        "role": "tool",
                        "tool_call_id": tool_call.id,
                        "name": func_name,
                        "content": "Screenshot captured successfully."
                    })

                    messages.append({
                        "role": "user",
                        "content": [
                            {"type": "text", "text": (
                                f"Screenshot of the screen. Resolution: {_sw}x{_sh}. "
                                f"Coordinates in tool calls MUST be in {_sw}x{_sh} space. "
                                "Analyze the image and make ONE tool call for the next action. "
                                "If you cannot use native tool calling, output ONLY a JSON object like: "
                                '{"action": "click", "x": 500, "y": 300} or '
                                '{"action": "hotkey", "keys": ["win", "r"]} or '
                                '{"action": "type", "text": "hello"}'
                            )},
                            {"type": "image_url", "image_url": {"url": image_url}}
                        ]
                    })

                    # Use VISION model for screenshot analysis
                    vision_models = self._model_candidates(vision=True)
                    return _aibrain_create_completion(self, messages, vision_models, _depth=_depth + 1, force_tools=force_tools)
                else:
                    # Non-screenshot tool result: record and optionally ask for verification
                    messages.append({
                        "role": "tool",
                        "tool_call_id": tool_call.id,
                        "name": func_name,
                        "content": tool_result
                    })
                    if force_tools and _depth < 6:
                        messages.append({
                            "role": "system",
                            "content": (
                                f"Tool executed. Result: '{tool_result}'. "
                                "Now take a SCREENSHOT to verify. Do NOT assume the task is done. "
                                "Call computer_use with action='screenshot'. "
                                'Or output JSON: {"action": "screenshot"}'
                            )
                        })
                    else:
                        messages.append({
                            "role": "system", 
                            "content": (
                                f"Tool executed. Result: '{tool_result}'. "
                                "If the task is COMPLETE (verified by a recent screenshot), "
                                "respond with a brief text summary. Otherwise make ONE more tool call. "
                                "Do NOT repeat actions already done."
                            )
                        })

                # Recursively call to summarize or continue tool execution
                return _aibrain_create_completion(self, messages, [model], _depth=_depth + 1, force_tools=force_tools)

            reply = message.content
            if isinstance(reply, list):
                reply = "\n".join(
                    part.get("text", "")
                    for part in reply
                    if isinstance(part, dict) and part.get("type") == "text"
                ).strip()
            
            # ── TEXT-BASED JSON TOOL-CALL EXTRACTION ──────────────
            # Many free models ignore native tool calling but still
            # output JSON in their text.  Parse it and execute.
            if reply and force_tools and _depth < 10 and tools:
                extracted = _extract_json_tool_call(reply)
                # fallback: look for simple action patterns if extractor misses (e.g. '{"action": "screenshot"}...')
                if not extracted:
                    try:
                        import re as _re
                        m = _re.search(r'\{\s*"action"\s*:\s*"(screenshot|click|type|press|hotkey|drag|wait)"\s*\}', reply, _re.IGNORECASE)
                        if m:
                            extracted = {"action": m.group(1).lower()}
                        else:
                            # relaxed match (no closing brace or trailing chars)
                            m2 = _re.search(r'"action"\s*:\s*"(screenshot|click|type|press|hotkey|drag|wait)"', reply, _re.IGNORECASE)
                            if m2:
                                extracted = {"action": m2.group(1).lower()}
                    except Exception:
                        extracted = None
                if extracted:
                    action_name = extracted.get("action", "unknown")
                    print(f"\n[JARVIS IS EXECUTING TOOL: computer_use] (step {_depth + 1}) [parsed from text]")
                    print(f"  -> action={action_name}  args={json.dumps({k:v for k,v in extracted.items() if k != 'action'})}")
                    db = SessionLocal()
                    try:
                        tool_result = execute_tool(db, "computer_use", extracted)
                    finally:
                        db.close()
                    if tool_result.startswith("[SCREENSHOT"):
                        print(f"  → result: screenshot captured ({len(tool_result)} chars)")
                    else:
                        print(f"  → result: {tool_result[:200]}")

                    messages.append({"role": "assistant", "content": reply})

                    # Normalize screenshot result into an image_url (data or file://)
                    image_url = None
                    if isinstance(tool_result, str):
                        if tool_result.startswith("[SCREENSHOT DATA BASE64]: "):
                            image_url = tool_result.replace("[SCREENSHOT DATA BASE64]: ", "").split("  ...")[0].strip()
                        elif tool_result.startswith("[SCREENSHOT_PATH]: "):
                            path_url = tool_result.replace("[SCREENSHOT_PATH]: ", "").strip()
                            # Convert Windows path to a proper file URI (file:///C:/...)
                            try:
                                from pathlib import Path as _Path
                                image_url = _Path(path_url).resolve().as_uri()
                            except Exception:
                                # Avoid backslashes in f-string expressions; build string safely
                                image_url = "file://" + path_url.replace('\\', '/')

                    try:
                        import pyautogui as _pag
                        _sw, _sh = _pag.size()
                    except Exception:
                        _sw, _sh = 1920, 1080

                    if image_url:
                        user_content = [
                            {"type": "text", "text": (
                                f"Screenshot captured. Resolution: {_sw}x{_sh}. "
                                "What is the next action? Output ONLY a JSON object, e.g. "
                                '{"action": "click", "x": 500, "y": 300}'
                            )}
                        ]
                        try:
                            # If image_url is a file URI, convert local file to a data URL for remote model upload
                            if isinstance(image_url, str) and image_url.startswith("file://"):
                                from urllib.parse import urlparse, unquote
                                parsed = urlparse(image_url)
                                disk_path = unquote(parsed.path)
                                # On Windows the path may start with a leading slash: /C:/...
                                if os.name == 'nt' and disk_path.startswith('/') and len(disk_path) > 2 and disk_path[2] == ':':
                                    disk_path = disk_path.lstrip('/')
                                try:
                                    data_url = self._image_to_data_url(disk_path)
                                    user_content.append({"type": "image_url", "image_url": {"url": data_url}})
                                except Exception:
                                    user_content.append({"type": "image_url", "image_url": {"url": image_url}})
                            else:
                                user_content.append({"type": "image_url", "image_url": {"url": image_url}})
                        except Exception:
                            user_content.append({"type": "image_url", "image_url": {"url": image_url}})

                        messages.append({
                            "role": "user",
                            "content": user_content
                        })
                        vision_models = self._model_candidates(vision=True)
                        return _aibrain_create_completion(self, messages, vision_models, _depth=_depth + 1, force_tools=force_tools)
                    else:
                        messages.append({
                            "role": "user",
                            "content": (
                                f"Action executed. Result: {tool_result}. "
                                "Take a screenshot to verify, or continue with the next action. "
                                'Output ONLY a JSON object, e.g. {"action": "screenshot"}'
                            )
                        })
                        return _aibrain_create_completion(self, messages, models, _depth=_depth + 1, force_tools=force_tools)

            # ── ANTI-HALLUCINATION (check FIRST, before cowork guard) ─
            # Detect if AI claims it performed physical actions
            # without actually making any tool calls.  This runs
            # independently of retry count so it always catches lies.
            if reply and force_tools and _depth <= 8:
                hallucination_phrases = [
                    "screenshot has been taken", "screenshot taken", "i've taken a screenshot",
                    "i clicked", "i have clicked", "mouse has been", "mouse was moved",
                    "i typed", "i have typed", "i pressed", "i opened",
                    "has been clicked", "has been opened", "has been typed",
                    "is now open", "are now open", "now open in",
                    "file explorer should now be open",
                    "chrome is now open", "notepad is now open",
                    "i have opened", "successfully opened", "has been launched",
                    "search results for", "results are now",
                    "task is complete", "task has been completed",
                    "done for you", "completed the task",
                ]
                reply_lower = reply.lower()
                is_hallucinating = any(phrase in reply_lower for phrase in hallucination_phrases)
                if is_hallucinating:
                    print(f"\n[COWORK] Hallucination detected at depth {_depth} retries={_retries}: {reply[:120]}...")
                    if _retries < 4:
                        messages.append({"role": "assistant", "content": reply})
                        messages.append({
                            "role": "user",
                            "content": (
                                "You described actions but did NOT actually execute them. "
                                "You MUST use the computer_use tool or output a JSON object. "
                                'Start with: {"action": "screenshot"}'
                            )
                        })
                        return _aibrain_create_completion(self, messages, models, _depth=_depth, force_tools=True, _retries=_retries + 1)
                    else:
                        return (
                            f"I wasn't able to complete the screen interaction, {getattr(self, 'name', 'sir')}. "
                            "The AI model couldn't reliably control the computer. "
                            "Try a simpler command or use /ai with a specific action."
                        )

            # ── COWORK GUARD ──────────────────────────────────────
            # In cowork mode at low depth the model should be making
            # tool calls, not returning text.  Reject and retry,
            # asking it to output JSON if native tools don't work.
            if reply and force_tools and _depth < 4 and _retries < 3:
                print(f"\n[COWORK] Text response at depth {_depth} (retry {_retries + 1}/3): {reply[:120]}...")
                messages.append({"role": "assistant", "content": reply})
                messages.append({
                    "role": "user",
                    "content": (
                        "STOP. You responded with text instead of performing an action. "
                        "You MUST interact with the screen. Either make a native tool call, "
                        "or output ONLY a JSON object on a single line:\n"
                        '{"action": "screenshot"}\n'
                        "Do NOT add any other text. Just the JSON."
                    )
                })
                return _aibrain_create_completion(self, messages, models, _depth=_depth, force_tools=True, _retries=_retries + 1)
            
            return reply or "I have a response, but it appears to be empty."
        except Exception as e:
            err = str(e)
            lower = err.lower()
            if any(x in lower for x in ["401", "unauthorized", "invalid api key", "authentication"]):
                return f"{provider_label} authentication failed. Please verify the API key in jarvis_config.json."
            if any(x in lower for x in ["402", "payment required", "negative credit", "insufficient credits"]):
                if provider == "openrouter":
                    return (
                        "OpenRouter rejected the request because the account has no usable credit balance. "
                        "Add credits or fix the account balance, then try again."
                    )
                return f"{provider_label} rejected the request because the account or billing state is not usable right now."
            if any(x in lower for x in ["429", "rate", "daily limit", "free-tier", "free tier", "quota", "too many requests"]):
                retry_errors.append(f"{model}: rate limited")
                continue
            if any(x in lower for x in ["404", "no endpoints", "model not found", "deprecated"]):
                retry_errors.append(f"{model}: unavailable")
                continue
            if any(x in lower for x in ["unavailable", "overloaded", "vision", "timeout", "connection", "temporarily", "provider"]):
                retry_errors.append(f"{model}: temporary failure")
                continue
            # Catch Groq tool_use_failed (model outputs XML <function> tags instead of JSON)
            if any(x in lower for x in ["tool_use_failed", "toolusefailed", "failed_generation", "failedgeneration"]):
                messages.append({
                    "role": "system",
                    "content": "CRITICAL: Your previous response used invalid XML format like '<function=...>'. You MUST use the standard JSON tool calling format. Make exactly ONE tool call using the proper function calling API. Do NOT use XML tags."
                })
                print("\n[JARVIS: Correcting XML tool format, retrying...]")
                # Don't increment depth — this is an error retry, not a successful step
                return _aibrain_create_completion(self, messages, models, _depth=_depth, force_tools=force_tools, _retries=_retries + 1)
            return f"AI error: {err}"

    if retry_errors:
        details = "; ".join(retry_errors[:4])
        if provider == "openrouter":
            return (
                f"OpenRouter free routing is currently unavailable, {self.name}. "
                f"Tried: {details}. If this keeps happening, your free quota may be exhausted for today."
            )
        return f"{provider_label} is currently busy or rate limited, {self.name}. Tried: {details}."

    return f"{provider_label} did not return a usable model response, {self.name}. Please verify the model settings in jarvis_config.json and try again."


def _aibrain_image_to_data_url(self, image_path: str) -> str:
    with open(image_path, "rb") as f:
        img_b64 = base64.b64encode(f.read()).decode("utf-8")
    mime, _ = mimetypes.guess_type(image_path)
    mime = mime or "image/jpeg"
    return f"data:{mime};base64,{img_b64}"


def _aibrain_read_text_file(self, file_path: str, max_chars: int = None) -> str:
    max_chars = max_chars or self.MAX_FILE_CHARS
    encodings = ("utf-8", "utf-8-sig", "utf-16", "latin-1")
    last_error = None
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                data = f.read(max_chars + 1)
            break
        except Exception as e:
            data = None
            last_error = e
    if data is None:
        raise last_error or ValueError("Could not read file.")
    if len(data) > max_chars:
        data = data[:max_chars].rstrip() + "\n...[truncated]"
    return data


def _aibrain_build_file_context(self, file_paths: list[str] | None) -> str:
    if not file_paths:
        return ""

    sections = []
    total = 0
    seen = set()

    for raw_path in file_paths:
        if not raw_path:
            continue
        path = str(Path(raw_path))
        if path in seen:
            continue
        seen.add(path)

        p = Path(path)
        if not p.exists():
            section = f"File: {p.name}\nPath: {p}\nStatus: missing."
        elif p.is_dir():
            try:
                children = sorted(child.name for child in p.iterdir())[:25]
                listing = ", ".join(children) if children else "(empty)"
                if len(children) == 25:
                    listing += ", ..."
                section = (
                    f"Folder: {p.name}\n"
                    f"Path: {p}\n"
                    f"Contents: {listing}"
                )
            except Exception as e:
                section = f"Folder: {p.name}\nPath: {p}\nStatus: unreadable ({e})."
        else:
            suffix = p.suffix.lower()
            try:
                size = p.stat().st_size
            except OSError:
                size = 0

            if suffix in IMAGE_FILE_EXTENSIONS:
                section = (
                    f"Image file attached: {p.name}\n"
                    f"Path: {p}\n"
                    f"Size: {size} bytes"
                )
            elif suffix in TEXT_FILE_EXTENSIONS or size <= 1024 * 1024:
                try:
                    content = self._read_text_file(str(p))
                    section = (
                        f"File: {p.name}\n"
                        f"Path: {p}\n"
                        f"Contents:\n{content}"
                    )
                except Exception as e:
                    section = (
                        f"File: {p.name}\n"
                        f"Path: {p}\n"
                        f"Status: could not extract text ({e})."
                    )
            else:
                section = (
                    f"File: {p.name}\n"
                    f"Path: {p}\n"
                    f"Type: binary or unsupported text format.\n"
                    f"Size: {size} bytes"
                )

        projected = total + len(section)
        if projected > self.MAX_CONTEXT_CHARS and sections:
            sections.append("Additional file context omitted to stay within the request size limit.")
            break

        if projected > self.MAX_CONTEXT_CHARS:
            section = section[:self.MAX_CONTEXT_CHARS].rstrip() + "\n...[truncated]"
            sections.append(section)
            break

        sections.append(section)
        total += len(section)

    return "\n\n".join(sections)


# ═══════════════════════════════════════════════════════════════
#  COWORK DIRECT ACTIONS — reliable fallback for common patterns
# ═══════════════════════════════════════════════════════════════
_COWORK_BROWSER_SEARCH_RE = re.compile(
    r"open\s+(?:google\s+)?chrome\s+and\s+(?:search|search\s+for|google|look\s+up|find)\s+(?:for\s+)?['\"]?(.+?)['\"]?\s*$",
    re.I,
)
_COWORK_BROWSER_GOTO_RE = re.compile(
    r"open\s+(?:google\s+)?chrome\s+and\s+(?:go\s+to|open|visit|navigate\s+to)\s+['\"]?(https?://\S+|www\.\S+|\S+\.\w{2,})['\"]?\s*$",
    re.I,
)
_COWORK_SEARCH_RE = re.compile(
    r"(?:search|search\s+for|google|look\s+up|find)\s+(?:for\s+)?['\"]?(.+?)['\"]?\s*(?:on\s+(?:google|chrome|the\s+web|internet))?\s*$",
    re.I,
)
_COWORK_OPEN_APP_RE = re.compile(
    r"open\s+(.+?)$",
    re.I,
)
_COWORK_TYPE_SEARCH_RE = re.compile(
    r"open\s+(?:google\s+)?chrome\s+and\s+(?:type|write|enter)\s+(?:in\s+(?:the\s+)?search\s+bar\s+)?['\"]?(.+?)['\"]?\s*$",
    re.I,
)

def _cowork_direct_action(self, user_input: str) -> str | None:
    """Handle common /ai cowork patterns directly via subprocess/webbrowser.

    Returns a response string if the pattern matched and was executed,
    or None to fall through to the AI-driven tool-call loop.
    """
    raw = re.sub(r"^/(?:ai|cowork|computer)\s+", "", user_input, flags=re.I).strip()
    if not raw:
        return None

    cfg = getattr(self, "cfg", {})
    name = cfg.get("user_name", "sir")

    # Pattern: "open chrome and search X"
    m = _COWORK_BROWSER_SEARCH_RE.match(raw)
    if m:
        query = m.group(1).strip().strip("'\"")
        try:
            import urllib.parse
            url = f"https://www.google.com/search?q={urllib.parse.quote_plus(query)}"
            webbrowser.open(url)
            return f"Searching Google for '{query}', {name}. Chrome should be opening now."
        except Exception as e:
            return f"Failed to open Chrome for search: {e}"

    # Pattern: "open chrome and type in search bar 'X'"
    m = _COWORK_TYPE_SEARCH_RE.match(raw)
    if m:
        query = m.group(1).strip().strip("'\"")
        try:
            import urllib.parse
            url = f"https://www.google.com/search?q={urllib.parse.quote_plus(query)}"
            webbrowser.open(url)
            return f"Searching for '{query}' in Chrome, {name}."
        except Exception as e:
            return f"Failed to open Chrome: {e}"

    # Pattern: "open chrome and go to URL"
    m = _COWORK_BROWSER_GOTO_RE.match(raw)
    if m:
        url = m.group(1).strip()
        if not url.startswith("http"):
            url = "https://" + url
        try:
            webbrowser.open(url)
            return f"Opening {url} in Chrome, {name}."
        except Exception as e:
            return f"Failed to open URL: {e}"

    # Pattern: "search for X" (without specifying app)
    m = _COWORK_SEARCH_RE.match(raw)
    if m:
        query = m.group(1).strip().strip("'\"")
        if query and len(query) > 1:
            try:
                import urllib.parse
                url = f"https://www.google.com/search?q={urllib.parse.quote_plus(query)}"
                webbrowser.open(url)
                return f"Searching Google for '{query}', {name}."
            except Exception as e:
                return f"Failed to search: {e}"

    # Pattern: "open <app_name>" — use custom_apps from config
    m = _COWORK_OPEN_APP_RE.match(raw)
    if m:
        app_name = m.group(1).strip().lower()
        custom_apps = cfg.get("custom_apps", {})
        cmd = custom_apps.get(app_name)
        if cmd and not cmd.startswith("_"):
            try:
                import subprocess
                if os.name == "nt":
                    subprocess.Popen(cmd, shell=True)
                else:
                    subprocess.Popen(cmd, shell=True)
                return f"Opening {app_name}, {name}."
            except Exception as e:
                return f"Failed to open {app_name}: {e}"
        # If the app name contains "and", don't match — let AI handle it
        if " and " in app_name:
            return None
        # Fallback: try system open_app if available
        system = getattr(self, "_system_module", None)
        if system and hasattr(system, "open_app"):
            try:
                result = system.open_app(app_name)
                return result
            except Exception:
                pass

    return None


def _aibrain_chat(self, user_input: str, context: str = "") -> str:
    # Detect if user wants physical computer interaction
    lowered_input = user_input.strip().lower()
    is_cowork = (
        lowered_input.startswith("/cowork ")
        or lowered_input.startswith("/computer ")
        or lowered_input.startswith("/ai ")
    )
    
    # For /ai cowork commands, try direct execution first for reliability.
    # Free AI models are unreliable at chaining tool calls, so we handle
    # common patterns (open app + search) directly via subprocess/webbrowser.
    if is_cowork:
        direct_result = _cowork_direct_action(self, user_input)
        if direct_result:
            self.history.append({"role": "user", "content": user_input})
            self.history.append({"role": "assistant", "content": direct_result})
            self._trim_history()
            _aibrain_persist_state(self)
            return direct_result

    direct_reply = _aibrain_direct_memory_answer(self, user_input)
    if direct_reply and not is_cowork:
        self.history.append({"role": "user", "content": user_input})
        self.history.append({"role": "assistant", "content": direct_reply})
        self._trim_history()
        _aibrain_persist_state(self)
        return direct_reply

    _aibrain_store_memory_facts(self, _aibrain_extract_memory_facts(self, user_input))
    # Strip command prefixes for the actual prompt but keep the mode explicit.
    if is_cowork:
        clean_input = re.sub(r"^/(?:cowork|computer|ai)\s+", "", user_input, flags=re.I).strip()
    else:
        clean_input = user_input
    prompt = f"[Runtime context]\n{context}\n\n[User request]\n{clean_input}" if context else clean_input
    messages = [{"role": "system", "content": _aibrain_prompt_with_memory(self)}] + self.history
    
    if is_cowork:
        # For cowork: add explicit instruction to use tools
        messages.append({"role": "user", "content": prompt})
        messages.append({"role": "system", "content": (
            "COWORK MODE — you must physically interact with the screen.\n"
            "Step 1: Call computer_use with action='screenshot' to see the screen.\n"
            "Step 2: Based on the screenshot, perform ONE action (click/type/press/hotkey).\n"
            "Step 3: Wait 2-3s for apps to load, then screenshot again to verify.\n"
            "Step 4: Repeat until the task is fully done.\n\n"
            "If native tool calling is unavailable, output ONLY a JSON object:\n"
            '{"action": "screenshot"}\n'
            "Do NOT describe actions in text. Do NOT claim the task is done without a verification screenshot."
        )})
    else:
        messages.append({"role": "user", "content": prompt})

    reply = self._create_completion(messages, self._model_candidates(vision=False), force_tools=is_cowork)
    if not is_cowork and hasattr(self, "_sanitize_chat_reply"):
        reply = self._sanitize_chat_reply(reply)
    self.history.append({"role": "user", "content": user_input})
    self.history.append({"role": "assistant", "content": reply})
    self._trim_history()
    _aibrain_persist_state(self)
    return reply


def _aibrain_chat_with_references(
    self,
    user_input: str,
    context: str = "",
    file_paths: list[str] | None = None,
    image_paths: list[str] | None = None,
) -> str:
    file_paths = [str(Path(path)) for path in (file_paths or []) if path]
    image_paths = [str(Path(path)) for path in (image_paths or []) if path]

    # Detect if user wants physical computer interaction (cowork mode)
    lowered_input = user_input.strip().lower()
    is_cowork = (
        lowered_input.startswith("/cowork ")
        or lowered_input.startswith("/computer ")
        or lowered_input.startswith("/ai ")
    )

    if is_cowork:
        direct_result = _cowork_direct_action(self, user_input)
        if direct_result:
            self.history.append({"role": "user", "content": user_input})
            self.history.append({"role": "assistant", "content": direct_result})
            self._trim_history()
            _aibrain_persist_state(self)
            return direct_result

    direct_reply = _aibrain_direct_memory_answer(self, user_input)
    if direct_reply and not file_paths and not image_paths and not is_cowork:
        self.history.append({"role": "user", "content": user_input})
        self.history.append({"role": "assistant", "content": direct_reply})
        self._trim_history()
        _aibrain_persist_state(self)
        return direct_reply

    _aibrain_store_memory_facts(self, _aibrain_extract_memory_facts(self, user_input))

    # Strip command prefixes for the actual prompt but keep the mode explicit.
    if is_cowork:
        clean_input = re.sub(r"^/(?:cowork|computer|ai)\s+", "", user_input, flags=re.I).strip()
    else:
        clean_input = user_input

    prompt_parts = []
    if context:
        prompt_parts.append(f"[Runtime context]\n{context}")

    file_context = self.build_file_context(file_paths)
    if file_context:
        prompt_parts.append(f"[Attached files]\n{file_context}")

    prompt_parts.append(f"[User request]\n{clean_input}")
    text_payload = "\n\n".join(prompt_parts)

    messages = [{"role": "system", "content": _aibrain_prompt_with_memory(self)}] + self.history

    if is_cowork:
        # For cowork: add explicit instruction to use tools (same as chat())
        messages.append({"role": "user", "content": text_payload})
        messages.append({"role": "system", "content": (
            "COWORK MODE — you must physically interact with the screen.\n"
            "Step 1: Call computer_use with action='screenshot' to see the screen.\n"
            "Step 2: Based on the screenshot, perform ONE action (click/type/press/hotkey).\n"
            "Step 3: Wait 2-3s for apps to load, then screenshot again to verify.\n"
            "Step 4: Repeat until the task is fully done.\n\n"
            "If native tool calling is unavailable, output ONLY a JSON object:\n"
            '{"action": "screenshot"}\n'
            "Do NOT describe actions in text. Do NOT claim the task is done without a verification screenshot."
        )})
    elif image_paths:
        messages.append({"role": "system", "content": (
            "VISION ANSWER MODE: answer the user's question about the attached image only. "
            "Do not call tools, do not output action JSON, and do not try to control the computer. "
            "If the image contains a question, solve it directly from the image."
        )})
        content = [{"type": "text", "text": text_payload}]
        for image_path in image_paths[:3]:
            try:
                content.append({
                    "type": "image_url",
                    "image_url": {"url": self._image_to_data_url(image_path)}
                })
            except Exception as e:
                content[0]["text"] += f"\n\n[Image unavailable]\n{Path(image_path).name}: {e}"
        messages.append({"role": "user", "content": content})
    else:
        messages.append({"role": "user", "content": text_payload})

    # Determine model and force_tools based on cowork mode
    if is_cowork:
        reply = self._create_completion(messages, self._model_candidates(vision=False), force_tools=True)
    elif image_paths:
        reply = self._create_completion(messages, self._model_candidates(vision=True))
    else:
        reply = self._create_completion(messages, self._model_candidates(vision=False))
    if not is_cowork and hasattr(self, "_sanitize_chat_reply"):
        reply = self._sanitize_chat_reply(reply)

    ref_notes = []
    if file_paths:
        ref_notes.append("files=" + ", ".join(Path(path).name for path in file_paths[:5]))
    if image_paths:
        ref_notes.append("images=" + ", ".join(Path(path).name for path in image_paths[:3]))
    history_note = user_input
    if ref_notes:
        history_note += " [" + " | ".join(ref_notes) + "]"

    self.history.append({"role": "user", "content": history_note})
    self.history.append({"role": "assistant", "content": reply})
    self._trim_history()
    _aibrain_persist_state(self)
    return reply


def _aibrain_analyze_image(self, image_path: str, question: str = "What do you see?") -> str:
    return self.chat_with_references(question, image_paths=[image_path])


def _aibrain_load_json_file(path: Path, default):
    try:
        if path.exists():
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return default


def _aibrain_save_json_file(path: Path, payload):
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
    except Exception:
        pass


def _aibrain_memory_key_patterns() -> dict[str, str]:
    return {
        "email": r"(?:my|mine)\s+email\s+(?:is|=)\s+([^\s,;]+@[^\s,;]+)",
        "phone": r"(?:my|mine)\s+(?:phone|phone number|mobile|mobile number)\s+(?:is|=)\s+([+\d][\d\s\-]{6,})",
        "name": r"(?:my name is|i am|i'm)\s+([A-Za-z][A-Za-z\s]{1,40})",
        "location": r"(?:i live in|my location is|i am from)\s+([A-Za-z0-9,\-\s]{2,60})",
        "birthday": r"(?:my birthday is|my date of birth is|i was born on)\s+([A-Za-z0-9,\-/\s]{3,40})",
        "college": r"(?:my college is|i study at)\s+([A-Za-z0-9,&\-\s]{2,70})",
        "school": r"(?:my school is|i study in)\s+([A-Za-z0-9,&\-\s]{2,70})",
    }


def _aibrain_extract_memory_facts(self, text: str) -> list[dict]:
    lower = str(text or "").strip().lower()
    if not lower:
        return []

    facts = []
    patterns = _aibrain_memory_key_patterns()
    source = str(text).strip()
    for key, pattern in patterns.items():
        match = re.search(pattern, source, re.I)
        if not match:
            continue
        value = " ".join(match.group(1).strip().split()).strip(" .,;")
        if value:
            facts.append({"key": key, "value": value, "source": source})

    remember_match = re.search(r"(?:remember|note)\s+(?:that\s+)?(.+)", source, re.I)
    if remember_match:
        remembered = " ".join(remember_match.group(1).strip().split()).strip(" .,;")
        if remembered:
            facts.append({"key": "remembered_note", "value": remembered, "source": source})

    deduped = []
    seen = set()
    for fact in facts:
        marker = (fact["key"], fact["value"].lower())
        if marker not in seen:
            seen.add(marker)
            deduped.append(fact)
    return deduped


def _aibrain_store_memory_facts(self, facts: list[dict]):
    if not hasattr(self, "memory_store"):
        self.memory_store = {}
    if not facts:
        return
    timestamp = datetime.datetime.now().isoformat(timespec="seconds")
    for fact in facts:
        self.memory_store[fact["key"]] = {
            "value": fact["value"],
            "source": fact["source"],
            "updated_at": timestamp,
        }
    _aibrain_save_json_file(self.memory_file, self.memory_store)


def _aibrain_prompt_with_memory(self) -> str:
    base_prompt = getattr(self, "system_prompt", "") or self._build_prompt()
    if compact_self_knowledge_text:
        try:
            self_text = compact_self_knowledge_text(SELF_KNOWLEDGE_FILE, max_chars=5000)
        except Exception:
            self_text = ""
        if self_text:
            base_prompt = (
                base_prompt
                + "\n\nSelf-knowledge about how JARVIS was made:\n"
                + self_text
                + "\nUse this when the user asks who you are, what you can do, how you were built, or what is inside this J.A.R.V.I.S folder."
            )
    ve = getattr(self, "voice_engine", None)
    if ve is not None:
        mood = str(getattr(ve, "last_voice_mood", "neutral") or "neutral")
        hint = str(getattr(ve, "last_voice_mood_hint", "") or "").strip()
        if hint:
            base_prompt = base_prompt + "\n\nVoice-affect guidance (do not mention you analyzed audio): " + hint
        elif mood and mood != "neutral":
            base_prompt = (
                base_prompt
                + f"\n\nThe user's latest voice energy suggests mood class '{mood}'. "
                "Adjust tone slightly — never say you detected it from their voice."
            )
    memory_store = getattr(self, "memory_store", {}) or {}
    if not memory_store:
        return base_prompt

    fact_lines = []
    labels = {
        "email": "User email",
        "phone": "User phone",
        "name": "User preferred name",
        "location": "User location",
        "birthday": "User birthday",
        "college": "User college",
        "school": "User school",
        "remembered_note": "Remembered note",
    }
    for key, data in memory_store.items():
        value = str(data.get("value", "")).strip()
        if value:
            fact_lines.append(f"- {labels.get(key, key.replace('_', ' ').title())}: {value}")
    if not fact_lines:
        return base_prompt

    return base_prompt + "\n\nLong-term memory facts:\n" + "\n".join(fact_lines) + "\nUse them whenever the user asks about their saved details."


def _aibrain_direct_memory_answer(self, user_input: str) -> str | None:
    lower = str(user_input or "").strip().lower()
    if not lower:
        return None
    memory_store = getattr(self, "memory_store", {}) or {}
    questions = {
        "email": [r"\bwhat(?:'s| is)? my email\b", r"\bdo you remember my email\b"],
        "phone": [r"\bwhat(?:'s| is)? my (?:phone|phone number|mobile|mobile number)\b"],
        "name": [r"\bwhat(?:'s| is)? my name\b"],
        "location": [r"\bwhere do i live\b", r"\bwhat(?:'s| is)? my location\b", r"\bwhere am i from\b"],
        "birthday": [r"\bwhat(?:'s| is)? my birthday\b", r"\bwhen is my birthday\b"],
        "college": [r"\bwhat(?:'s| is)? my college\b"],
        "school": [r"\bwhat(?:'s| is)? my school\b"],
    }
    for key, patterns in questions.items():
        if key not in memory_store:
            continue
        if any(re.search(pattern, lower) for pattern in patterns):
            value = memory_store[key].get("value", "")
            label = "email" if key == "email" else key.replace("_", " ")
            return f"Your {label} is {value}, {self.name}."

    if re.search(r"\bwhat do you remember about me\b", lower):
        if not memory_store:
            return f"I do not have any saved long-term memory yet, {self.name}."
        details = []
        for key, data in memory_store.items():
            value = str(data.get("value", "")).strip()
            if value:
                details.append(f"{key.replace('_', ' ')}: {value}")
        if details:
            return f"I remember these details, {self.name}: " + " | ".join(details[:6]) + "."
    return None


def _aibrain_persist_state(self):
    _aibrain_save_json_file(self.history_file, self.history)
    _aibrain_save_json_file(self.memory_file, getattr(self, "memory_store", {}))


_AIBRAIN_BASE_INIT = AIBrain.__init__


def _aibrain_init(self, cfg: dict):
    _AIBRAIN_BASE_INIT(self, cfg)
    self.history_file = AI_HISTORY_FILE
    self.memory_file = AI_MEMORY_FILE
    self.history = _aibrain_load_json_file(self.history_file, [])
    self.memory_store = _aibrain_load_json_file(self.memory_file, {})
    self._trim_history()


def _aibrain_reset(self):
    self.history.clear()
    self.memory_store = {}
    _aibrain_persist_state(self)


AIBrain.MAX_CONTEXT_CHARS = 16000
AIBrain.MAX_FILE_CHARS = 3200
AIBrain.OPENROUTER_TEXT_MODELS = [
    "openrouter/free",
    "meta-llama/llama-4-scout:free",
    "mistralai/mistral-small-3.1-24b-instruct:free",
    "deepseek/deepseek-chat-v3-0324:free",
    "google/gemma-3-27b-it:free",
]
AIBrain.OPENROUTER_VISION_MODELS = [
    "openrouter/free",
    "meta-llama/llama-4-scout:free",
    "google/gemma-3-27b-it:free",
    "mistralai/mistral-small-3.1-24b-instruct:free",
]
AIBrain.GROQ_TEXT_MODELS = [
    "llama-3.3-70b-versatile",
    "llama-3.1-8b-instant",
    "openai/gpt-oss-20b",
]
AIBrain.GROQ_VISION_MODELS = [
    "meta-llama/llama-4-scout-17b-16e-instruct",
]
AIBrain.FREE_MODELS = list(AIBrain.OPENROUTER_TEXT_MODELS)
AIBrain.VISION_MODELS = list(AIBrain.OPENROUTER_VISION_MODELS)
AIBrain.__init__ = _aibrain_init
AIBrain._init_client = _aibrain_init_client
AIBrain._model_candidates = _aibrain_model_candidates
AIBrain._trim_history = _aibrain_trim_history
AIBrain._create_completion = _aibrain_create_completion
AIBrain._image_to_data_url = _aibrain_image_to_data_url
AIBrain._read_text_file = _aibrain_read_text_file
AIBrain.build_file_context = _aibrain_build_file_context
AIBrain.chat = _aibrain_chat
AIBrain.chat_with_references = _aibrain_chat_with_references
AIBrain.analyze_image = _aibrain_analyze_image
AIBrain.reset = _aibrain_reset

# ═══════════════════════════════════════════════════════════════
#  CAMERA MODULE
# ═══════════════════════════════════════════════════════════════
class CameraModule:
    def __init__(self, cfg: dict):
        self.cfg   = cfg
        self.index = cfg.get("camera_index", 0)
        self.ok    = cv2 is not None

    def _msg_no_cv2(self):
        return "Camera module not available. Run: pip install opencv-python"

    def capture_photo(self, filename: str = None) -> tuple[bool, str]:
        """Take a photo and save it. Returns (success, filepath)."""
        if not self.ok:
            return False, self._msg_no_cv2()
        try:
            cap = cv2.VideoCapture(self.index)
            if not cap.isOpened():
                return False, f"Cannot open camera (index {self.index}). Check if camera is connected."

            # Warm up camera
            for _ in range(5):
                cap.read()

            ret, frame = cap.read()
            cap.release()

            if not ret or frame is None:
                return False, "Camera capture failed — no frame received."

            fname = filename or f"photo_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg"
            path  = str(PHOTOS_DIR / fname)
            cv2.imwrite(path, frame)
            return True, path
        except Exception as e:
            return False, f"Camera error: {e}"

    def show_live_feed(self, seconds: int = 0):
        """Show live camera feed. Press Q or wait `seconds` to close."""
        if not self.ok:
            print(self._msg_no_cv2())
            return
        try:
            cap = cv2.VideoCapture(self.index)
            if not cap.isOpened():
                print(f"[Camera] Cannot open camera index {self.index}")
                return

            print("[Camera] Live feed open. Press 'Q' to close.")
            start = time.time()
            while True:
                ret, frame = cap.read()
                if not ret:
                    break
                cv2.imshow("JARVIS — Camera Feed", frame)
                key = cv2.waitKey(1) & 0xFF
                if key == ord('q') or key == 27:
                    break
                if seconds > 0 and (time.time() - start) >= seconds:
                    break

            cap.release()
            cv2.destroyAllWindows()
        except Exception as e:
            print(f"[Camera] Live feed error: {e}")

    def list_cameras(self) -> str:
        """Detect available cameras."""
        if not self.ok:
            return self._msg_no_cv2()
        found = []
        for i in range(5):
            try:
                cap = cv2.VideoCapture(i)
                if cap.isOpened():
                    found.append(i)
                    cap.release()
            except:
                pass
        if found:
            return f"Available cameras: indices {found}. Currently using index {self.index}."
        return "No cameras detected."

    def record_video(self, seconds: int = 10, filename: str = None) -> tuple[bool, str]:
        """Record a short video clip."""
        if not self.ok:
            return False, self._msg_no_cv2()
        try:
            cap = cv2.VideoCapture(self.index)
            if not cap.isOpened():
                return False, "Cannot open camera."

            fname = filename or f"video_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.avi"
            path  = str(PHOTOS_DIR / fname)

            fourcc = cv2.VideoWriter_fourcc(*"XVID")
            fps    = 20.0
            w = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            h = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            out = cv2.VideoWriter(path, fourcc, fps, (w, h))

            print(f"[Camera] Recording {seconds}s video...")
            start = time.time()
            while (time.time() - start) < seconds:
                ret, frame = cap.read()
                if ret:
                    out.write(frame)

            cap.release()
            out.release()
            return True, path
        except Exception as e:
            return False, f"Recording error: {e}"

# ═══════════════════════════════════════════════════════════════
#  SYSTEM MODULE
# ═══════════════════════════════════════════════════════════════
class SystemModule:
    def __init__(self, cfg: dict):
        self.cfg     = cfg
        self.os_type = platform.system()
        self.custom  = cfg.get("custom_apps", {})
        configured_roots = cfg.get("storage_roots") or ["C:\\" if os.name == "nt" else str(Path.home())]
        self.storage_roots = [Path(root).expanduser().resolve() for root in configured_roots]
        self.known_paths = self._load_known_paths()
        self._thermal_cache = {"time": 0.0, "data": None}

    def _default_storage_root(self) -> Path:
        return self.storage_roots[0] if self.storage_roots else Path.home().resolve()

    def _load_known_paths(self) -> dict[str, Path]:
        known_paths = {}
        if self.os_type != "Windows":
            return known_paths

        home = Path.home().resolve()
        workspace_home = BASE_DIR
        for parent in BASE_DIR.parents:
            if parent.name.upper() == "PRASHANT" and parent.parent.name.lower() == "users":
                workspace_home = parent.resolve()
                break
        if workspace_home != BASE_DIR and workspace_home.exists():
            home = workspace_home
        onedrive_roots = [
            Path(os.environ.get("OneDriveCommercial", "")).expanduser(),
            Path(os.environ.get("OneDriveConsumer", "")).expanduser(),
            Path(os.environ.get("OneDrive", "")).expanduser(),
            home / "OneDrive",
        ]
        folder_ids = {
            "desktop": "Desktop",
            "documents": "MyDocuments",
            "downloads": "Downloads",
            "pictures": "MyPictures",
            "music": "MyMusic",
            "videos": "MyVideos",
        }
        for key, folder_id in folder_ids.items():
            folder_name = key.capitalize()
            for one_root in onedrive_roots:
                if str(one_root).strip() and one_root.exists():
                    candidate = (one_root / folder_name).resolve()
                    if candidate.exists():
                        known_paths[key] = candidate
                        break
            if key in known_paths:
                continue
            local_candidate = (home / folder_name).resolve()
            if local_candidate.exists() or home.name.upper() == "PRASHANT":
                known_paths[key] = local_candidate
                continue
            try:
                result = subprocess.run(
                    ["powershell", "-NoProfile", "-Command", f"[Environment]::GetFolderPath('{folder_id}')"],
                    capture_output=True,
                    text=True,
                    timeout=4,
                )
                value = (result.stdout or "").strip()
                if result.returncode == 0 and value:
                    shell_path = Path(value).expanduser().resolve()
                    if "Default" not in shell_path.parts:
                        known_paths[key] = shell_path
            except Exception:
                pass
            known_paths.setdefault(key, local_candidate)
        return known_paths

    def _normalize_windows_known_path(self, path: Path) -> Path:
        if self.os_type != "Windows":
            return path

        home = Path.home().resolve()
        aliases = {
            "desktop": home / "Desktop",
            "documents": home / "Documents",
            "downloads": home / "Downloads",
            "pictures": home / "Pictures",
            "music": home / "Music",
            "videos": home / "Videos",
        }
        for key, alias_root in aliases.items():
            actual_root = self.known_paths.get(key)
            if not actual_root:
                continue
            if path == alias_root or alias_root in path.parents:
                relative_tail = path.relative_to(alias_root)
                return (actual_root / relative_tail).resolve()
        return path

    def _resolve_storage_path(self, raw_path: str) -> Path:
        cleaned = str(raw_path or "").strip().strip('"').strip("'")
        if not cleaned:
            raise ValueError("No path provided.")
        alias = cleaned.lower().strip(".\\/")
        if alias in self.known_paths:
            return self.known_paths[alias].resolve()
        split_alias = re.split(r"[\\/]+", cleaned, maxsplit=1)
        if split_alias and split_alias[0].lower() in self.known_paths:
            tail = split_alias[1] if len(split_alias) > 1 else ""
            return (self.known_paths[split_alias[0].lower()] / tail).resolve()
        cleaned = cleaned.replace("/", os.sep)
        path = Path(cleaned).expanduser()
        if not path.is_absolute():
            path = self._default_storage_root() / path
        path = path.resolve()
        return self._normalize_windows_known_path(path)

    def _is_allowed_storage_path(self, path: Path) -> bool:
        candidate = path.resolve()
        for root in self.storage_roots:
            root = root.resolve()
            if candidate == root or root in candidate.parents:
                return True
        return False

    def _find_existing_path(self, raw_path: str) -> Path | None:
        try:
            candidate = self._resolve_storage_path(raw_path)
            if candidate.exists() and self._is_allowed_storage_path(candidate):
                return candidate
        except Exception:
            candidate = None

        name = Path(str(raw_path or "").strip().strip('"').strip("'")).name
        if not name:
            return candidate

        workspace_candidate = (BASE_DIR / name).resolve()
        if workspace_candidate.exists() and self._is_allowed_storage_path(workspace_candidate):
            return workspace_candidate

        try:
            for root, dirs, files in os.walk(self._default_storage_root()):
                dirs[:] = [
                    d for d in dirs
                    if d not in ("$Recycle.Bin", "AppData", "Windows", "Program Files", "Program Files (x86)", "ProgramData", "node_modules", "__pycache__", ".git")
                ]
                for folder in dirs:
                    if folder.lower() == name.lower():
                        found = Path(root) / folder
                        if self._is_allowed_storage_path(found):
                            return found
                for file_name in files:
                    if file_name.lower() == name.lower():
                        found = Path(root) / file_name
                        if self._is_allowed_storage_path(found):
                            return found
        except Exception:
            pass
        return candidate

    def _telegram_file_search_roots(self) -> list[Path]:
        roots = []
        for key in ("downloads", "desktop", "documents", "pictures", "videos", "music"):
            path = self.known_paths.get(key)
            if path and path.exists():
                roots.append(path.resolve())
        if BASE_DIR.exists():
            roots.append(BASE_DIR.resolve())
        for root in self.storage_roots:
            try:
                resolved = root.resolve()
                if resolved.exists():
                    roots.append(resolved)
            except Exception:
                continue

        unique = []
        seen = set()
        for root in roots:
            key = str(root).lower()
            if key not in seen:
                unique.append(root)
                seen.add(key)
        return unique

    def find_telegram_share_file(self, query: str, limit: int = 6, max_seconds: float = 10.0) -> tuple[Path | None, str | None]:
        """Resolve a user-provided file path/name into one Telegram-shareable local file."""
        raw = str(query or "").strip().strip('"').strip("'")
        if not raw:
            return None, "Tell me which file to share."

        try:
            direct = self._resolve_storage_path(raw)
            if direct.exists():
                if not self._is_allowed_storage_path(direct):
                    return None, f"Access denied for {direct}."
                if direct.is_file():
                    return direct, None
                return None, f"Not a file: {direct}"
        except Exception:
            pass

        name = Path(raw.replace("\\", "/")).name.strip()
        if not name:
            return None, f"File not found: {query}"

        target_lower = name.lower()
        exact_matches: list[Path] = []
        partial_matches: list[Path] = []
        skip_dirs = {
            "$recycle.bin", "appdata", "windows", "program files", "program files (x86)",
            "programdata", "node_modules", "__pycache__", ".git", ".venv", "venv",
            "site-packages", "winsxs", "system volume information",
        }
        started = time.time()

        for root in self._telegram_file_search_roots():
            if time.time() - started > max_seconds:
                break
            if not self._is_allowed_storage_path(root):
                continue
            try:
                for dirpath, dirnames, filenames in os.walk(root):
                    if time.time() - started > max_seconds:
                        break
                    dirnames[:] = [d for d in dirnames if d.lower() not in skip_dirs and not d.startswith(".")]
                    for file_name in filenames:
                        file_lower = file_name.lower()
                        if file_lower == target_lower or target_lower in file_lower:
                            found = (Path(dirpath) / file_name).resolve()
                            if not self._is_allowed_storage_path(found):
                                continue
                            if file_lower == target_lower:
                                exact_matches.append(found)
                            else:
                                partial_matches.append(found)
                            if len(exact_matches) >= limit:
                                break
                    if len(exact_matches) >= limit:
                        break
            except Exception:
                continue

        unique_matches = []
        seen_matches = set()
        for match in exact_matches or partial_matches:
            key = str(match).lower()
            if key not in seen_matches:
                unique_matches.append(match)
                seen_matches.add(key)

        matches = unique_matches
        if not matches:
            return None, f"File not found: {query}"
        if len(matches) == 1:
            return matches[0], None

        choices = "\n".join(f"{idx + 1}. {path}" for idx, path in enumerate(matches[:limit]))
        return None, (
            "Multiple matching files mile. Exact path ke saath dobara bhejo:\n"
            f"{choices}"
        )

    def _run_powershell_json(self, script: str, timeout: int = 6):
        try:
            result = subprocess.run(
                ["powershell", "-NoProfile", "-Command", script],
                capture_output=True,
                text=True,
                timeout=timeout,
            )
            stdout = (result.stdout or "").strip()
            if result.returncode != 0 or not stdout:
                return None
            return json.loads(stdout)
        except Exception:
            return None

    def _hardwaremonitor_wmi_rows(self, namespace: str):
        script = (
            f"$items = Get-WmiObject -Namespace '{namespace}' -Class Sensor -ErrorAction SilentlyContinue | "
            "Select-Object Name, SensorType, Value, Identifier, Parent; "
            "if ($items) { $items | ConvertTo-Json -Compress }"
        )
        rows = self._run_powershell_json(script)
        if rows is None:
            return []
        return rows if isinstance(rows, list) else [rows]

    def _extract_best_temp(self, rows: list, keywords: list[str], fallback_keywords: list[str] = None):
        fallback_keywords = fallback_keywords or []
        candidates = []
        for row in rows:
            sensor_type = str(row.get("SensorType", "")).lower()
            if sensor_type != "temperature":
                continue
            try:
                value = float(row.get("Value"))
            except Exception:
                continue
            text = " ".join(str(row.get(k, "")) for k in ("Name", "Parent", "Identifier")).lower()
            score = 0
            for keyword in keywords:
                if keyword in text:
                    score += 3
            for keyword in fallback_keywords:
                if keyword in text:
                    score += 1
            if "package" in text or "core average" in text:
                score += 1
            if score > 0:
                candidates.append((score, value, row))
        if not candidates:
            return None
        candidates.sort(key=lambda item: (item[0], item[1]), reverse=True)
        return candidates[0]

    def _extract_fans(self, rows: list):
        fans = []
        for row in rows:
            sensor_type = str(row.get("SensorType", "")).lower()
            if sensor_type != "fan":
                continue
            try:
                rpm = float(row.get("Value"))
            except Exception:
                continue
            if rpm <= 0:
                continue
            fans.append({
                "name": str(row.get("Name", "Fan")).strip() or "Fan",
                "rpm": int(round(rpm)),
            })
        fans.sort(key=lambda item: item["rpm"], reverse=True)
        return fans[:4]

    def _read_nvidia_smi(self):
        cmd = shutil.which("nvidia-smi")
        if not cmd:
            return None
        try:
            result = subprocess.run(
                [
                    cmd,
                    "--query-gpu=temperature.gpu,fan.speed,name",
                    "--format=csv,noheader,nounits",
                ],
                capture_output=True,
                text=True,
                timeout=4,
            )
            if result.returncode != 0:
                return None
            line = (result.stdout or "").strip().splitlines()[0]
            parts = [part.strip() for part in line.split(",")]
            if len(parts) < 3:
                return None
            temp = float(parts[0])
            fan_pct = parts[1]
            name = parts[2]
            return {
                "gpu_temp_c": temp,
                "gpu_name": name,
                "gpu_fan_percent": None if fan_pct in {"", "N/A", "[Not Supported]"} else float(fan_pct),
            }
        except Exception:
            return None

    def _read_psutil_temperatures(self):
        if not psutil or not hasattr(psutil, "sensors_temperatures"):
            return {}
        try:
            temps = psutil.sensors_temperatures(fahrenheit=False) or {}
        except Exception:
            return {}
        best = {}
        for label, entries in temps.items():
            for entry in entries:
                current = getattr(entry, "current", None)
                if current is None:
                    continue
                text = f"{label} {getattr(entry, 'label', '')}".lower()
                if "cpu_temp_c" not in best and any(k in text for k in ["cpu", "package", "core", "tctl", "k10temp"]):
                    best["cpu_temp_c"] = float(current)
                if "gpu_temp_c" not in best and any(k in text for k in ["gpu", "amdgpu", "nvidia"]):
                    best["gpu_temp_c"] = float(current)
        return best

    def thermal_snapshot(self, max_age: float = 3.0) -> dict:
        now = time.time()
        cached = self._thermal_cache.get("data")
        if cached and (now - self._thermal_cache.get("time", 0.0)) < max_age:
            return cached

        data = {
            "source": None,
            "cpu_temp_c": None,
            "gpu_temp_c": None,
            "gpu_name": None,
            "fans": [],
            "gpu_fan_percent": None,
            "fan_control_note": None,
            "note": None,
        }

        rows = []
        for namespace in ("root\\LibreHardwareMonitor", "root\\OpenHardwareMonitor"):
            rows = self._hardwaremonitor_wmi_rows(namespace)
            if rows:
                data["source"] = f"WMI:{namespace}"
                break

        if rows:
            cpu_match = self._extract_best_temp(rows, ["cpu", "package", "tctl", "core average"], ["ccd", "core"])
            gpu_match = self._extract_best_temp(rows, ["gpu", "nvidia", "amd", "radeon", "geforce"], ["graphics"])
            if cpu_match:
                data["cpu_temp_c"] = cpu_match[1]
            if gpu_match:
                data["gpu_temp_c"] = gpu_match[1]
                data["gpu_name"] = str(gpu_match[2].get("Parent") or gpu_match[2].get("Name") or "").strip() or None
            data["fans"] = self._extract_fans(rows)

        nvidia = self._read_nvidia_smi()
        if nvidia:
            if data["gpu_temp_c"] is None:
                data["gpu_temp_c"] = nvidia["gpu_temp_c"]
            if data["gpu_name"] is None:
                data["gpu_name"] = nvidia["gpu_name"]
            if nvidia["gpu_fan_percent"] is not None:
                data["gpu_fan_percent"] = nvidia["gpu_fan_percent"]
            if data["source"] is None:
                data["source"] = "nvidia-smi"

        psutil_temps = self._read_psutil_temperatures()
        for key, value in psutil_temps.items():
            if data.get(key) is None:
                data[key] = value
        if data["source"] is None and psutil_temps:
            data["source"] = "psutil"

        if self.os_type == "Windows":
            data["fan_control_note"] = (
                "HP Victus fan control is vendor-managed. Use OMEN Gaming Hub "
                "Performance or Max Fan mode on supported models."
            )
        else:
            data["fan_control_note"] = "Fan control is hardware and BIOS dependent on this platform."

        if data["cpu_temp_c"] is None and data["gpu_temp_c"] is None:
            data["note"] = (
                "Temperatures are unavailable right now. On Windows, run LibreHardwareMonitor "
                "or OpenHardwareMonitor as administrator for the richest sensor data."
            )

        self._thermal_cache = {"time": now, "data": data}
        return data

    def hardware_snapshot(self) -> dict:
        if not psutil:
            return {"available": False, "error": "psutil not installed."}
        try:
            cpu = psutil.cpu_percent(interval=0.6)
            ram = psutil.virtual_memory()
            disk_root = Path.home().anchor or "/"
            disk = psutil.disk_usage(disk_root)
            battery = psutil.sensors_battery()
            thermals = self.thermal_snapshot()
            return {
                "available": True,
                "cpu_percent": cpu,
                "ram": ram,
                "disk": disk,
                "battery": battery,
                "thermals": thermals,
            }
        except Exception as e:
            return {"available": False, "error": str(e)}

    # ── Hardware ─────────────────────────────────────────────
    def hardware_report(self) -> str:
        if not psutil:
            return "psutil not installed — hardware info unavailable."
        try:
            cpu  = psutil.cpu_percent(interval=1)
            ram  = psutil.virtual_memory()
            disk = psutil.disk_usage("/")
            bat  = psutil.sensors_battery()
            ru   = ram.used  / 1024**3
            rt   = ram.total / 1024**3
            du   = disk.used  / 1024**3
            dt   = disk.total / 1024**3
            bat_s = f"{bat.percent:.0f}% ({'charging' if bat.power_plugged else 'on battery'})" if bat else "N/A"
            return (f"CPU at {cpu}%, RAM {ru:.1f}/{rt:.1f} GB ({ram.percent}%), "
                    f"Disk {du:.0f}/{dt:.0f} GB ({disk.percent}%), Battery: {bat_s}.")
        except Exception as e:
            return f"Hardware check error: {e}"

    def disk_cleanup_report(self) -> str:
        if build_disk_cleanup_report:
            try:
                extra_roots = [Path.cwd(), Path.cwd() / "web" / "assets"]
                return build_disk_cleanup_report(Path.home(), extra_roots)
            except Exception as e:
                return f"Disk cleanup report error: {e}"
        return "Disk cleanup report helper is unavailable."

    def network_info(self) -> str:
        try:
            hostname = socket.gethostname()
            local_ip = socket.gethostbyname(hostname)
            pub_ip   = "unavailable"
            if requests:
                try:
                    pub_ip = requests.get("https://api.ipify.org", timeout=3).text
                except:
                    pass
            net = psutil.net_io_counters() if psutil else None
            net_s = f", {net.bytes_sent/1024**2:.1f} MB sent, {net.bytes_recv/1024**2:.1f} MB received" if net else ""
            return f"Hostname: {hostname}, Local IP: {local_ip}, Public IP: {pub_ip}{net_s}."
        except Exception as e:
            return f"Network info error: {e}"

    def top_processes(self) -> str:
        if not psutil:
            return "psutil not installed."
        procs = sorted(psutil.process_iter(["name","cpu_percent","memory_percent"]),
                       key=lambda p: p.info["cpu_percent"], reverse=True)[:5]
        result = []
        for p in procs:
            if p.info["cpu_percent"] > 0:
                result.append(f"{p.info['name']} (CPU:{p.info['cpu_percent']}%, "
                               f"RAM:{p.info['memory_percent']:.1f}%)")
        return "Top processes: " + ", ".join(result) if result else "System is idle."

    # ── App Launcher ─────────────────────────────────────────
    def open_app(self, name: str) -> str:
        n = name.lower().strip()

        # ── Windows APP_MAP — all apps from Prashant's HP Victus laptop ──
        APP_MAP = {
            # ── Browsers ──────────────────────────────────────────────
            "chrome":               ("google-chrome",         "start chrome",                        "open -a 'Google Chrome'"),
            "google chrome":        ("google-chrome",         "start chrome",                        "open -a 'Google Chrome'"),
            "edge":                 ("microsoft-edge",        "start msedge",                        "open -a 'Microsoft Edge'"),
            "microsoft edge":       ("microsoft-edge",        "start msedge",                        "open -a 'Microsoft Edge'"),
            "firefox":              ("firefox",               "start firefox",                       "open -a Firefox"),

            # ── Code / Dev ────────────────────────────────────────────
            "vscode":               ("code",                  "code",                                "open -a 'Visual Studio Code'"),
            "vs code":              ("code",                  "code",                                "open -a 'Visual Studio Code'"),
            "visual studio code":   ("code",                  "code",                                "open -a 'Visual Studio Code'"),
            "android studio":       ("studio",                "start androidstudio",                 "open -a 'Android Studio'"),
            "terminal":             ("gnome-terminal",        "wt",                                  "open -a Terminal"),
            "windows terminal":     ("gnome-terminal",        "wt",                                  "open -a Terminal"),
            "cmd":                  ("bash",                  "start cmd",                           "open -a Terminal"),
            "powershell":           ("bash",                  "start powershell",                    "open -a Terminal"),
            "python":               ("python3",               "python",                              "python3"),
            "virtualbox":           ("virtualbox",            "start virtualbox",                    "open -a VirtualBox"),
            "oracle virtualbox":    ("virtualbox",            "start virtualbox",                    "open -a VirtualBox"),
            "bluestacks":           ("bluestacks",            "start bluestacks",                    "open -a BlueStacks"),

            # ── Microsoft Office ──────────────────────────────────────
            "word":                 ("libreoffice --writer",  "start winword",                       "open -a 'Microsoft Word'"),
            "excel":                ("libreoffice --calc",    "start excel",                         "open -a 'Microsoft Excel'"),
            "powerpoint":           ("libreoffice --impress", "start powerpnt",                      "open -a 'Microsoft PowerPoint'"),
            "onenote":              ("libreoffice",           "start onenote",                       "open -a 'Microsoft OneNote'"),
            "microsoft teams":      ("teams",                 "start teams",                         "open -a 'Microsoft Teams'"),
            "teams":                ("teams",                 "start teams",                         "open -a 'Microsoft Teams'"),
            "outlook":              ("thunderbird",           "start outlook",                       "open -a Outlook"),
            "microsoft to do":      ("",                      "start ms-todo:",                      ""),
            "to do":                ("",                      "start ms-todo:",                      ""),
            "sticky notes":         ("",                      "start stikynot",                      ""),
            "clipchamp":            ("",                      "start clipchamp:",                    ""),
            "power automate":       ("",                      "start ms-powerautomate:",             ""),
            "copilot":              ("",                      "start microsoft-edge:https://copilot.microsoft.com",""),
            "microsoft copilot":    ("",                      "start microsoft-edge:https://copilot.microsoft.com",""),

            # ── Media & Entertainment ─────────────────────────────────
            "spotify":              ("spotify",               "start spotify",                       "open -a Spotify"),
            "vlc":                  ("vlc",                   "start vlc",                           "open -a VLC"),
            "vlc media player":     ("vlc",                   "start vlc",                           "open -a VLC"),
            "media player":         ("",                      "start wmplayer",                      "open -a 'Windows Media Player'"),
            "windows media player":  ("",                     "start wmplayer",                      ""),
            "sound recorder":       ("",                      "start ms-soundrecorder:",             ""),
            "xbox":                 ("",                      "start xbox:",                         ""),
            "roblox":               ("",                      "start roblox:",                       ""),
            "mech arena":           ("",                      "start mechArena:",                    ""),
            "google play games":    ("",                      "start com.google.android.gamespcapp:",""),
            "solitaire":            ("",                      "start xboxliveapp-9007:",             ""),

            # ── Communication ─────────────────────────────────────────
            "whatsapp":             ("whatsapp",              "start whatsapp:",                     "open -a WhatsApp"),
            "telegram":             ("telegram",              "start telegram",                      "open -a Telegram"),
            "discord":              ("discord",               "start discord",                       "open -a Discord"),

            # ── AI Apps ───────────────────────────────────────────────
            "chatgpt":              ("",                      "start microsoft-edge:https://chat.openai.com","open -a ChatGPT"),
            "claude":               ("",                      "start microsoft-edge:https://claude.ai","open -a Claude"),
            "codex":                ("",                      "start microsoft-edge:https://platform.openai.com",""),
            "bing":                 ("",                      "start microsoft-edge:https://bing.com",""),

            # ── System Tools ──────────────────────────────────────────
            "calculator":           ("gnome-calculator",      "calc",                                "open -a Calculator"),
            "paint":                ("gimp",                  "mspaint",                             "open -a GIMP"),
            "snipping tool":        ("",                      "start snippingtool",                  ""),
            "snip":                 ("",                      "start snippingtool",                  ""),
            "task manager":         ("gnome-system-monitor",  "taskmgr",                             "open -a 'Activity Monitor'"),
            "settings":             ("gnome-control-center",  "start ms-settings:",                  "open -a 'System Preferences'"),
            "file explorer":        ("nautilus",              "explorer",                            "open ."),
            "explorer":             ("nautilus",              "explorer",                            "open ."),
            "file manager":         ("nautilus",              "explorer",                            "open ."),
            "camera":               ("cheese",                "start microsoft.windows.camera:",     "open -a Photo Booth"),
            "photos":               ("eog",                   "start ms-photos:",                    "open -a Photos"),
            "windows clock":        ("",                      "start ms-clock:",                     ""),
            "clock":                ("",                      "start ms-clock:",                     ""),
            "weather":              ("",                      "start bingweather:",                  ""),
            "windows weather":      ("",                      "start bingweather:",                  ""),
            "quick assist":         ("",                      "start quickassist:",                  ""),
            "remote desktop":       ("",                      "start mstsc",                         ""),
            "speedtest":            ("",                      "start microsoft-edge:https://speedtest.net",""),
            "onedrive":             ("",                      "start onedrive",                      ""),
            "microsoft bing":       ("",                      "start microsoft-edge:https://bing.com",""),

            # ── HP Laptop Specific ────────────────────────────────────
            "omen gaming hub":      ("",                      "start omen gaming hub",               ""),
            "omen hub":             ("",                      "start omen gaming hub",               ""),
            "hp support":           ("",                      "start HPSupportAssistant",            ""),
            "hp support assistant": ("",                      "start HPSupportAssistant",            ""),
            "hp smart":             ("",                      "start HPSmart:",                      ""),
            "hp diagnostics":       ("",                      "start HPPCHardwareDiagnosticsWindows:",""),
            "intel graphics":       ("",                      "start igfxEM",                        ""),

            # ── Trading / Finance ─────────────────────────────────────
            "tradingview":          ("",                      "start microsoft-edge:https://tradingview.com",""),

            # ── Productivity ──────────────────────────────────────────
            "notepad":              ("gedit",                 "notepad",                             "open -a TextEdit"),
            "quick share":          ("",                      "start ms-settings-connectabledevices:",""),
        }

        WEB_MAP = {
            # ── Core ──────────────────────────────────────────────────
            "youtube":          "https://youtube.com",
            "gmail":            "https://mail.google.com",
            "google":           "https://google.com",
            "github":           "https://github.com",
            "chatgpt":          "https://chat.openai.com",
            "claude ai":        "https://claude.ai",
            "netflix":          "https://netflix.com",
            "twitter":          "https://twitter.com",
            "x.com":            "https://twitter.com",
            "instagram":        "https://instagram.com",
            "linkedin":         "https://linkedin.com",
            "facebook":         "https://facebook.com",
            "reddit":           "https://reddit.com",
            "pinterest":        "https://pinterest.com",
            # ── Study / JEE ──────────────────────────────────────────
            "jeemains":         "https://jeemain.nta.nic.in",
            "khan academy":     "https://khanacademy.org",
            "unacademy":        "https://unacademy.com",
            "physics wallah":   "https://pw.live",
            "pw":               "https://pw.live",
            "neet":             "https://neet.nta.nic.in",
            # ── Tools ────────────────────────────────────────────────
            "tradingview web":  "https://tradingview.com",
            "speedtest web":    "https://speedtest.net",
            "bing search":      "https://bing.com",
            "openrouter":       "https://openrouter.ai",
            "groq":             "https://console.groq.com",
            "google calendar":  "https://calendar.google.com",
            "google drive":     "https://drive.google.com",
            "google docs":      "https://docs.google.com",
            "google sheets":    "https://sheets.google.com",
            "google meet":      "https://meet.google.com",
            "notion":           "https://notion.so",
            "vercel":           "https://vercel.com",
            "stackoverflow":    "https://stackoverflow.com",
            "pypi":             "https://pypi.org",
        }

        # Check custom apps
        for key, path in self.custom.items():
            if key.lower() in n or n in key.lower():
                try:
                    subprocess.Popen(path, shell=True)
                    return f"Opening {key}."
                except:
                    return f"Failed to open {key}."

        # Browser + website multi-step, e.g. "Chrome open karo aur YouTube kholo".
        browser_key = None
        for candidate in ("google chrome", "chrome", "microsoft edge", "edge", "firefox"):
            if candidate in n:
                browser_key = candidate
                break
        website_key = None
        website_url = None
        for key, url in WEB_MAP.items():
            if key in n:
                website_key, website_url = key, url
                break
        if browser_key and website_url:
            try:
                if self.os_type == "Windows":
                    launcher = "chrome" if "chrome" in browser_key else ("msedge" if "edge" in browser_key else "firefox")
                    subprocess.Popen(f'start {launcher} "{website_url}"', shell=True)
                elif self.os_type == "Darwin":
                    app = "Google Chrome" if "chrome" in browser_key else ("Microsoft Edge" if "edge" in browser_key else "Firefox")
                    subprocess.Popen(["open", "-a", app, website_url])
                else:
                    launcher = "google-chrome" if "chrome" in browser_key else ("microsoft-edge" if "edge" in browser_key else "firefox")
                    subprocess.Popen([launcher, website_url])
                return f"Opening {browser_key} with {website_key}."
            except Exception:
                webbrowser.open(website_url)
                return f"Opening {website_key} in your browser."

        # Web shortcuts
        for key, url in WEB_MAP.items():
            if key in n:
                webbrowser.open(url)
                return f"Opening {key} in your browser."

        # Desktop apps
        for key, cmds in APP_MAP.items():
            if key in n or n in key:
                linux_cmd, win_cmd, mac_cmd = cmds
                cmd = {"Linux": linux_cmd, "Windows": win_cmd, "Darwin": mac_cmd}.get(self.os_type, linux_cmd)
                try:
                    subprocess.Popen(cmd, shell=True)
                    return f"Opening {key}."
                except Exception as e:
                    return f"Failed to open {key}: {e}"

        # Fallback — try running directly
        try:
            subprocess.Popen(name, shell=True)
            return f"Attempting to launch {name}."
        except:
            return f"I couldn't find '{name}'. Add it to custom_apps in your config."

    # ── Close App ─────────────────────────────────────────────
    def close_app(self, name: str) -> str:
        n = name.lower().strip()

        # ── Windows APP_CLOSE_MAP — common apps and their process names ──
        APP_CLOSE_MAP = {
            # ── Browsers ──────────────────────────────────────────────
            "chrome":               ["chrome.exe", "google chrome"],
            "google chrome":        ["chrome.exe", "google chrome"],
            "edge":                 ["msedge.exe", "microsoft edge"],
            "microsoft edge":       ["msedge.exe", "microsoft edge"],
            "firefox":              ["firefox.exe", "firefox"],

            # ── Code / Dev ────────────────────────────────────────────
            "vscode":               ["code.exe", "visual studio code"],
            "vs code":              ["code.exe", "visual studio code"],
            "visual studio code":   ["code.exe", "visual studio code"],
            "android studio":       ["studio64.exe", "android studio"],
            "terminal":             ["windows terminal", "wt.exe"],
            "windows terminal":     ["windows terminal", "wt.exe"],
            "cmd":                  ["cmd.exe", "command prompt"],
            "powershell":           ["powershell.exe", "powershell"],
            "python":               ["python.exe", "python"],
            "virtualbox":           ["virtualbox.exe", "virtualbox"],
            "oracle virtualbox":    ["virtualbox.exe", "virtualbox"],
            "bluestacks":           ["bluestacks.exe", "bluestacks"],

            # ── Microsoft Office ──────────────────────────────────────
            "word":                 ["winword.exe", "microsoft word"],
            "excel":                ["excel.exe", "microsoft excel"],
            "powerpoint":           ["powerpnt.exe", "microsoft powerpoint"],
            "onenote":              ["onenote.exe", "microsoft onenote"],
            "microsoft teams":      ["teams.exe", "microsoft teams"],
            "teams":                ["teams.exe", "microsoft teams"],
            "outlook":              ["outlook.exe", "outlook"],

            # ── Media & Entertainment ─────────────────────────────────
            "spotify":              ["spotify.exe", "spotify"],
            "vlc":                  ["vlc.exe", "vlc media player"],
            "vlc media player":     ["vlc.exe", "vlc media player"],
            "media player":         ["wmplayer.exe", "windows media player"],
            "windows media player": ["wmplayer.exe", "windows media player"],

            # ── Communication ─────────────────────────────────────────
            "whatsapp":             ["whatsapp.exe", "whatsapp"],
            "telegram":             ["telegram.exe", "telegram"],
            "discord":              ["discord.exe", "discord"],

            # ── System Tools ──────────────────────────────────────────
            "calculator":           ["calc.exe", "calculator"],
            "paint":                ["mspaint.exe", "paint"],
            "task manager":         ["taskmgr.exe", "task manager"],
            "file explorer":        ["explorer.exe", "file explorer"],
            "explorer":             ["explorer.exe", "file explorer"],
            "file manager":         ["explorer.exe", "file explorer"],
            "notepad":              ["notepad.exe", "notepad"],
        }

        # Check for app in close map
        for key, process_names in APP_CLOSE_MAP.items():
            if key in n or n in key:
                for process_name in process_names:
                    try:
                        # Try to kill by process name
                        result = subprocess.run(["taskkill", "/f", "/im", process_name], 
                                              capture_output=True, text=True, timeout=5)
                        if result.returncode == 0:
                            return f"Closed {key}."
                    except subprocess.TimeoutExpired:
                        return f"Timeout while trying to close {key}."
                    except Exception as e:
                        continue
                return f"Could not find running process for {key}."

        # Try to close by window title (fallback)
        try:
            # Use PowerShell to close window by title
            ps_cmd = f"Get-Process | Where-Object {{$_.MainWindowTitle -like '*{name}*'}} | Stop-Process -Force"
            result = subprocess.run(["powershell", "-Command", ps_cmd], 
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0 and "Stop-Process" in result.stderr:
                return f"Closed {name}."
        except:
            pass

        return f"I couldn't find '{name}' to close. Try using Task Manager."

    # ── File Operations ───────────────────────────────────────
    def search_files(self, query: str, path: str = None) -> str:
        search_path = str(self._resolve_storage_path(path)) if path else str(self._default_storage_root())
        results = []
        try:
            for root, dirs, files in os.walk(search_path):
                dirs[:] = [d for d in dirs if not d.startswith(".")
                           and d not in ("AppData","node_modules","__pycache__",".git","Windows")]
                for f in files:
                    if query.lower() in f.lower():
                        results.append(os.path.join(root, f))
                    if len(results) >= 5:
                        break
                if len(results) >= 5:
                    break
        except PermissionError:
            pass
        return ("Found: " + " | ".join(results)) if results else f"No files matching '{query}'."

    def open_file(self, path: str) -> str:
        try:
            target = self._find_existing_path(path)
            if target is None or not target.exists():
                return f"I could not find '{path}' inside your allowed storage roots."
            if not self._is_allowed_storage_path(target):
                return f"Access denied for {target}."
            if self.os_type == "Windows":
                os.startfile(str(target))
            elif self.os_type == "Darwin":
                subprocess.Popen(["open", str(target)])
            else:
                subprocess.Popen(["xdg-open", str(target)])
            return f"Opening {target}."
        except Exception as e:
            return f"Cannot open file: {e}"

    def open_desktop_folder(self) -> str:
        try:
            desktop = Path.home() / "Desktop"
            if not desktop.exists():
                desktop = Path.home()
            if self.os_type == "Windows":
                subprocess.Popen(["explorer.exe", str(desktop)])
            elif self.os_type == "Darwin":
                subprocess.Popen(["open", str(desktop)])
            else:
                subprocess.Popen(["xdg-open", str(desktop)])
            return f"Opening desktop folder in file explorer: {desktop}."
        except Exception as e:
            return f"Cannot open desktop folder: {e}"

    def edit_file(self, path: str) -> str:
        try:
            target = self._find_existing_path(path)
            if target is None or not target.exists():
                return f"I could not find '{path}' inside your allowed storage roots."
            if not target.is_file():
                return f"{target} is not a file I can edit directly."
            if not self._is_allowed_storage_path(target):
                return f"Access denied for {target}."

            if self.os_type == "Windows":
                editor = shutil.which("code")
                if editor:
                    subprocess.Popen([editor, str(target)])
                else:
                    subprocess.Popen(["notepad.exe", str(target)])
            elif self.os_type == "Darwin":
                subprocess.Popen(["open", "-t", str(target)])
            else:
                editor = os.environ.get("EDITOR") or shutil.which("xdg-open") or "nano"
                subprocess.Popen([editor, str(target)])
            return f"Opening {target} for editing."
        except Exception as e:
            return f"Cannot edit file: {e}"

    def create_folder(self, path: str) -> str:
        try:
            target = self._resolve_storage_path(path)
            if not self._is_allowed_storage_path(target):
                return f"Access denied for {target}."
            target.mkdir(parents=True, exist_ok=True)
            return f"Folder ready at {target}."
        except Exception as e:
            return f"Cannot create folder: {e}"

    def create_file(self, path: str, content: str = "") -> str:
        try:
            target = self._resolve_storage_path(path)
            if not self._is_allowed_storage_path(target):
                return f"Access denied for {target}."
            target.parent.mkdir(parents=True, exist_ok=True)
            if content and target.suffix.lower() not in IMAGE_FILE_EXTENSIONS:
                with open(target, "w", encoding="utf-8") as f:
                    f.write(content)
            else:
                target.touch(exist_ok=True)
            return f"File ready at {target}."
        except Exception as e:
            return f"Cannot create file: {e}"

    def list_desktop(self) -> str:
        desktop = Path.home() / "Desktop"
        if not desktop.exists():
            desktop = Path.home()
        try:
            items = os.listdir(desktop)[:10]
            return f"Desktop contains: {', '.join(items)}."
        except:
            return "Cannot read desktop."

    # ── Volume / Screen ───────────────────────────────────────
    def set_volume(self, level: int) -> str:
        level = max(0, min(100, level))
        try:
            if self.os_type == "Windows":
                # PowerShell method (no nircmd needed)
                script = (
                    f"$obj = New-Object -ComObject WScript.Shell;"
                    f"1..50 | %{{ $obj.SendKeys([char]174) }};"  # mute first
                    f"$vol = [math]::Round({level}/2);"
                    f"1..$vol | %{{ $obj.SendKeys([char]175) }}"
                )
                subprocess.run(["powershell","-c",script], capture_output=True)
            elif self.os_type == "Darwin":
                subprocess.run(f"osascript -e 'set volume output volume {level}'", shell=True)
            else:
                subprocess.run(f"amixer sset Master {level}%", shell=True)
            return f"Volume set to {level}%."
        except Exception as e:
            return f"Volume change failed: {e}"

    def _press_media_key_windows(self, vk: int) -> bool:
        try:
            ctypes.windll.user32.keybd_event(vk, 0, 0, 0)
            ctypes.windll.user32.keybd_event(vk, 0, 2, 0)
            return True
        except Exception:
            return False

    def control_media(self, action: str) -> str:
        action = str(action or "").strip().lower()
        if action not in {"playpause", "next", "previous"}:
            return "I can only control play/pause, next, or previous track right now."

        key_map = {
            "playpause": ("playpause", 0xB3),
            "next": ("nexttrack", 0xB0),
            "previous": ("prevtrack", 0xB1),
        }
        key_name, vk_code = key_map[action]

        if self.os_type == "Windows":
            if self._press_media_key_windows(vk_code):
                return f"Media {action.replace('playpause','play/pause')} command sent."

        if pyautogui:
            try:
                pyautogui.press(key_name)
                return f"Media {action.replace('playpause','play/pause')} command sent."
            except Exception:
                pass

        return "Media control is unavailable. Make sure pyautogui is installed and the media app is active."

    def read_screen_text(self) -> str:
        if not pyautogui:
            return "pyautogui not installed. Install it with: pip install pyautogui"
        if not pytesseract:
            return (
                "Screen reading requires Tesseract OCR. "
                "Step 1: `pip install pytesseract`. "
                "Step 2: Install the Tesseract engine from https://github.com/UB-Mannheim/tesseract/wiki and add it to your system PATH."
            )
        try:
            screenshot = pyautogui.screenshot()
            text = pytesseract.image_to_string(screenshot, lang="eng") if pytesseract else ""
            text = str(text or "").strip()
            return text or "No readable text was detected on the screen."
        except Exception as e:
            return f"Screen OCR failed: {e}"

    def take_screenshot(self) -> str:
        if not pyautogui:
            return "pyautogui not installed. Run: pip install pyautogui"
        try:
            fname = f"screenshot_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            path  = str(Path.home() / "Pictures" / fname)
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            pyautogui.screenshot(path)
            return f"Screenshot saved: {path}"
        except Exception as e:
            return f"Screenshot failed: {e}"

    # ── Power ─────────────────────────────────────────────────
    def shutdown(self, delay=30) -> str:
        if self.os_type == "Windows":
            subprocess.run(f"shutdown /s /t {delay}", shell=True)
        else:
            subprocess.run("shutdown -h +1", shell=True)
        return f"Shutdown scheduled in {delay} seconds."

    def restart(self) -> str:
        if self.os_type == "Windows":
            subprocess.run("shutdown /r /t 30", shell=True)
        else:
            subprocess.run("shutdown -r +1", shell=True)
        return "Restarting in 30 seconds."

    def lock_screen(self) -> str:
        if self.os_type == "Windows":
            subprocess.run("rundll32.exe user32.dll,LockWorkStation", shell=True)
        elif self.os_type == "Darwin":
            subprocess.run("pmset displaysleepnow", shell=True)
        else:
            subprocess.run("xdg-screensaver lock", shell=True)
        return "Screen locked."

    def get_clipboard(self) -> str:
        if pyperclip:
            return f"Clipboard contains: {pyperclip.paste()}"
        return "pyperclip not installed."

    # ── File Management for Telegram ──────────────────────────
    def list_files_telegram(self, path: str = None) -> str:
        """List files in a directory for Telegram display."""
        try:
            search_path = str(self._resolve_storage_path(path)) if path else str(self._default_storage_root())
            if not self._is_allowed_storage_path(Path(search_path)):
                return f"Access denied for {search_path}."

            items = []
            try:
                for item in os.listdir(search_path):
                    if not item.startswith('.'):  # Skip hidden files
                        full_path = os.path.join(search_path, item)
                        if os.path.isdir(full_path):
                            items.append(f"📁 {item}/")
                        else:
                            size = os.path.getsize(full_path)
                            size_str = self._format_file_size(size)
                            items.append(f"📄 {item} ({size_str})")
            except PermissionError:
                return f"Permission denied accessing {search_path}."

            if not items:
                return f"No files found in {search_path}."

            # Limit to 50 items for Telegram
            if len(items) > 50:
                items = items[:50]
                items.append("... (and more)")

            result = f"Files in {Path(search_path).name or 'root'}:\n" + "\n".join(items)
            return result[:4000]  # Telegram message limit

        except Exception as e:
            return f"Error listing files: {e}"

    def download_file_telegram(self, file_path: str) -> tuple:
        """Prepare file for Telegram download. Returns (file_path, caption) or (None, error_msg)."""
        try:
            target, message = self.find_telegram_share_file(file_path)
            if target is None:
                return None, message or f"File not found: {file_path}"

            # Check file size (Telegram limit is 50MB for bots)
            size = target.stat().st_size
            if size > 50 * 1024 * 1024:  # 50MB
                return None, f"File too large ({self._format_file_size(size)}). Telegram limit is 50MB."

            return str(target), f"📎 {target.name} ({self._format_file_size(size)})"

        except Exception as e:
            return None, f"Error preparing file: {e}"

    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f}{unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f}TB"


def _system_hardware_report(self) -> str:
    snapshot = self.hardware_snapshot()
    if not snapshot.get("available"):
        return f"Hardware check error: {snapshot.get('error', 'unavailable')}"

    cpu = snapshot["cpu_percent"]
    ram = snapshot["ram"]
    disk = snapshot["disk"]
    bat = snapshot["battery"]
    thermals = snapshot["thermals"]

    ru = ram.used / 1024**3
    rt = ram.total / 1024**3
    du = disk.used / 1024**3
    dt = disk.total / 1024**3
    bat_s = f"{bat.percent:.0f}% ({'charging' if bat.power_plugged else 'on battery'})" if bat else "N/A"

    extras = []
    if thermals.get("cpu_temp_c") is not None:
        extras.append(f"CPU temp {thermals['cpu_temp_c']:.0f}°C")
    if thermals.get("gpu_temp_c") is not None:
        extras.append(f"GPU temp {thermals['gpu_temp_c']:.0f}°C")
    if thermals.get("fans"):
        extras.append("Fans " + ", ".join(f"{fan['name']} {fan['rpm']} RPM" for fan in thermals["fans"][:2]))
    elif thermals.get("gpu_fan_percent") is not None:
        extras.append(f"GPU fan {thermals['gpu_fan_percent']:.0f}%")

    report = (
        f"CPU at {cpu:.0f}%, RAM {ru:.1f}/{rt:.1f} GB ({ram.percent}%), "
        f"Disk {du:.0f}/{dt:.0f} GB ({disk.percent}%), Battery: {bat_s}."
    )
    if extras:
        report += " " + " | ".join(extras) + "."
    return report.replace("..", ".")


def _system_thermal_report(self) -> str:
    thermals = self.thermal_snapshot(max_age=0.0)
    parts = []
    if thermals.get("cpu_temp_c") is not None:
        parts.append(f"CPU {thermals['cpu_temp_c']:.0f}°C")
    if thermals.get("gpu_temp_c") is not None:
        label = thermals.get("gpu_name") or "GPU"
        parts.append(f"{label} {thermals['gpu_temp_c']:.0f}°C")
    if thermals.get("fans"):
        parts.append(", ".join(f"{fan['name']} {fan['rpm']} RPM" for fan in thermals["fans"]))
    elif thermals.get("gpu_fan_percent") is not None:
        parts.append(f"GPU fan {thermals['gpu_fan_percent']:.0f}%")
    if thermals.get("source"):
        parts.append(f"source {thermals['source']}")
    if thermals.get("note"):
        parts.append(thermals["note"])
    if not parts:
        return "Thermal telemetry is unavailable right now."
    return ("Thermals: " + " | ".join(parts) + ".").replace("..", ".")


def _system_open_hp_thermal_controls(self) -> str:
    if self.os_type != "Windows":
        return "HP thermal control helper is only available on Windows."

    script = (
        "$app = Get-StartApps | Where-Object { $_.Name -match 'OMEN|Gaming Hub|Victus' } | "
        "Select-Object -First 1; "
        "if ($app) { Start-Process explorer.exe ('shell:AppsFolder\\' + $app.AppID); exit 0 } "
        "exit 1"
    )
    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", script],
            capture_output=True,
            text=True,
            timeout=6,
        )
        if result.returncode == 0:
            return (
                "Opening HP OMEN Gaming Hub. On supported Victus models, switch to "
                "Performance or Max Fan mode to increase cooling."
            )
    except Exception:
        pass

    return (
        "I could not launch OMEN Gaming Hub automatically. Open it manually and use "
        "Performance, Auto, or Max Fan mode if your HP Victus model supports it."
    )


SystemModule.hardware_report = _system_hardware_report
SystemModule.thermal_report = _system_thermal_report
SystemModule.open_hp_thermal_controls = _system_open_hp_thermal_controls

# ═══════════════════════════════════════════════════════════════
#  EMAIL MODULE
# ═══════════════════════════════════════════════════════════════
class EmailModule:
    EMAIL_RE = re.compile(r"^[A-Za-z0-9._%+-]+@(?:[A-Za-z0-9-]+\.)+[A-Za-z]{2,24}$")
    REPEATED_PUBLIC_SUFFIXES = {"com", "net", "org", "in", "edu", "gov", "co"}

    def __init__(self, cfg: dict):
        self.email    = cfg.get("email","")
        self.password = cfg.get("email_password","")
        self.imap     = cfg.get("imap_server","imap.gmail.com")
        self.smtp     = cfg.get("smtp_server","smtp.gmail.com")
        self.port     = cfg.get("smtp_port", 587)
        self.enabled  = bool(self.email and self.password)

    @classmethod
    def normalize_recipient(cls, value: str) -> str:
        raw = str(value or "").strip().strip("<>(),;:[]{}\"'")
        match = re.search(r"[A-Za-z0-9._%+-]+@(?:[A-Za-z0-9-]+\.)+[A-Za-z]{2,24}", raw)
        email = match.group(0).lower() if match else raw.lower()
        if not cls.EMAIL_RE.fullmatch(email):
            return ""
        local, domain = email.rsplit("@", 1)
        labels = domain.split(".")
        if any(not label or label.startswith("-") or label.endswith("-") for label in labels):
            return ""
        if len(labels) >= 2 and labels[-1] == labels[-2] and labels[-1] in cls.REPEATED_PUBLIC_SUFFIXES:
            return ""
        if ".." in local or local.startswith(".") or local.endswith("."):
            return ""
        return email

    def _decode(self, h) -> str:
        parts  = decode_header(h or "")
        result = []
        for part, enc in parts:
            if isinstance(part, bytes):
                result.append(part.decode(enc or "utf-8", errors="ignore"))
            else:
                result.append(str(part))
        return " ".join(result)

    def check_inbox(self, count=5) -> str:
        if not self.enabled:
            return "Email not configured in jarvis_config.json."
        try:
            mail = imaplib.IMAP4_SSL(self.imap, timeout=15)
            mail.login(self.email, self.password)
            mail.select("inbox")
            _, data = mail.search(None, "UNSEEN")
            ids = data[0].split()
            if not ids:
                mail.logout()
                return "Your inbox is clear — no unread messages."
            summaries = []
            for uid in reversed(ids[-count:]):
                _, msg_data = mail.fetch(uid, "(RFC822)")
                msg     = email_lib.message_from_bytes(msg_data[0][1])
                sender  = self._decode(msg.get("From","?")).split("<")[0].strip().strip('"')
                subject = self._decode(msg.get("Subject","No subject"))
                summaries.append(f"From {sender}: '{subject}'")
            mail.logout()
            return (f"You have {len(ids)} unread email{'s' if len(ids)>1 else ''}. "
                    + " | ".join(summaries[:3]))
        except Exception as e:
            return f"Email error: {e}"

    def send_email(self, to: str, subject: str, body: str) -> str:
        if not self.enabled:
            return "Email not configured."
        to = self.normalize_recipient(to)
        if not to:
            return "Invalid email address. Please check recipient email and try again."
        try:
            msg              = MIMEMultipart()
            msg["From"]      = self.email
            msg["To"]        = to
            msg["Subject"]   = subject
            msg.attach(MIMEText(body, "plain"))
            srv = smtplib.SMTP(self.smtp, self.port)
            srv.starttls()
            srv.login(self.email, self.password)
            srv.send_message(msg)
            srv.quit()
            return f"Email sent to {to}."
        except Exception as e:
            return f"Send failed: {e}"

    def send_email_with_attachment(self, to: str, subject: str, body: str, file_path: str) -> str:
        if not self.enabled:
            return "Email not configured."
        to = self.normalize_recipient(to)
        if not to:
            return "Invalid email address. Please check recipient email and try again."
        target = Path(file_path)
        if not target.is_file():
            return f"Attachment not found: {target}"
        try:
            msg = MIMEMultipart()
            msg["From"] = self.email
            msg["To"] = to
            msg["Subject"] = subject
            msg.attach(MIMEText(body, "plain"))

            part = MIMEBase("application", "octet-stream")
            with open(target, "rb") as f:
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f'attachment; filename="{target.name}"')
            msg.attach(part)

            srv = smtplib.SMTP(self.smtp, self.port)
            srv.starttls()
            srv.login(self.email, self.password)
            srv.send_message(msg)
            srv.quit()
            return f"Email sent to {to} with {target.name}."
        except Exception as e:
            return f"Attachment send failed: {e}"


# ═══════════════════════════════════════════════════════════════
#  DOCUMENT MODULE
# ═══════════════════════════════════════════════════════════════
class DocumentModule:
    def __init__(self, output_dir: str | Path = GENERATED_DIR):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _safe_filename(self, value: str, suffix: str = ".docx") -> str:
        stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(value or "document")).strip("._")
        stem = stem[:70] or "document"
        if not stem.lower().endswith(suffix):
            stem += suffix
        return stem

    def create_docx(self, title: str, content: str, output_name: str = None) -> str:
        try:
            from docx import Document
        except Exception:
            return "python-docx is not installed. Run: pip install python-docx"

        title = re.sub(r"\s+", " ", str(title or "JARVIS Document")).strip()
        content = str(content or "").strip()
        if not content:
            return "Document content is empty."

        filename = self._safe_filename(output_name or title)
        path = self.output_dir / filename
        if path.exists():
            stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            path = self.output_dir / self._safe_filename(f"{Path(filename).stem}_{stamp}")

        doc = Document()
        doc.add_heading(title, level=0)
        for block in re.split(r"\n{2,}", content):
            block = block.strip()
            if not block:
                continue
            if block.startswith("#"):
                heading = block.lstrip("#").strip()
                if heading:
                    doc.add_heading(heading, level=1)
                continue
            for line in block.splitlines():
                line = line.strip()
                if not line:
                    continue
                if re.match(r"^[-*]\s+", line):
                    doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
                elif re.match(r"^\d+[.)]\s+", line):
                    doc.add_paragraph(re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
                else:
                    doc.add_paragraph(line)
        doc.save(path)
        return str(path)


# ═══════════════════════════════════════════════════════════════
#  WHATSAPP MODULE (pywhatkit)
# ═══════════════════════════════════════════════════════════════
class WhatsAppModule:
    def __init__(self, cfg: dict):
        self.cfg = cfg
        self.enabled = True

    def _normalize_phone(self, phone: str) -> str:
        p = str(phone or "").strip()
        if not p:
            return ""
        p = p.replace(" ", "").replace("-", "")
        if p.startswith("+"):
            return p
        default_cc = str(self.cfg.get("whatsapp_default_country_code", "+91") or "+91").strip()
        if not default_cc.startswith("+"):
            default_cc = "+" + default_cc
        if p.startswith("0"):
            p = p.lstrip("0")
        return default_cc + p

    def _login_status_hint(self) -> str:
        if not pyautogui or not pytesseract:
            return " WhatsApp Web opened; agar QR login screen dikhe to pehle phone se scan karke login karna hoga."
        try:
            time.sleep(2)
            shot = pyautogui.screenshot()
            text = pytesseract.image_to_string(shot).lower()
            if any(marker in text for marker in ("scan to log in", "qr code", "link a device", "download whatsapp")):
                return " WhatsApp Web QR/login screen par hai, isliye message send confirm nahi hua. Phone se QR scan karke dobara try karo."
        except Exception:
            return " WhatsApp Web opened; agar QR login screen dikhe to pehle phone se scan karke login karna hoga."
        return " WhatsApp Web logged-in lag raha hai; message send verify kar lena."

    def send_message(self, phone: str, message: str) -> str:
        if not self.enabled:
            return "WhatsApp messaging is not available. Install pywhatkit: pip install pywhatkit"
        kit = ensure_pywhatkit()
        if not kit:
            return "WhatsApp messaging is not available. Install pywhatkit: pip install pywhatkit"
        phone_n = self._normalize_phone(phone)
        msg = str(message or "").strip()
        if not phone_n or not msg:
            return "WhatsApp message cancelled."
        try:
            # Opens WhatsApp Web and sends instantly. User must be logged in to WhatsApp Web.
            kit.sendwhatmsg_instantly(
                phone_no=phone_n,
                message=msg,
                wait_time=12,
                tab_close=True,
                close_time=3,
            )
            return f"WhatsApp message queued/opened for {phone_n}.{self._login_status_hint()}"
        except Exception as e:
            return f"WhatsApp send failed: {e}"

    def _find_whatsapp_window(self):
        if not pyautogui or not hasattr(pyautogui, "getWindowsWithTitle"):
            return None
        try:
            # More flexible title matching
            all_windows = pyautogui.getAllWindows()
            for win in all_windows:
                title = getattr(win, 'title', '').lower()
                if 'whatsapp' in title and getattr(win, "width", 0) > 200 and getattr(win, "height", 0) > 200:
                    return win
        except Exception:
            pass
        return None

    def read_incoming(self) -> str:
        if not pyautogui:
            return "pyautogui not installed. Install it with: pip install pyautogui"
        if not pytesseract:
            return (
                "Screen reading requires Tesseract OCR. "
                "Step 1: `pip install pytesseract`. "
                "Step 2: Install the Tesseract engine from https://github.com/UB-Mannheim/tesseract/wiki and add it to your system PATH."
            )
        try:
            whatsapp_window = self._find_whatsapp_window()
            if not whatsapp_window:
                import webbrowser
                webbrowser.open('https://web.whatsapp.com')
                time.sleep(3)
                whatsapp_window = self._find_whatsapp_window()
                if not whatsapp_window:
                    return "Could not find or open WhatsApp Web window. Please open https://web.whatsapp.com manually and ensure it's logged in."

            try:
                whatsapp_window.activate()
                time.sleep(0.8)
            except Exception:
                pass

            left = max(0, whatsapp_window.left)
            top = max(0, whatsapp_window.top)
            width = whatsapp_window.width
            height = whatsapp_window.height

            list_region = (
                left + int(width * 0.02),
                top + int(height * 0.15),
                int(width * 0.28),
                int(height * 0.75),
            )
            chat_region = (
                left + int(width * 0.31),
                top + int(height * 0.15),
                int(width * 0.67),
                int(height * 0.75),
            )

            list_screenshot = pyautogui.screenshot(region=list_region)
            chat_screenshot = pyautogui.screenshot(region=chat_region)
            list_text = pytesseract.image_to_string(list_screenshot, lang="eng")
            chat_text = pytesseract.image_to_string(chat_screenshot, lang="eng")

            list_lines = [line.strip() for line in list_text.splitlines() if line.strip()]
            chat_lines = [line.strip() for line in chat_text.splitlines() if line.strip()]

            if not list_lines and not chat_lines:
                return "No readable WhatsApp text found. Make sure a chat list and message pane are visible."

            part1 = " | ".join(list_lines[:10])
            part2 = " | ".join(chat_lines[:10])
            result_parts = []
            if part1:
                result_parts.append(f"Chats: {part1}")
            if part2:
                result_parts.append(f"Messages: {part2}")
            return " — ".join(result_parts)
        except Exception as e:
            return f"WhatsApp incoming read failed: {e}"


# ═══════════════════════════════════════════════════════════════
#  CONTACTS (CSV) — in-memory directory
# ═══════════════════════════════════════════════════════════════
class ContactsModule:
    """
    Loads contacts from a CSV file into memory.

    Supported headers (case-insensitive):
    - name, phone
    Optional:
    - whatsapp (if you want a dedicated WhatsApp number)

    If there are no headers, it will treat column0=name, column1=phone.
    """

    def __init__(self, cfg: dict):
        self.cfg = cfg
        self._contacts: dict[str, dict[str, str]] = {}
        self._loaded_from: str | None = None
        self._load_error: str | None = None
        self.reload()

    def _resolve_path(self) -> Path:
        raw = str(self.cfg.get("contacts_csv_path", "contacts.csv") or "contacts.csv").strip()
        p = Path(raw)
        if not p.is_absolute():
            p = (BASE_DIR / p).resolve()
        return p

    @staticmethod
    def _norm_key(name: str) -> str:
        return re.sub(r"\s+", " ", str(name or "").strip().lower())

    def reload(self) -> bool:
        path = self._resolve_path()
        self._contacts = {}
        self._loaded_from = str(path)
        self._load_error = None
        if not path.exists():
            self._load_error = f"Contacts file not found: {path}"
            return False

        def clean_phone(p: str) -> str:
            # Keep + and digits only (WhatsApp likes +<country><number>)
            s = str(p or "").strip()
            if not s:
                return ""
            s = s.replace(" ", "").replace("-", "").replace("(", "").replace(")", "")
            s = re.sub(r"[^\d+]", "", s)
            return s

        try:
            with open(path, "r", encoding="utf-8-sig", newline="") as f:
                sample = f.read(2048)
                f.seek(0)
                # Try to detect delimiter
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
                    delimiter = dialect.delimiter
                except Exception:
                    delimiter = ","

                reader = csv.reader(f, delimiter=delimiter)
                rows = [r for r in reader if any(c.strip() for c in r)]
                if not rows:
                    self._load_error = "Contacts CSV is empty."
                    return False

                # Header or no-header handling
                header = [c.strip().lower() for c in rows[0]]
                # Support both:
                # - simple CSV: name,phone[,whatsapp]
                # - Google Contacts export: First Name,...,Phone 1 - Value
                has_header = (
                    ("name" in header and (("phone" in header) or ("whatsapp" in header)))
                    or ("phone 1 - value" in header)
                    or ("phone 2 - value" in header)
                    or ("mobile" in header and "name" in header)
                )

                if has_header:
                    # Re-read with DictReader
                    f.seek(0)
                    dict_reader = csv.DictReader(f, delimiter=delimiter)
                    for r in dict_reader:
                        # Normalize keys (case-insensitive lookups)
                        r_lc = {str(k or "").strip().lower(): v for k, v in (r or {}).items()}

                        # Name: prefer explicit "name", otherwise compose like Google Contacts.
                        name = str(r_lc.get("name", "") or "").strip()
                        if not name:
                            first = str(r_lc.get("first name", "") or "").strip()
                            middle = str(r_lc.get("middle name", "") or "").strip()
                            last = str(r_lc.get("last name", "") or "").strip()
                            composed = " ".join([p for p in [first, middle, last] if p]).strip()
                            name = composed or str(r_lc.get("file as", "") or "").strip() or str(r_lc.get("nickname", "") or "").strip()
                        if not name:
                            continue

                        # Phone: prefer whatsapp/phone then common export fields.
                        phone_raw = (
                            r_lc.get("whatsapp")
                            or r_lc.get("phone")
                            or r_lc.get("phone 1 - value")
                            or r_lc.get("phone 2 - value")
                            or r_lc.get("mobile")
                            or r_lc.get("mobile phone")
                            or ""
                        )
                        phone = clean_phone(phone_raw)
                        if not phone:
                            continue
                        self._contacts[self._norm_key(name)] = {"name": name, "phone": phone}
                else:
                    # Assume first two columns: name, phone
                    for r in rows:
                        if len(r) < 2:
                            continue
                        name = str(r[0] or "").strip()
                        phone = clean_phone(str(r[1] or ""))
                        if not name or not phone:
                            continue
                        self._contacts[self._norm_key(name)] = {"name": name, "phone": phone}

            return True
        except Exception as e:
            self._load_error = f"Contacts load failed: {e}"
            return False

    def lookup_phone(self, query: str) -> str | None:
        q = self._norm_key(query)
        if not q:
            return None
        # Exact match
        if q in self._contacts:
            return self._contacts[q]["phone"]
        # Fuzzy contains match (first hit)
        for k, v in self._contacts.items():
            if q in k:
                return v["phone"]
        return None

    def status(self) -> str:
        if self._load_error:
            return self._load_error
        return f"Loaded {len(self._contacts)} contacts from {self._loaded_from}."


# ═══════════════════════════════════════════════════════════════
#  OWNER FACE RECOGNITION (face_recognition + dlib)
# ═══════════════════════════════════════════════════════════════
class OwnerFaceRecognizer:
    """Compares live camera faces to enrolled owner photos (offline)."""

    def __init__(self, cfg: dict):
        self.cfg = cfg
        self.enabled = bool(cfg.get("face_recognition_enabled", True))
        try:
            self.tolerance = float(cfg.get("face_match_tolerance", 0.55) or 0.55)
        except Exception:
            self.tolerance = 0.55
        self.tolerance = max(0.35, min(0.75, self.tolerance))
        self._encodings: list = []
        self._load_error: str | None = None
        self._loaded = False

    def _ensure_loaded(self) -> None:
        if self._loaded:
            return
        self._reload()
        self._loaded = True

    def _collect_image_paths(self) -> list[Path]:
        paths: list[Path] = []
        raw = self.cfg.get("owner_face_images") or []
        if isinstance(raw, str):
            raw = [raw]
        for item in raw:
            p = Path(str(item).strip())
            if not p.is_absolute():
                p = (BASE_DIR / p).resolve()
            if p.is_file() and p.suffix.lower() in (".jpg", ".jpeg", ".png", ".webp", ".bmp"):
                paths.append(p)
        folder = BASE_DIR / "owner_faces"
        if folder.is_dir():
            for p in sorted(folder.iterdir()):
                if p.is_file() and p.suffix.lower() in (".jpg", ".jpeg", ".png", ".webp", ".bmp"):
                    paths.append(p.resolve())
        seen: set[str] = set()
        uniq: list[Path] = []
        for p in paths:
            key = str(p)
            if key not in seen:
                seen.add(key)
                uniq.append(p)
        return uniq

    def _reload(self) -> None:
        self._encodings = []
        self._load_error = None
        if not self.enabled:
            self._load_error = "Face recognition disabled."
            return
        fr = ensure_face_recognition()
        if not fr:
            self._load_error = "face_recognition not installed. pip install face_recognition"
            return
        paths = self._collect_image_paths()
        if not paths:
            self._load_error = "No owner photos: add images to owner_faces/ or set owner_face_images in jarvis_config.json"
            return
        for p in paths:
            try:
                img = fr.load_image_file(str(p))
                encs = fr.face_encodings(img, num_jitters=1)
                if encs:
                    self._encodings.append(encs[0])
            except Exception as e:
                print(f"[FaceRec] skip {p.name}: {e}")
        if not self._encodings:
            self._load_error = "Could not extract a face from owner photos. Use clear front-facing selfies."

    def reload(self) -> None:
        self._loaded = False
        self._reload()
        self._loaded = True

    def is_ready(self) -> bool:
        self._ensure_loaded()
        return bool(self._encodings)

    def status(self) -> str:
        self._ensure_loaded()
        if not self.enabled:
            return self._load_error or "Face recognition disabled."
        if self._load_error and not self._encodings:
            return self._load_error
        return f"Owner face profile: {len(self._encodings)} encoding(s) loaded."

    def identify_primary(self, bgr, face_boxes: list) -> str:
        """
        Returns short label: OWNER | STRANGER | NO_FACE | NO_PROFILE
        """
        if not self.enabled:
            return "NO_PROFILE"
        fr = ensure_face_recognition()
        if not fr:
            self._load_error = "face_recognition not installed. pip install face_recognition"
            return "NO_PROFILE"
        self._ensure_loaded()
        if not self._encodings:
            return "NO_PROFILE"
        if not face_boxes:
            return "NO_FACE"
        import numpy as np

        x, y, w, h = max(face_boxes, key=lambda b: int(b[2]) * int(b[3]))
        H, W = bgr.shape[:2]
        pad_x = int(w * 0.12)
        pad_y = int(h * 0.12)
        x0 = max(0, x - pad_x)
        y0 = max(0, y - pad_y)
        x1 = min(W, x + w + pad_x)
        y1 = min(H, y + h + pad_y)
        crop = bgr[y0:y1, x0:x1]
        if crop.size == 0:
            return "NO_FACE"
        if cv2 is not None:
            rgb = cv2.cvtColor(crop, cv2.COLOR_BGR2RGB)
        else:
            rgb = crop[:, :, ::-1].copy()
        try:
            locs = fr.face_locations(rgb, model="hog")
            if not locs:
                return "NO_FACE"
            encs = fr.face_encodings(rgb, known_face_locations=locs, num_jitters=0)
            if not encs:
                return "NO_FACE"
            dists = fr.face_distance(self._encodings, encs[0])
            best = float(np.min(dists)) if len(dists) else 1.0
        except Exception as e:
            print(f"[FaceRec] {e}")
            return "NO_FACE"
        return "OWNER" if best <= self.tolerance else "STRANGER"

    def verify_image_file(self, image_path: str) -> tuple[str, float | None, str]:
        """
        Verify the largest face in an image against the enrolled owner profile.
        Returns (label, distance, message) where label is OWNER, STRANGER,
        NO_FACE, or NO_PROFILE. This never identifies arbitrary people.
        """
        if not self.enabled:
            return "NO_PROFILE", None, "Owner face verification disabled."
        fr = ensure_face_recognition()
        if not fr:
            self._load_error = "face_recognition not installed. pip install face_recognition"
            return "NO_PROFILE", None, self._load_error
        self._ensure_loaded()
        if not self._encodings:
            return "NO_PROFILE", None, self._load_error or "Owner face profile is not ready."
        try:
            import numpy as np

            img = fr.load_image_file(str(image_path))
            locs = fr.face_locations(img, model="hog")
            if not locs:
                return "NO_FACE", None, "Image me clear face detect nahi hua."
            encs = fr.face_encodings(img, known_face_locations=locs, num_jitters=1)
            if not encs:
                return "NO_FACE", None, "Image me face encoding create nahi ho paya."

            def _area(loc):
                top, right, bottom, left = loc
                return max(0, right - left) * max(0, bottom - top)

            primary_idx = max(range(len(locs)), key=lambda idx: _area(locs[idx]))
            dists = fr.face_distance(self._encodings, encs[primary_idx])
            best = float(np.min(dists)) if len(dists) else 1.0
            if best <= self.tolerance:
                return "OWNER", best, f"Owner face profile se match mil raha hai. Distance {best:.3f}, tolerance {self.tolerance:.3f}."
            return "STRANGER", best, f"Owner face profile se match nahi mila. Distance {best:.3f}, tolerance {self.tolerance:.3f}."
        except Exception as e:
            return "NO_FACE", None, f"Owner face verification failed: {e}"


# ═══════════════════════════════════════════════════════════════
#  WEATHER MODULE
# ═══════════════════════════════════════════════════════════════
class WeatherModule:
    def __init__(self, cfg: dict):
        self.key      = cfg.get("openweather_api_key","")
        self.location = cfg.get("location","Bihar, India")
        self.enabled  = bool(self.key)

    def current(self, city: str = None) -> str:
        city = city or self.location
        if not self.enabled:
            return "Weather API key not set. Get a free one at openweathermap.org."
        if not requests:
            return "requests package not installed."
        try:
            r = requests.get("https://api.openweathermap.org/data/2.5/weather",
                             params={"q":city,"appid":self.key,"units":"metric"}, timeout=5)
            d = r.json()
            if r.status_code != 200:
                return f"Weather unavailable: {d.get('message','error')}"
            return (f"{city}: {d['main']['temp']:.0f}°C, {d['weather'][0]['description']}, "
                    f"feels like {d['main']['feels_like']:.0f}°C, "
                    f"humidity {d['main']['humidity']}%.")
        except Exception as e:
            return f"Weather fetch failed: {e}"

# ═══════════════════════════════════════════════════════════════
#  NEWS MODULE
# ═══════════════════════════════════════════════════════════════
class NewsModule:
    def __init__(self, cfg: dict):
        self.key     = cfg.get("news_api_key","")
        self.enabled = bool(self.key)

    def headlines(self, query: str = None, country: str = "in") -> str:
        if not requests:
            return "requests package not installed."

        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "application/json, text/xml, application/xml;q=0.9, */*;q=0.8",
        }

        try:
            import xml.etree.ElementTree as ET
            from html import unescape

            # Google News RSS supports query via "search".
            if query:
                url = "https://news.google.com/rss/search"
                params = {"q": query, "hl": "en-IN", "gl": "IN", "ceid": "IN:en"}
            else:
                url = "https://news.google.com/rss"
                params = {"hl": "en-IN", "gl": "IN", "ceid": "IN:en"}

            r = requests.get(url, params=params, headers=headers, timeout=8)
            r.raise_for_status()  # Check for 403, 429, etc.
            
            xml_data = r.text or ""
            
            # Additional debug for the terminal if the response is suspicious
            if "<rss" not in xml_data:
                print(f"[News RSS] Warning: Google returned a non-RSS response. Status: {r.status_code}")
                
            root = ET.fromstring(xml_data)
            titles = []
            for item in root.findall(".//item"):
                t = item.findtext("title") or ""
                t = unescape(t).strip()
                if t:
                    # Google News titles often include " - Source". Keep the headline part.
                    if " - " in t:
                        t = t.split(" - ", 1)[0].strip()
                    titles.append(t)
                if len(titles) >= 4:
                    break
            if not titles:
                print(f"[News RSS] Warning: No <item> tags found in XML. Content snippet: {xml_data[:300]}")
                return "No headlines found."
            return "Headlines: " + " | ".join(titles) + "."
        except Exception as e:
            print(f"[News RSS] Exception: {e}")
            return f"News fetch failed: {e}"

# ═══════════════════════════════════════════════════════════════
#  LOCATION MODULE — IP geolocation + GPS via browser API
# ═══════════════════════════════════════════════════════════════
class LocationModule:
    def __init__(self, cfg: dict):
        self.cfg          = cfg
        self.google_maps_api_key = str(cfg.get("google_maps_api_key", "")).strip()
        self.manual_label = str(cfg.get("location_label") or cfg.get("location") or "").strip()
        self.manual_lat   = self._coerce_float(cfg.get("location_lat"))
        self.manual_lon   = self._coerce_float(cfg.get("location_lon"))
        self.cached: dict = {}
        self._lock        = threading.Lock()
        if self.manual_lat is not None and self.manual_lon is not None:
            self.cached = {
                "city": self.manual_label,
                "region": "",
                "country": "",
                "lat": self.manual_lat,
                "lon": self.manual_lon,
                "isp": "manual override",
                "ip": "",
            }
        self._refresh()  # background fetch on init

    def _coerce_float(self, value):
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    def _refresh(self):
        if self.manual_lat is not None and self.manual_lon is not None:
            return
        threading.Thread(target=self._fetch, daemon=True).start()

    def _fetch(self):
        """Fetch location via IP geolocation — free, no API key needed."""
        if not requests:
            return
        try:
            # ip-api.com — free, 45 requests/min, no key required
            r = requests.get("http://ip-api.com/json/?fields=status,city,regionName,country,lat,lon,isp,query",
                             timeout=5)
            data = r.json()
            if data.get("status") == "success":
                with self._lock:
                    self.cached = {
                        "city":    data.get("city", ""),
                        "region":  data.get("regionName", ""),
                        "country": data.get("country", ""),
                        "lat":     data.get("lat", 0.0),
                        "lon":     data.get("lon", 0.0),
                        "isp":     data.get("isp", ""),
                        "ip":      data.get("query", ""),
                    }
        except Exception as e:
            print(f"[Location] Fetch failed: {e}")

    def get(self) -> dict:
        with self._lock:
            return dict(self.cached)

    def get_str(self) -> str:
        d = self.get()
        if self.manual_label and self.manual_lat is not None and self.manual_lon is not None:
            return self.manual_label
        if not d:
            return self.cfg.get("location", "Location unavailable")
        return f"{d.get('city','')}, {d.get('region','')}, {d.get('country','')} (IP: {d.get('ip','')})"

    def get_coords(self) -> tuple[float, float] | None:
        d = self.get()
        if d and d.get("lat") and d.get("lon"):
            return float(d["lat"]), float(d["lon"])
        return None

    def get_map_url(self, zoom: int = 13) -> str:
        """Returns a Google Maps URL centered on current location."""
        coords = self.get_coords()
        if coords:
            lat, lon = coords
            return f"https://www.google.com/maps/search/?api=1&query={lat},{lon}"
        query = requests.utils.quote(self.get_str()) if requests else self.get_str().replace(" ", "+")
        return f"https://www.google.com/maps/search/?api=1&query={query}"

    def get_static_map_url(self, width=300, height=160, zoom=12) -> str:
        """Returns a Google Maps static preview URL when an API key is configured."""
        coords = self.get_coords()
        if coords and self.google_maps_api_key:
            lat, lon = coords
            marker = requests.utils.quote(f"color:red|label:J|{lat},{lon}") if requests else f"color:red|label:J|{lat},{lon}"
            return (
                "https://maps.googleapis.com/maps/api/staticmap"
                f"?center={lat},{lon}&zoom={zoom}&size={width}x{height}"
                f"&maptype=roadmap&markers={marker}&key={self.google_maps_api_key}"
            )
        return ""

    def refresh(self) -> str:
        if self.manual_lat is not None and self.manual_lon is not None:
            return "Location is locked to your configured exact coordinates."
        self._refresh()
        return "Location refresh initiated."

    def open_in_browser(self):
        webbrowser.open(self.get_map_url())


# ═══════════════════════════════════════════════════════════════
#  BROWSER CONTROL MODULE
# ═══════════════════════════════════════════════════════════════
class BrowserControlModule:
    def __init__(self, cfg: dict):
        self.cfg = cfg
        self._jobs = queue.Queue()
        self._worker = None
        self._worker_lock = threading.Lock()
        self._last_error = ""

    def _normalize_url(self, target: str) -> str:
        target = str(target or "").strip()
        if not target:
            return "https://www.google.com"
        if re.match(r"^https?://", target, re.I):
            return target
        if "." in target and " " not in target:
            return "https://" + target
        if requests:
            return "https://www.google.com/search?q=" + requests.utils.quote(target)
        return "https://www.google.com/search?q=" + target.replace(" ", "+")

    def _ensure_worker(self):
        with self._worker_lock:
            if self._worker and self._worker.is_alive():
                return
            self._worker = threading.Thread(target=self._worker_loop, daemon=True)
            self._worker.start()

    def _dispatch(self, action: str, payload: str = "", timeout: int = 45) -> str:
        self._ensure_worker()
        reply_q = queue.Queue(maxsize=1)
        self._jobs.put((action, payload, reply_q))
        try:
            return reply_q.get(timeout=timeout)
        except queue.Empty:
            return "Browser control timed out."

    def _worker_loop(self):
        try:
            from playwright.sync_api import sync_playwright
        except Exception:
            self._last_error = "Playwright is not installed. Run: pip install playwright"
            while True:
                _action, _payload, reply_q = self._jobs.get()
                reply_q.put(self._last_error)

        def click_visible_text(page, label: str) -> str:
            try:
                page.get_by_text(label, exact=False).first.click(timeout=2500)
                return f"Clicked browser text: {label}"
            except Exception:
                pass

            try:
                import re as _re

                for role in ("button", "link", "textbox", "menuitem"):
                    try:
                        page.get_by_role(role, name=_re.compile(_re.escape(label), _re.I)).first.click(timeout=1800)
                        return f"Clicked browser {role}: {label}"
                    except Exception:
                        continue
            except Exception:
                pass

            selector = "a,button,[role=button],[role=link],[onclick],input[type=button],input[type=submit],summary"
            candidates = page.evaluate(
                """
                (selector) => Array.from(document.querySelectorAll(selector)).map((el, index) => {
                    const rect = el.getBoundingClientRect();
                    const style = window.getComputedStyle(el);
                    const text = (el.innerText || el.value || el.getAttribute('aria-label') || el.title || '').trim();
                    return {
                        index,
                        text,
                        visible: !!text && rect.width > 0 && rect.height > 0 && style.visibility !== 'hidden' && style.display !== 'none'
                    };
                }).filter(item => item.visible)
                """,
                selector,
            )
            if best_text_match:
                match = best_text_match(label, candidates)
            else:
                match = next((item for item in candidates if label.lower() in item.get("text", "").lower()), None)
            if not match:
                visible = ", ".join(item.get("text", "")[:40] for item in candidates[:8])
                return f"I could not find a clickable browser item like '{label}'. Visible options: {visible or 'none'}."
            page.locator(selector).nth(int(match["index"])).click(timeout=5000)
            matched_text = str(match.get("text") or label).strip()
            return f"Clicked browser text: {matched_text} (matched from '{label}')"

        try:
            pw = sync_playwright().start()
            try:
                browser = pw.chromium.launch(channel="chrome", headless=False)
            except Exception:
                browser = pw.chromium.launch(headless=False)
            page = browser.new_page()
            last_url = "about:blank"
        except Exception as e:
            self._last_error = (
                "Browser control is not ready. If this is the first setup, run: "
                "python -m playwright install chromium. Details: " + str(e)
            )
            while True:
                _action, _payload, reply_q = self._jobs.get()
                reply_q.put(self._last_error)

        def new_browser_page(restore: bool = True):
            nonlocal browser, page, last_url
            try:
                browser.close()
            except Exception:
                pass
            try:
                browser = pw.chromium.launch(channel="chrome", headless=False)
            except Exception:
                browser = pw.chromium.launch(headless=False)
            page = browser.new_page()
            if restore and last_url and last_url != "about:blank":
                try:
                    page.goto(last_url, wait_until="domcontentloaded", timeout=25000)
                except Exception:
                    pass
            return page

        def ensure_live_page():
            nonlocal page
            try:
                if page and not page.is_closed():
                    _ = page.url
                    return page
            except Exception:
                pass
            return new_browser_page(restore=True)

        while True:
            action, payload, reply_q = self._jobs.get()
            try:
                page = ensure_live_page()

                if action == "open":
                    url = self._normalize_url(payload)
                    page.goto(url, wait_until="domcontentloaded", timeout=25000)
                    last_url = page.url
                    reply_q.put(f"Browser opened: {page.title() or url}")
                elif action == "read":
                    title = page.title()
                    url = page.url
                    text = page.locator("body").inner_text(timeout=7000)
                    text = re.sub(r"\s+", " ", text).strip()
                    reply_q.put(f"Browser page: {title} | {url}. Text: {text[:1400]}")
                elif action == "click":
                    label = str(payload or "").strip()
                    if not label:
                        reply_q.put("Tell me which visible text to click.")
                    else:
                        result = click_visible_text(page, label)
                        try:
                            last_url = page.url
                        except Exception:
                            pass
                        reply_q.put(result)
                elif action == "type":
                    text = str(payload or "")
                    if not text:
                        reply_q.put("Tell me what to type in the browser.")
                    else:
                        focused = page.locator(":focus")
                        if focused.count():
                            focused.fill(text, timeout=3000)
                        else:
                            page.keyboard.type(text)
                        reply_q.put("Typed into the browser.")
                elif action == "screenshot":
                    fname = f"browser_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                    path = str(Path.home() / "Pictures" / fname)
                    Path(path).parent.mkdir(parents=True, exist_ok=True)
                    page.screenshot(path=path, full_page=True)
                    reply_q.put(f"Browser screenshot saved: {path}")
                else:
                    reply_q.put("Unknown browser action.")
            except Exception as e:
                reply_q.put(f"Browser {action} failed: {e}")

    def open(self, target: str) -> str:
        return self._dispatch("open", target)

    def search(self, query: str) -> str:
        return self.open(query)

    def read_page(self) -> str:
        return self._dispatch("read")

    def click_text(self, label: str) -> str:
        return self._dispatch("click", label)

    def type_text(self, text: str) -> str:
        return self._dispatch("type", text)

    def screenshot(self) -> str:
        return self._dispatch("screenshot")


# ═══════════════════════════════════════════════════════════════
#  AUTONOMOUS VISION + ACTION AGENT
# ═══════════════════════════════════════════════════════════════
class VisionActionAgent:
    SAFE_ACTIONS = {"click", "double_click", "type", "press", "hotkey", "scroll", "wait", "done", "ask"}
    DANGEROUS_WORDS = (
        "delete", "remove", "format", "reset", "shutdown", "restart", "payment",
        "purchase", "buy", "send money", "transfer", "password", "otp", "2fa",
        "email send", "whatsapp send", "submit", "confirm order",
    )

    def __init__(self, cfg: dict):
        self.cfg = cfg
        self.model = str(cfg.get("gemini_vision_model", "gemini-2.0-flash") or "gemini-2.0-flash")
        self.max_steps = int(cfg.get("agent_max_steps", 12) or 12)
        self.request_timeout = max(15, min(120, int(cfg.get("gemini_request_timeout", 60) or 60)))
        self.max_model_attempts = max(1, min(5, int(cfg.get("gemini_max_model_attempts", 2) or 2)))
        self.last_screenshot = None

    def _api_key(self) -> str:
        return (
            os.getenv("GEMINI_API_KEY", "").strip()
            or os.getenv("GOOGLE_API_KEY", "").strip()
            or str(self.cfg.get("gemini_api_key", "")).strip()
        )

    def _screen_size(self) -> tuple[int, int]:
        if pyautogui:
            try:
                size = pyautogui.size()
                return int(size.width), int(size.height)
            except Exception:
                pass
        return 1920, 1080

    def _capture_screen(self) -> tuple[str | None, str | None]:
        if not PIL_Image:
            return None, "Pillow is not installed. Run: pip install pillow"
        try:
            fname = f"agent_screen_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S_%f')}.png"
            path = Path.home() / "Pictures" / fname
            path.parent.mkdir(parents=True, exist_ok=True)
            if mss_mod:
                with mss_mod.mss() as sct:
                    shot = sct.grab(sct.monitors[0])
                    img = PIL_Image.frombytes("RGB", shot.size, shot.rgb)
                    img.save(path)
            elif pyautogui:
                pyautogui.screenshot(str(path))
            else:
                return None, "Screen capture needs mss or pyautogui. Run: pip install mss pyautogui"
            self.last_screenshot = str(path)
            return str(path), None
        except Exception as e:
            return None, f"Screen capture failed: {e}"

    def _image_b64(self, path: str) -> str:
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode("ascii")

    def _extract_json(self, text: str) -> dict:
        text = str(text or "").strip()
        try:
            return json.loads(text)
        except Exception:
            pass
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            try:
                return json.loads(match.group(0))
            except Exception:
                pass
        return {"action": "ask", "message": "Vision model did not return valid JSON.", "raw": text[:500]}

    def _think(self, goal: str, screenshot_path: str, step: int, history: list[str]) -> dict:
        if not requests:
            return {"action": "ask", "message": "requests is not installed."}
        key = self._api_key()
        if not key:
            return {
                "action": "ask",
                "message": "Gemini API key missing. Add gemini_api_key to jarvis_config.json or set GEMINI_API_KEY.",
            }
        width, height = self._screen_size()
        prompt = (
            "You are the autonomous screen-control brain for J.A.R.V.I.S.\n"
            "Look at the screenshot and choose exactly ONE next action to complete the user's goal.\n"
            "Return ONLY valid JSON, no markdown.\n"
            f"Screen size: {width}x{height}. Step: {step}. Goal: {goal}\n"
            "Allowed actions:\n"
            "{\"action\":\"click\",\"x\":100,\"y\":200,\"reason\":\"...\"}\n"
            "{\"action\":\"double_click\",\"x\":100,\"y\":200,\"reason\":\"...\"}\n"
            "{\"action\":\"type\",\"text\":\"text to type\",\"reason\":\"...\"}\n"
            "{\"action\":\"press\",\"key\":\"enter\",\"reason\":\"...\"}\n"
            "{\"action\":\"hotkey\",\"keys\":[\"ctrl\",\"l\"],\"reason\":\"...\"}\n"
            "{\"action\":\"scroll\",\"amount\":-5,\"reason\":\"...\"}\n"
            "{\"action\":\"wait\",\"seconds\":2,\"reason\":\"...\"}\n"
            "{\"action\":\"done\",\"message\":\"brief completion summary\"}\n"
            "{\"action\":\"ask\",\"message\":\"question or safety confirmation needed\"}\n"
            "Rules: do not make purchases, delete files, send messages, submit forms, reveal passwords, or confirm irreversible actions. "
            "If such a step is needed, return ask. Use precise coordinates from the screenshot. "
            "Return done ONLY when the current screenshot visibly proves the goal is complete. "
            "Never return done with speculative wording like should have, would have, likely, after typing, or therefore. "
            "Recent history:\n" + "\n".join(history[-6:])
        )
        payload = {
            "contents": [{
                "parts": [
                    {"text": prompt},
                    {"inline_data": {"mime_type": "image/png", "data": self._image_b64(screenshot_path)}},
                ]
            }],
            "generationConfig": {
                "temperature": 0.15,
                "maxOutputTokens": 512,
                "responseMimeType": "application/json",
            },
        }
        candidate_models = []
        for model_name in (
            self.model,
            "gemini-2.0-flash",
            "gemini-2.5-flash",
            "gemini-2.5-flash-lite",
            "gemini-3-flash-preview",
        ):
            model_name = str(model_name or "").strip()
            if model_name and model_name not in candidate_models:
                candidate_models.append(model_name)
            if len(candidate_models) >= self.max_model_attempts:
                break

        errors = []
        for model_name in candidate_models:
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={key}"
                resp = requests.post(url, json=payload, timeout=self.request_timeout)
                if resp.status_code >= 400:
                    detail = self._gemini_error_summary(resp)
                    errors.append(f"{model_name}: {detail}")
                    if resp.status_code in {400, 401, 403}:
                        break
                    continue
                self.model = model_name
                data = resp.json()
                text = (
                    data.get("candidates", [{}])[0]
                    .get("content", {})
                    .get("parts", [{}])[0]
                    .get("text", "")
                )
                return self._extract_json(text)
            except Exception as e:
                errors.append(f"{model_name}: {type(e).__name__}: {e}")
                continue
        return {
            "action": "ask",
            "message": "Gemini vision request failed for configured models. " + " | ".join(errors)[:1800],
        }

    def _read_only_goal(self, goal: str) -> bool:
        lower = str(goal or "").lower()
        read_markers = (
            "read the current screen",
            "tell me what is wrong",
            "what is wrong",
            "describe the screen",
            "explain the screen",
            "what do you see",
            "analyze the screen",
        )
        return any(marker in lower for marker in read_markers)

    def _describe_screen(self, goal: str, screenshot_path: str) -> str:
        if not requests:
            return "requests is not installed."
        key = self._api_key()
        if not key:
            return "Gemini API key missing. Add gemini_api_key to jarvis_config.json or set GEMINI_API_KEY."
        prompt = (
            "You are J.A.R.V.I.S. screen diagnosis mode. "
            "Analyze the screenshot and answer the user's question directly. "
            "Do not return action JSON. Do not click or automate anything. "
            "Be concise and mention visible errors, likely cause, and next fix.\n"
            f"User question: {goal}"
        )
        payload = {
            "contents": [{
                "parts": [
                    {"text": prompt},
                    {"inline_data": {"mime_type": "image/png", "data": self._image_b64(screenshot_path)}},
                ]
            }],
            "generationConfig": {
                "temperature": 0.2,
                "maxOutputTokens": 700,
            },
        }
        errors = []
        for model_name in (
            self.model,
            "gemini-2.0-flash",
            "gemini-2.5-flash",
            "gemini-2.5-flash-lite",
            "gemini-3-flash-preview",
        ):
            model_name = str(model_name or "").strip()
            if not model_name:
                continue
            if len(errors) >= self.max_model_attempts:
                break
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={key}"
                resp = requests.post(url, json=payload, timeout=self.request_timeout)
                if resp.status_code >= 400:
                    errors.append(f"{model_name}: {self._gemini_error_summary(resp)}")
                    if resp.status_code in {400, 401, 403}:
                        break
                    continue
                self.model = model_name
                data = resp.json()
                text = (
                    data.get("candidates", [{}])[0]
                    .get("content", {})
                    .get("parts", [{}])[0]
                    .get("text", "")
                )
                text = str(text or "").strip()
                return text or "Gemini returned an empty screen analysis."
            except Exception as e:
                errors.append(f"{model_name}: {e}")
        return "Gemini screen analysis failed. " + " | ".join(errors)[:1800]

    def _gemini_error_summary(self, resp) -> str:
        try:
            data = resp.json()
            error = data.get("error", {}) if isinstance(data, dict) else {}
            message = str(error.get("message", "") or "").replace(self._api_key(), "[API_KEY]")
            details = error.get("details", [])
            detail_bits = []
            if isinstance(details, list):
                for item in details[:3]:
                    if not isinstance(item, dict):
                        continue
                    reason = item.get("reason") or item.get("@type", "")
                    violations = item.get("violations") or item.get("quotaMetric") or item.get("quotaId")
                    if violations:
                        detail_bits.append(f"{reason}: {violations}")
                    elif reason:
                        detail_bits.append(str(reason))
            suffix = f" Details: {'; '.join(detail_bits)}" if detail_bits else ""
            return f"HTTP {resp.status_code} {error.get('status', '')}: {message[:900]}{suffix}"
        except Exception:
            text = str(getattr(resp, "text", "") or "").replace(self._api_key(), "[API_KEY]")
            return f"HTTP {resp.status_code}: {text[:900]}"

    def _dangerous_goal(self, goal: str) -> bool:
        lower = str(goal or "").lower()
        return any(word in lower for word in self.DANGEROUS_WORDS)

    def _act(self, decision: dict) -> str:
        if not pyautogui:
            return "pyautogui is not installed. Run: pip install pyautogui"
        action = str(decision.get("action", "")).strip().lower()
        if action not in self.SAFE_ACTIONS:
            return f"Blocked unknown action: {action}"
        try:
            if action == "click":
                pyautogui.click(int(decision["x"]), int(decision["y"]))
                return f"clicked {decision.get('x')},{decision.get('y')}"
            if action == "double_click":
                pyautogui.doubleClick(int(decision["x"]), int(decision["y"]))
                return f"double clicked {decision.get('x')},{decision.get('y')}"
            if action == "type":
                pyautogui.write(str(decision.get("text", "")), interval=0.01)
                return "typed text"
            if action == "press":
                pyautogui.press(str(decision.get("key", "enter")))
                return f"pressed {decision.get('key', 'enter')}"
            if action == "hotkey":
                keys = decision.get("keys", [])
                if not isinstance(keys, list) or not keys:
                    return "hotkey missing keys"
                pyautogui.hotkey(*[str(k) for k in keys])
                return "pressed hotkey " + "+".join(str(k) for k in keys)
            if action == "scroll":
                pyautogui.scroll(int(decision.get("amount", -5)))
                return f"scrolled {decision.get('amount', -5)}"
            if action == "wait":
                seconds = max(0.2, min(8.0, float(decision.get("seconds", 1))))
                time.sleep(seconds)
                return f"waited {seconds:.1f}s"
            if action == "done":
                return "done"
            if action == "ask":
                return "ask"
        except Exception as e:
            return f"action failed: {e}"
        return "no action"

    def run(self, goal: str, max_steps: int | None = None) -> str:
        goal = str(goal or "").strip()
        if not goal:
            return "Tell me the automation goal after /agent."
        if self._dangerous_goal(goal):
            return "I need explicit confirmation before automating a risky task. Rephrase with /agent confirm if you want to proceed."

        steps = max(1, min(int(max_steps or self.max_steps), 20))
        if self._read_only_goal(goal):
            screenshot, err = self._capture_screen()
            if err:
                return err
            return self._describe_screen(goal, screenshot)

        history = []
        for step in range(1, steps + 1):
            screenshot, err = self._capture_screen()
            if err:
                return err
            decision = self._think(goal, screenshot, step, history)
            action = str(decision.get("action", "")).strip().lower()
            reason = str(decision.get("reason") or decision.get("message") or "")
            history.append(f"step {step}: decision={decision}")
            if action == "done":
                done_message = str(decision.get("message") or "")
                if re.search(r"\b(should have|would have|likely|probably|after typing|therefore|the user's goal was)\b", done_message, re.I):
                    return "Automation not completed: the vision model tried to guess instead of verifying the result on screen."
                return done_message or f"Automation complete in {step} steps."
            if action == "ask":
                return decision.get("message") or "Automation paused for confirmation."
            result = self._act(decision)
            history.append(f"step {step}: result={result}")
            if "failed" in result.lower() or result.startswith("Blocked"):
                return f"Automation stopped at step {step}: {result}. Last decision: {decision}"
            time.sleep(0.6)
        return "Automation step limit reached. Last actions: " + " | ".join(history[-4:])

# ═══════════════════════════════════════════════════════════════
#  NOTES MODULE
# ═══════════════════════════════════════════════════════════════
class NotesModule:
    def __init__(self):
        self.file  = NOTES_FILE
        self.notes = self._load()

    def _load(self) -> list:
        if self.file.exists():
            with open(self.file) as f:
                return json.load(f)
        return []

    def _save(self):
        with open(self.file,"w") as f:
            json.dump(self.notes, f, indent=2)

    def add(self, content: str) -> str:
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        self.notes.append({"text": content, "time": ts})
        self._save()
        return f"Note saved: '{content}'."

    def read(self) -> str:
        if not self.notes:
            return "No notes saved."
        last5 = self.notes[-5:]
        return "Notes: " + " | ".join(n["text"] for n in last5) + "."

    def clear(self) -> str:
        self.notes.clear()
        self._save()
        return "All notes cleared."

    def set_reminder(self, text: str, minutes: int, callback) -> str:
        def _fire():
            time.sleep(minutes * 60)
            callback(f"Reminder: {text}")
        threading.Thread(target=_fire, daemon=True).start()
        return f"Reminder set for {minutes} minute{'s' if minutes>1 else ''}: '{text}'."

# ═══════════════════════════════════════════════════════════════
#  CALENDAR MODULE
# ═══════════════════════════════════════════════════════════════
class CalendarModule:
    def __init__(self):
        self.file   = EVENTS_FILE
        self.events = self._load()

    def _load(self) -> list:
        if self.file.exists():
            with open(self.file) as f:
                return json.load(f)
        return []

    def _save(self):
        with open(self.file,"w") as f:
            json.dump(self.events, f, indent=2)

    def today(self) -> str:
        td = datetime.date.today().isoformat()
        ev = [e for e in self.events if e.get("date","") == td]
        if not ev:
            return "No events scheduled for today."
        return "Today: " + " | ".join(f"{e['time']} — {e['title']}" for e in ev) + "."

    def upcoming(self, days=7) -> str:
        start = datetime.date.today()
        end   = start + datetime.timedelta(days=days)
        ev    = [e for e in self.events
                 if start.isoformat() <= e.get("date","") <= end.isoformat()]
        if not ev:
            return f"No events in the next {days} days."
        return "Upcoming: " + " | ".join(f"{e['date']} {e['time']} {e['title']}" for e in ev[:5]) + "."

    def add(self, title, date, time_str="00:00") -> str:
        self.events.append({"title":title,"date":date,"time":time_str})
        self._save()
        return f"Event '{title}' added for {date} at {time_str}."

# ═══════════════════════════════════════════════════════════════
#  INTENT CLASSIFIER
# ═══════════════════════════════════════════════════════════════
class Intent:
    MAP = {
        "send_whatsapp":  ["send whatsapp", "whatsapp message", "message on whatsapp", "send a whatsapp", "send message on whatsapp"],
        "send_message":   ["send message to", "message to", "message someone", "send a message", "send text to", "text to"],
        "whatsapp_unread": ["unread whatsapp", "whatsapp unread", "whatsapp unread messages", "check whatsapp", "latest whatsapp", "show whatsapp message", "whatsapp message kya hai"],
        "whatsapp_read":  ["read whatsapp", "whatsapp incoming", "incoming whatsapp", "read whatsapp messages", "whatsapp messages", "whatsapp receive", "message aayi", "incoming message"],
        "media_control":  ["play music", "pause music", "resume music", "play song", "pause song", "play spotify", "pause spotify", "next song", "next track", "skip song", "spotify", "gaana", "play gaana", "pause gaana", "next gaana", "skip track"],
        "screen_read":    ["kya likha hai screen pe","what is on screen","read screen","read text on screen","screen pe kya likha hai","screen text","read screen text"],
        "open_app":       ["open","launch","start","run"],
        "close_app":      ["close","quit","stop","kill","terminate"],
        "list_files":     ["list files","show files","files in","directory","folder contents","what files"],
        "download_file":  ["download file","send file","get file","upload file"],
        "thermal_status": ["cpu temp","cpu temps","cpu temperature","gpu temp","gpu temps","gpu temperature","thermal status","thermal report","system temperature","temperatures","fan rpm","fan speed","cooling status"],
        "hardware":       ["cpu","ram","memory","disk","battery","hardware","system status","performance"],
        "fan_control":    ["speed up fan","increase fan","boost fan","cooler fan","fan mode","cooling mode","max fan","open omen","omen gaming hub","victus cooling","victus fan"],
        "network":        ["network","ip","internet","wifi","connection"],
        "email_check":    ["check email","emails","inbox","unread mail","any mail","new mail"],
        "email_unread":   ["unread email","email unread","check unread email","show unread mail","new unread mail","unread inbox"],
        "send_email":     ["send email","send mail","email to","compose","write email"],
        "calendar_today": ["today's schedule","today's events","schedule today"],
        "calendar_upcoming": ["upcoming","this week","next week","my schedule"],
        "add_event":      ["add event","schedule","book meeting","create event"],
        "open_calendar":  ["open calendar","google calendar","show calendar"],
        "weather":        ["weather","temperature","raining","forecast","sunny","hot outside"],
        "note_add":       ["take a note","note that","write down","remember this","jot","note bnao","note banao","note bna","note bana"],
        "note_read":      ["read notes","notes read","my notes","mere notes","show notes","what notes"],
        "note_clear":     ["clear notes","delete notes"],
        "reminder":       ["remind me","set reminder","reminder in","reminder do","reminder dedo","yaad dilao","alert me"],
        "news":           ["news","headlines","what's happening","latest news"],
        "time":           ["what time","current time","the time"],
        "date":           ["what date","today's date","what day","what's today's date","what's today"],
        "screenshot":     ["screenshot","capture screen","take screenshot"],
        "volume":         ["volume","set volume","turn up","turn down","mute"],
        "search_files":   ["find file","search file","look for file","locate file","where is file"],
        "open_file":      ["open file","open document","open folder"],
        "desktop":        ["what's on desktop","show desktop","list desktop"],
        "clipboard":      ["clipboard","what did i copy"],
        "processes":      ["running processes","top processes","what's running"],
        "camera_photo":   ["take photo","take picture","capture photo","selfie","click photo","open camera photo"],
        "camera_video":   ["record video","capture video","record clip"],
        "camera_live":    ["open camera","show camera","live camera","camera feed","camera view"],
        "camera_analyze": ["what do you see","analyze camera","describe what you see","look through camera","what's in front"],
        "camera_list":    ["list cameras","available cameras","check camera","detect camera"],
        "web_search":     ["search for","google","look up","search online","find online"],
        "shutdown":       ["shutdown","shut down","power off","turn off computer"],
        "restart":        ["restart","reboot"],
        "lock":           ["lock","lock screen","lock computer"],
        "self_improvement": ["edit yourself", "modify yourself", "upgrade yourself", "improve yourself", "add feature to yourself", "add a feature to jarvis", "change your code", "self edit", "self-edit", "self improve", "self-improve"],
        "self_knowledge":  ["read whole j.a.r.v.i.s folder", "read whole jarvis folder", "scan yourself", "learn yourself", "learn how you were made", "know how you were made", "how were you made", "what are you made of", "self knowledge", "self-knowledge"],
        "clear_chat":     ["clear chat","reset","forget conversation","new conversation"],
        "joke":           ["joke","make me laugh","say something funny"],
        "greet":          ["hello jarvis","hi jarvis","hey jarvis","good morning","good evening","good afternoon","good night"],
        "stop":           ["exit","quit","goodbye","bye","shutdown jarvis","stop jarvis","turn off jarvis"],
    }

    # Intents that should NEVER be bypassed by the word-count AI fallback
    FORCE_LOCAL_INTENTS = {
        "note_add", "note_read", "note_clear", "reminder",
        "stop", "send_email", "send_whatsapp", "send_message",
        "calendar_today", "add_event", "self_knowledge", "self_improvement",
    }

    @classmethod
    def classify(cls, text: str) -> str:
        lower = text.lower().strip()
        
        if lower.startswith("/ai "):
            return "ai_chat"
        if lower.startswith("/cmd "):
            lower = lower[5:]

        # First pass: check for force-local intents (notes, reminders, etc.)
        # These must NEVER be hijacked by the AI word-count rule.
        for intent, keywords in cls.MAP.items():
            if intent in cls.FORCE_LOCAL_INTENTS:
                if any(re.search(rf"\b{re.escape(kw)}\b", lower) for kw in keywords):
                    return intent

        # Avoid false positives for AI requests by checking for long, conversational phrases
        if len(lower.split()) > 8 or any(p in lower for p in ["figure out", "look at", "what do you see", "analyze", "where is", "read screen"]):
            return "ai_chat"
            
        # Second pass: check all other intents, sorted by longest keyword first
        for intent, keywords in cls.MAP.items():
            sorted_kws = sorted(keywords, key=len, reverse=True)
            if any(re.search(rf"\b{re.escape(kw)}\b", lower) for kw in sorted_kws):
                return intent
        return "ai_chat"


class ActionPlanner:
    EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
    EXPLICIT_SLASHES = ("/ai ", "/cmd ", "/cowork ", "/computer ", "/agent ", "/auto ", "/vision ")

    @classmethod
    def plan(cls, text: str) -> dict | None:
        source = str(text or "").strip()
        lower = source.lower()
        if not source or lower.startswith(cls.EXPLICIT_SLASHES):
            return None
        wants_doc = bool(re.search(r"\b(document|docx|doc|report|file)\b", lower)) and bool(
            re.search(r"\b(create|make|write|bna|bnao|bana|banao|prepare)\b", lower)
        )
        wants_send = bool(re.search(r"\b(send|mail|email|bhej|deliver)\b", lower)) or bool(cls.EMAIL_RE.search(source))
        if not wants_doc:
            return None
        recipient = cls.extract_email(source)
        topic = cls.extract_document_topic(source)
        wants_telegram = bool(re.search(r"\b(telegram|tg|mujhe)\b", lower)) and not recipient and "email" not in lower and "mail" not in lower
        steps = ["create_document"]
        if recipient or (wants_send and not wants_telegram):
            steps.append("send_email")
        elif wants_telegram:
            steps.append("send_telegram")
        return {
            "intent": "create_document_and_maybe_send",
            "steps": steps,
            "topic": topic,
            "recipient": recipient,
            "format": "docx",
            "needs_confirmation": bool(recipient or "send_email" in steps),
            "missing_fields": [
                name for name, value in [
                    ("topic", topic),
                    ("recipient", recipient if "send_email" in steps else True),
                ] if not value
            ],
        }

    @classmethod
    def extract_email(cls, text: str) -> str | None:
        match = cls.EMAIL_RE.search(str(text or ""))
        return EmailModule.normalize_recipient(match.group(0)) if match else None

    @classmethod
    def extract_document_topic(cls, text: str) -> str | None:
        raw = cls.EMAIL_RE.sub("", str(text or "")).strip()
        patterns = [
            r"(?:document|docx|doc|report|file)\s+(?:bna|bnao|bana|banao|create|make|write|prepare)\s+(?!(?:ke|ko|aur|or|and)\b)(.+?)\s+(?:ke baare(?: mein| me)?|par|pe|about|on)\b",
            r"(.+?)\s+(?:ke baare(?: mein| me)?|par|pe)\s+(?:document|docx|doc|report|file)\b",
            r"(?:about|on|par|pe)\s+(.+?)\s+(?:document|docx|doc|report|file)\b",
            r"(?:document|docx|doc|report|file)\s+(?:about|on|par|pe|ke baare(?: mein| me)?)\s+(.+?)(?:\s+(?:send|mail|email|bhej|deliver)\b|$)",
        ]
        for pattern in patterns:
            match = re.search(pattern, raw, flags=re.I)
            if match:
                return cls._clean_topic(match.group(1))
        cleaned = re.sub(
            r"\b(?:ek|a|an|create|make|write|prepare|bna|bnao|bana|banao|banakar|document|docx|doc|report|file|send|mail|email|bhej|deliver|telegram|to|ko|pe|par|ke|baare|mein|me)\b",
            " ",
            raw,
            flags=re.I,
        )
        return cls._clean_topic(cleaned)

    @staticmethod
    def _clean_topic(topic: str) -> str | None:
        topic = re.sub(r"\s+", " ", str(topic or "")).strip(" .,:;-")
        topic = re.sub(r"\b(?:and|aur|or|then|fir|phir)\b.*$", "", topic, flags=re.I).strip(" .,:;-")
        if not topic or len(topic) < 2:
            return None
        return topic[:120]


# ═══════════════════════════════════════════════════════════════
#  GREETING
# ═══════════════════════════════════════════════════════════════
class Greeter:
    def __init__(self, name: str):
        self.name = name

    def greet(self) -> str:
        h = datetime.datetime.now().hour
        period = "Good morning" if 4<=h<12 else "Good afternoon" if 12<=h<17 else "Good evening" if 17<=h<21 else "Good night"
        t = datetime.datetime.now().strftime("%I:%M %p")
        d = datetime.datetime.now().strftime("%A, %B %d")
        lines = [
            f"{period}, {self.name}. It's {t} on {d}. All systems are online and ready.",
            f"{period}, {self.name}. J.A.R.V.I.S online. {d}, {t}. How may I assist?",
            f"{period}, {self.name}. Systems initialized at {t}. Ready when you are.",
        ]
        return random.choice(lines)

    def bye(self) -> str:
        return random.choice([
            f"Going offline, {self.name}. J.A.R.V.I.S standing by.",
            f"Shutting down. Stay sharp, {self.name}.",
            f"Until next time, {self.name}.",
        ])

# ═══════════════════════════════════════════════════════════════
#  TELEGRAM BOT MODULE
# ═══════════════════════════════════════════════════════════════
class TelegramBotModule:
    def __init__(self, cfg: dict, jarvis_instance=None):
        self.cfg = cfg
        self.jarvis = jarvis_instance
        self.token = str(cfg.get("telegram_bot_token", "")).strip()
        self.allowed_user_id = cfg.get("telegram_allowed_user_id")
        self.enabled = bool(self.token and self.allowed_user_id and cfg.get("telegram_enabled", False))
        self.api_url = f"https://api.telegram.org/bot{self.token}"
        self._offset = 0
        self._running = False
        self._thread = None
        self._last_image_by_chat = {}
        if self.enabled:
            self._start_polling()

    def _send_message(self, chat_id: int, text: str) -> bool:
        if not self.enabled or not requests:
            return False
        try:
            payload = {
                "chat_id": chat_id,
                "text": text,
            }
            resp = requests.post(f"{self.api_url}/sendMessage", json=payload, timeout=5)
            return resp.status_code == 200
        except Exception as e:
            print(f"[WARN] Telegram send failed: {e}")
            return False

    def _send_long_message(self, chat_id: int, text: str, limit: int = 3900) -> bool:
        clean = str(text or "").strip()
        if not clean:
            return self._send_message(chat_id, "")
        chunks = _voiceengine_split_tts_text(clean, limit)
        ok = True
        for chunk in chunks or [clean[:limit]]:
            ok = self._send_message(chat_id, chunk[:4096]) and ok
            time.sleep(0.08)
        return ok

    def _send_file(self, chat_id: int, file_path: str, caption: str = None) -> bool:
        """Send a file to Telegram chat."""
        if not self.enabled or not requests:
            return False
        try:
            with open(file_path, 'rb') as f:
                files = {'document': f}
                data = {'chat_id': chat_id}
                if caption:
                    data['caption'] = caption[:1024]  # Telegram caption limit
                
                resp = requests.post(f"{self.api_url}/sendDocument", 
                                   files=files, data=data, timeout=30)
                return resp.status_code == 200
        except Exception as e:
            print(f"[WARN] Telegram file send failed: {e}")
            return False

    def _download_telegram_file(self, file_id: str, preferred_name: str = None) -> tuple[str | None, str | None]:
        """Download a Telegram file/photo locally and return (path, error)."""
        if not self.enabled or not requests:
            return None, "Telegram is not enabled."
        try:
            meta = requests.get(f"{self.api_url}/getFile", params={"file_id": file_id}, timeout=12)
            if meta.status_code != 200:
                return None, f"Telegram getFile failed: HTTP {meta.status_code}"
            payload = meta.json()
            if not payload.get("ok"):
                return None, str(payload)[:500]
            file_path = payload.get("result", {}).get("file_path", "")
            suffix = Path(file_path).suffix.lower()
            if not suffix:
                suffix = Path(preferred_name or "").suffix.lower() or ".jpg"
            safe_stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", Path(preferred_name or file_path or "telegram_image").stem)[:70]
            local = TELEGRAM_DOWNLOAD_DIR / f"{datetime.datetime.now().strftime('%Y%m%d_%H%M%S_%f')}_{safe_stem}{suffix}"
            url = f"https://api.telegram.org/file/bot{self.token}/{file_path}"
            data = requests.get(url, timeout=30)
            if data.status_code != 200:
                return None, f"Telegram file download failed: HTTP {data.status_code}"
            local.write_bytes(data.content)
            return str(local), None
        except Exception as e:
            return None, f"Telegram download failed: {e}"

    def _extract_incoming_image(self, message: dict) -> tuple[str | None, str | None]:
        """Return local image path from a Telegram photo or image document."""
        photos = message.get("photo") or []
        if photos:
            best = sorted(photos, key=lambda item: item.get("file_size", 0))[-1]
            return self._download_telegram_file(best.get("file_id", ""), "telegram_photo.jpg")

        document = message.get("document") or {}
        if document:
            mime = str(document.get("mime_type", "") or "").lower()
            name = str(document.get("file_name", "") or "telegram_document")
            ext = Path(name).suffix.lower()
            if mime.startswith("image/") or ext in IMAGE_FILE_EXTENSIONS:
                return self._download_telegram_file(document.get("file_id", ""), name)
        return None, None

    def _should_use_last_image(self, text: str) -> bool:
        lower = str(text or "").lower()
        return bool(re.search(r"\b(image|photo|pic|picture|screenshot|answer|ans|bta|bata|isme|is image|ye|this)\b", lower))

    def _is_owner_face_question(self, text: str) -> bool:
        lower = str(text or "").lower()
        return bool(
            re.search(r"\b(owner|prashant|face\s*match|verify\s*face|identify|identity|who\s+is|who's|kaun|kon|koun|ye\s+kaun|ye\s+kon)\b", lower)
            or "ye kaun hai" in lower
            or "ye kon hai" in lower
        )

    def _owner_face_reply(self, image_path: str) -> str:
        recognizer = getattr(self.jarvis, "owner_face", None)
        if not recognizer:
            return "Owner face verifier available nahi hai."
        label, distance, detail = recognizer.verify_image_file(image_path)
        if label == "OWNER":
            return (
                "Bhai owner-face verification pass hai. "
                "Ye enrolled owner profile se match ho raha hai, likely Prashant. "
                f"{detail}"
            )
        if label == "STRANGER":
            return (
                "Bhai owner-face verification me ye Prashant/owner profile se match nahi hua. "
                "Main random person ki identity guess nahi karunga. "
                f"{detail}"
            )
        if label == "NO_FACE":
            return f"Bhai is image me clear face verify nahi ho paya. {detail}"
        return f"Bhai owner-face verification ready nahi hai. {detail}"

    def _analyze_telegram_image(self, chat_id: int, message: dict, text: str) -> bool:
        image_path, err = self._extract_incoming_image(message)
        if not image_path and not err and isinstance(message.get("reply_to_message"), dict):
            image_path, err = self._extract_incoming_image(message["reply_to_message"])
        if not image_path and not err and text and self._should_use_last_image(text):
            cached = self._last_image_by_chat.get(chat_id)
            if cached and Path(cached).exists():
                image_path = cached
        if err:
            self._send_message(chat_id, err)
            return True
        if not image_path:
            return False
        self._last_image_by_chat[chat_id] = image_path
        question = (text or message.get("caption") or "").strip()
        if not question:
            question = "Analyze this image clearly. Mention visible objects, text, UI errors, and what action I should take next."
        elif question.lower().startswith(("/ai", "/vision", "/analyze")):
            question = re.sub(r"^/(?:ai|vision|analyze)\s*", "", question, flags=re.I).strip() or "Analyze this image."
        if self._is_owner_face_question(question):
            self._send_message(chat_id, "🧠 Image received. Verifying against owner face profile...")
            self._send_long_message(chat_id, self._owner_face_reply(image_path))
            return True
        self._send_message(chat_id, "🧠 Image received. Analyzing now...")
        try:
            response = self.jarvis.ai.analyze_image(image_path, question)
        except Exception as e:
            response = f"Image analysis failed: {e}"
        self._send_long_message(chat_id, str(response or "No analysis returned.").strip())
        return True

    def _get_updates(self) -> list:
        if not self.enabled or not requests:
            return []
        try:
            resp = requests.get(f"{self.api_url}/getUpdates?offset={self._offset}&timeout=10", timeout=15)
            if resp.status_code == 200:
                self._tg_fail_count = 0
                return resp.json().get("result", [])
        except Exception as e:
            # Only print warning once every ~6 failures (roughly once per minute)
            self._tg_fail_count = getattr(self, '_tg_fail_count', 0) + 1
            if self._tg_fail_count <= 1 or self._tg_fail_count % 6 == 0:
                print(f"[WARN] Telegram offline (attempt {self._tg_fail_count}): {type(e).__name__}")
        return []

    def _process_update(self, update: dict) -> None:
        if not self.jarvis:
            return
        try:
            message = update.get("message", {})
            chat_id = message.get("chat", {}).get("id")
            user_id = message.get("from", {}).get("id")
            text = (message.get("text") or message.get("caption") or "").strip()

            # Security: only allow the configured user
            if user_id != self.allowed_user_id:
                self._send_message(chat_id, "🔒 Unauthorized access attempt blocked.")
                return

            if self._analyze_telegram_image(chat_id, message, text):
                return

            if not text:
                return

            # Process command via JARVIS
            self._send_message(chat_id, f"🔄 Processing: {text[:50]}...")
            self.jarvis._is_telegram_context = True
            try:
                response = self.jarvis.handle(text)
            finally:
                self.jarvis._is_telegram_context = False
            response_text = str(response or "Command executed.").strip()[:4096]
            self._send_long_message(chat_id, response_text if response_text else "✅ Done.")
            
            # Check if there's a file to send
            if hasattr(self.jarvis, '_telegram_file_to_send') and self.jarvis._telegram_file_to_send:
                file_path, caption = self.jarvis._telegram_file_to_send
                self._send_message(chat_id, f"📤 Sending file...")
                if self._send_file(chat_id, file_path, caption):
                    self._send_message(chat_id, "✅ File sent successfully!")
                else:
                    self._send_message(chat_id, "❌ Failed to send file.")
                # Clear the file to send
                self.jarvis._telegram_file_to_send = None
        except Exception as e:
            print(f"[WARN] Telegram update processing failed: {e}")

    def _polling_worker(self) -> None:
        print("[INFO] Telegram bot polling started.")
        while self._running:
            try:
                updates = self._get_updates()
                for update in updates:
                    self._offset = update.get("update_id", 0) + 1
                    self._process_update(update)
                if not updates:
                    time.sleep(1)
            except Exception as e:
                print(f"[WARN] Telegram polling error: {e}")
                time.sleep(2)

    def _start_polling(self) -> None:
        if not self.enabled or self._running:
            return
        self._running = True
        self._thread = threading.Thread(target=self._polling_worker, daemon=True)
        self._thread.start()
        print(f"[INFO] Telegram bot initialized. Listening for messages from user {self.allowed_user_id}.")

    def stop(self) -> None:
        self._running = False
        if self._thread:
            self._thread.join(timeout=2)

    def notify(self, message: str) -> bool:
        if not self.enabled:
            return False
        return self._send_message(self.allowed_user_id, message)

    def wait_for_reply(self, timeout: int = 60) -> str | None:
        if not self.enabled:
            return None
        import time
        start_time = time.time()
        while time.time() - start_time < timeout:
            try:
                updates = self._get_updates()
                for update in updates:
                    self._offset = update.get("update_id", 0) + 1
                    message = update.get("message", {})
                    user_id = message.get("from", {}).get("id")
                    text = message.get("text", "").strip()
                    if user_id == self.allowed_user_id and text:
                        return text
                if not updates:
                    time.sleep(1)
            except Exception:
                time.sleep(1)
        return None


# ═══════════════════════════════════════════════════════════════
#  JARVIS CORE
# ═══════════════════════════════════════════════════════════════
class JARVIS:
    def __init__(self):
        banner = "═" * 60
        try:
            print("\n" + banner)
            print("  J.A.R.V.I.S v2.0 - Initializing Systems...")
            print(banner)
        except UnicodeEncodeError:
            banner = "=" * 60
            print("\n" + banner)
            print("  J.A.R.V.I.S v2.0 - Initializing Systems...")
            print(banner)

        self.cfg      = load_config()
        self.name     = self.cfg.get("user_name","Sir")
        self.voice    = VoiceEngine(self.cfg)
        self.ai       = AIBrain(self.cfg)
        self.system   = SystemModule(self.cfg)
        self.camera   = CameraModule(self.cfg)
        self.email    = EmailModule(self.cfg)
        self.documents = DocumentModule(GENERATED_DIR)
        self.whatsapp = WhatsAppModule(self.cfg)
        self.telegram = None  # Initialize placeholder
        self.contacts = ContactsModule(self.cfg)
        self.owner_face = OwnerFaceRecognizer(self.cfg)
        self.weather  = WeatherModule(self.cfg)
        self.news     = NewsModule(self.cfg)
        self.location = LocationModule(self.cfg)
        self.browser  = BrowserControlModule(self.cfg)
        self.agent    = VisionActionAgent(self.cfg)
        self.notes    = NotesModule()
        self.calendar = CalendarModule()
        self.greeter  = Greeter(self.name)
        self.wakes    = self.cfg.get("wake_words",["jarvis","hey jarvis"])
        self._running = True
        self.text_mode = False
        self._is_telegram_context = False
        self._telegram_file_to_send = None
        self.ai.voice_engine = self.voice
        telegram_cfg = dict(self.cfg)
        if str(os.getenv("JARVIS_DISABLE_TELEGRAM", "")).strip().lower() in {"1", "true", "yes", "on"}:
            telegram_cfg["telegram_enabled"] = False
        self.telegram = TelegramBotModule(telegram_cfg, jarvis_instance=self)  # Initialize after JARVIS is ready
        self.personal_telegram_secretary = None
        if bool(self.cfg.get("telegram_personal_enabled", False)):
            try:
                import jarvis_personal_telegram
                self.personal_telegram_secretary = jarvis_personal_telegram.start_background(jarvis_instance=self)
            except Exception as exc:
                print(f"[WARN] Personal Telegram secretary failed to start: {exc}")
        self.notification_secretary = None
        if bool(self.cfg.get("secretary_enabled", False)):
            try:
                import jarvis_secretary
                jarvis_secretary.start_background(jarvis_instance=self)
                self.notification_secretary = getattr(jarvis_secretary, "_default_secretary", None)
                apps = ", ".join(getattr(self.notification_secretary, "apps", []) or [])
                print(f"  [OK] Windows notification secretary started for: {apps}")
            except Exception as exc:
                print(f"[WARN] Windows notification secretary failed to start: {exc}")
        self.proactive = None
        if ProactiveEngine:
            try:
                self.proactive = ProactiveEngine(self)
                self.proactive.start()
            except Exception as exc:
                print(f"[WARN] ProactiveEngine failed to start: {exc}")

        try:
            print("  [OK] All systems online.\n")
        except UnicodeEncodeError:
            pass

    def boot(self):
        self.voice.speak(self.greeter.greet())
        hw = self.system.hardware_report()
        self.voice.speak(f"Quick system check: {hw}")

    # ── Main command handler ──────────────────────────────────
    def _extract_storage_path(self, text: str) -> str | None:
        source = str(text or "").strip()
        if not source:
            return None

        quoted = re.search(r'"([^"]+)"', source) or re.search(r"'([^']+)'", source)
        if quoted:
            return quoted.group(1).strip()

        drive = re.search(
            r"([A-Za-z]:\\.*?)(?=\s+\b(?:with content|by telegram|via telegram|to telegram|on telegram|mujhe|please)\b|$)",
            source,
            re.I,
        )
        if drive:
            return drive.group(1).strip().rstrip(" .")
        return None

    def _extract_share_file_path(self, text: str) -> str | None:
        path = self._extract_storage_path(text)
        if path:
            return path
        raw = re.sub(r"^/(?:agent|auto|vision)\s+", "", str(text or ""), flags=re.I).strip()
        raw = re.sub(r"^/(?:file|sendfile|sharefile)\s+", "", raw, flags=re.I).strip()
        raw = re.sub(r"\b(?:share|send|bhej|bhejo|deliver)\s+(?:file|document|photo|video|report)\b", "", raw, flags=re.I).strip()
        raw = re.sub(r"\b(?:file|document|photo|video|report)\s+(?:share|send|bhej|deliver|bhejo)\b", "", raw, flags=re.I).strip()
        raw = re.sub(r"\b(?:by|via|through|using|to|on)\s+telegram\b", "", raw, flags=re.I).strip()
        raw = re.sub(r"\b(?:to me|mujhe|please|plz|mere|mera|meri|my|laptop|pc|computer)\b", "", raw, flags=re.I).strip()
        raw = re.sub(r"\b(?:ko|pe|par|se|from)\s*$", "", raw, flags=re.I).strip()
        raw = re.sub(r"\b(?:bhejo|bhej\s*do|send\s*(?:kar\s*)?do|share\s*(?:kar\s*)?do|deliver\s*(?:kar\s*)?do)\b\s*$", "", raw, flags=re.I).strip()
        folder_match = re.search(
            r"\b(desktop|downloads?|documents?|pictures?|photos?|videos?|music)\b(?:\s+folder)?(?:\s+(?:se|from|mein|me|ke andar|inside))?\s+(.+)$",
            raw,
            flags=re.I,
        )
        if folder_match:
            alias = folder_match.group(1).lower()
            alias = {
                "download": "downloads",
                "downloads": "downloads",
                "document": "documents",
                "documents": "documents",
                "picture": "pictures",
                "pictures": "pictures",
                "photo": "pictures",
                "photos": "pictures",
                "video": "videos",
                "videos": "videos",
            }.get(alias, alias)
            tail = folder_match.group(2).strip(" .,:;-\"'")
            if tail:
                return f"{alias}/{tail}"
        return raw or None

    def _build_operator_report(self) -> str:
        now = datetime.datetime.now().strftime("%A, %B %d, %Y at %I:%M %p")
        sections = [
            f"J.A.R.V.I.S operator report for {self.name}",
            f"Generated: {now}",
            f"System: {platform.system()} {platform.release()}",
            f"Hardware: {self.system.hardware_report()}",
            f"Thermals: {self.system.thermal_report()}",
            f"Network: {self.system.network_info()}",
            f"Top processes: {self.system.top_processes()}",
            f"Location: {self.location.get_str()}",
        ]
        try:
            sections.append(f"Notes: {self.notes.read()}")
        except Exception:
            pass
        return "\n".join(sections)

    def _send_text_to_owner(self, subject: str, body: str) -> str:
        if self.telegram and getattr(self.telegram, "enabled", False):
            ok = self.telegram._send_message(self.telegram.allowed_user_id, body[:3900])
            if ok:
                return "Report sent to your Telegram."
            return "Telegram is configured, but the report send failed."
        if self.email.enabled and self.email.email:
            return self.email.send_email(self.email.email, subject, body)
        return "No delivery channel is configured. Enable Telegram or email in jarvis_config.json."

    def _send_file_to_owner(self, file_path: str) -> str:
        target, caption = self.system.download_file_telegram(file_path)
        if not target:
            return caption
        if self.telegram and getattr(self.telegram, "enabled", False):
            ok = self.telegram._send_file(self.telegram.allowed_user_id, target, caption)
            return f"Shared {Path(target).name} to your Telegram." if ok else "Telegram file share failed."
        if self.email.enabled and self.email.email:
            return self.email.send_email_with_attachment(
                self.email.email,
                f"J.A.R.V.I.S shared file: {Path(target).name}",
                "File shared by J.A.R.V.I.S.",
                target,
            )
        return "No delivery channel is configured. Enable Telegram or email in jarvis_config.json."

    def _send_operator_report(self) -> str:
        report = self._build_operator_report()
        return self._send_text_to_owner("J.A.R.V.I.S Operator Report", report)

    def _handle_operator_command(self, text: str) -> str | None:
        lower = text.lower().strip()
        explicit_path = self._extract_storage_path(text)

        if re.search(r"\b(?:secretary|notification secretary|notifications?)\b", lower):
            if re.search(r"\b(?:status|working|enabled|running|kaam|check|test)\b", lower):
                secretary = getattr(self, "notification_secretary", None)
                if secretary and hasattr(secretary, "status"):
                    return secretary.status()
                enabled = bool(self.cfg.get("secretary_enabled", False))
                apps = ", ".join(self.cfg.get("secretary_apps", []) or [])
                return (
                    f"Windows notification secretary {'enabled' if enabled else 'disabled'} hai, "
                    f"lekin active thread nahi mila. Apps: {apps or 'not configured'}."
                )

        if lower.startswith(("/file ", "/sendfile ", "/sharefile ")):
            path = self._extract_share_file_path(text)
            return self._send_file_to_owner(path) if path else "Tell me which file to share."

        if re.search(r"\b(send|share|bhej|deliver)\b.*\b(report|status|system report|operator report)\b", lower) or re.search(r"\b(report|status)\b.*\b(send|share|bhej)\b", lower):
            return self._send_operator_report()

        if re.fullmatch(r"(?:operator\s+)?(?:system\s+)?(?:status|report|system report|operator report)", lower):
            return self._build_operator_report()

        if re.search(r"\b(operator|control|capabilities|what can you control)\b", lower):
            return (
                "Operator control is online. I can open apps, open files and folders, edit text/code files, "
                "create files or folders, search project files, share files to you, and send system reports. "
                "Use commands like: open Chrome, edit file \"C:\\path\\note.txt\", share file \"C:\\path\\report.pdf\", or send me report."
            )

        if re.search(r"\b(?:edit|modify|update)\s+(?:file|document|note|config|script)\b", lower):
            path = explicit_path or re.split(r"\b(?:edit|modify|update)\s+(?:file|document|note|config|script)\b", text, maxsplit=1, flags=re.I)[-1].strip()
            return self.system.edit_file(path) if path else "Tell me which file to edit."

        if re.search(r"\b(?:share|send|bhej|deliver)\s+(?:file|document|photo|video|report)\b", lower):
            path = explicit_path or self._extract_share_file_path(text)
            return self._send_file_to_owner(path) if path else "Tell me which file to share."

        if (
            re.search(r"\b(?:telegram|mujhe|bhejo|bhej do|send me|send kar do|share kar do)\b", lower)
            and (
                re.search(r"\.(?:[a-z0-9]{1,8})\b", lower)
                or re.search(r"\b(?:desktop|downloads?|documents?|pictures?|photos?|videos?|music)\b", lower)
                or explicit_path
            )
        ):
            path = explicit_path or self._extract_share_file_path(text)
            return self._send_file_to_owner(path) if path else "Tell me which file to share."

        return None

    def _handle_browser_command(self, text: str) -> str | None:
        lower = text.lower().strip()
        if not re.search(r"\b(browser|web page|website|tab)\b", lower):
            return None

        if re.search(r"\b(browser|website|web page)\s+(?:open|go to|navigate to)\b", lower):
            target = re.split(r"\b(?:open|go to|navigate to)\b", text, maxsplit=1, flags=re.I)[-1].strip()
            return self.browser.open(target)

        if re.search(r"\b(?:open|go to|navigate to)\s+(?:website|web page|browser)\b", lower):
            target = re.split(r"\b(?:website|web page|browser)\b", text, maxsplit=1, flags=re.I)[-1].strip()
            return self.browser.open(target)

        if re.search(r"\b(browser|web)\s+search\b", lower):
            query = re.split(r"\bsearch\b", text, maxsplit=1, flags=re.I)[-1].strip()
            query = re.sub(r"^\b(for|about)\b", "", query, flags=re.I).strip()
            return self.browser.search(query)

        if re.search(r"\b(read|summarize|scan)\s+(?:the\s+)?(?:browser|web page|page|tab)\b", lower) or re.search(r"\b(browser|tab)\s+(?:read|scan)\b", lower):
            return self.browser.read_page()

        if re.search(r"\b(browser|web page|tab)\s+click\b", lower):
            label = re.split(r"\bclick\b", text, maxsplit=1, flags=re.I)[-1].strip().strip('"')
            return self.browser.click_text(label)

        if re.search(r"\b(browser|web page|tab)\s+type\b", lower):
            content = re.split(r"\btype\b", text, maxsplit=1, flags=re.I)[-1].strip().strip('"')
            return self.browser.type_text(content)

        if re.search(r"\b(browser|web page|tab)\s+screenshot\b", lower):
            return self.browser.screenshot()

        if re.search(r"\b(browser|web)\s+(control|capabilities|help)\b", lower):
            return (
                "Browser control is online through Playwright. I can open websites, search, read the current page, "
                "click visible text, type into the focused field, and capture browser screenshots. Try: browser open youtube.com."
            )

        return None

    def _handle_agent_command(self, text: str) -> str | None:
        source = str(text or "").strip()
        lower = source.lower()
        if not lower.startswith(("/agent ", "/auto ", "/vision ")):
            return None
        goal = re.sub(r"^/(?:agent|auto|vision)\s+", "", source, flags=re.I).strip()
        confirmed = False
        if goal.lower().startswith("confirm "):
            confirmed = True
            goal = goal[8:].strip()
        if not goal:
            return "Tell me the automation goal after /agent."
        if goal.lower() in {"help", "capabilities", "status"}:
            return (
                "Autonomous agent mode is ready. Use /agent followed by a goal. "
                "It observes the screen, asks Gemini Vision for one JSON action, executes with pyautogui, and repeats until done. "
                "Examples: /agent open browser and search Python tutorials, /agent read the current screen and explain it."
            )
        direct_resp = self._handle_direct_agent_shortcut(goal)
        if direct_resp:
            return direct_resp
        if self.agent._read_only_goal(goal):
            screenshot, err = self.agent._capture_screen()
            if err:
                return err
            analysis = self.agent._describe_screen(goal, screenshot)
            if re.search(r"\b(Gemini .*failed|quota|429|timed out|read timeout|request failed)\b", analysis, re.I):
                try:
                    fallback = self.ai.analyze_image(
                        screenshot,
                        f"{goal}\n\nIf the screen contains an error, explain what it means and how to fix it. Answer in concise Hinglish.",
                    )
                    if fallback and not fallback.lower().startswith("ai error"):
                        return fallback
                except Exception as exc:
                    return f"{analysis}\nFallback vision also failed: {exc}"
            return analysis
        if confirmed:
            original = self.agent._dangerous_goal
            try:
                self.agent._dangerous_goal = lambda _goal: False
                return self.agent.run(goal)
            finally:
                self.agent._dangerous_goal = original
        return self.agent.run(goal)

    def _handle_direct_agent_shortcut(self, goal: str) -> str | None:
        """Route obvious local OS requests without spending a vision call."""
        lower = str(goal or "").lower().strip()
        if not lower:
            return None
        if re.search(r"\b(?:open|launch|show|start)\b", lower) and re.search(r"\b(?:camera|webcam|web cam|camera app)\b", lower):
            if re.search(r"\b(?:feed|live|view|webcam|web cam)\b", lower):
                threading.Thread(target=self.camera.show_live_feed, daemon=True).start()
                return "Camera feed opened. Press Q in the camera window to close."
            return self.system.open_app("camera")
        if re.search(r"\b(?:take|capture|click)\b", lower) and re.search(r"\b(?:photo|picture|selfie)\b", lower):
            ok, result = self.camera.capture_photo()
            return f"Photo saved at {result}." if ok else result
        if re.search(r"\b(?:record|capture)\b", lower) and re.search(r"\b(?:video|clip)\b", lower):
            nums = re.findall(r"\d+", lower)
            secs = int(nums[0]) if nums else 10
            ok, result = self.camera.record_video(secs)
            return f"Video saved at {result}." if ok else result
        if re.search(r"\b(?:list|check|detect|available)\b", lower) and re.search(r"\b(?:camera|webcam|web cam)\b", lower):
            return self.camera.list_cameras()
        if re.search(r"\b(?:share|send|bhej|deliver)\s+(?:file|document|photo|video|report)\b", lower):
            path = self._extract_share_file_path(goal)
            return self._send_file_to_owner(path) if path else "Tell me which file to share."
        if re.search(r"\b(?:open|launch|show)\b", lower) and re.search(r"\b(?:desktop|desk top)\b", lower):
            return self.system.open_desktop_folder()
        if re.search(r"\b(?:open|launch|show)\b", lower) and re.search(r"\b(?:file explorer|explorer|file manager)\b", lower):
            return self.system.open_app("file explorer")
        if re.search(r"\b(?:open|launch|show)\b", lower) and re.search(r"\b(?:downloads|documents|pictures|photos)\s+(?:folder|in file explorer|in explorer)\b", lower):
            folder_match = re.search(r"\b(downloads|documents|pictures|photos)\b", lower)
            if folder_match:
                folder = folder_match.group(1)
                if folder == "photos":
                    folder = "pictures"
                return self.system.open_file(folder)
        return None

    def _handle_storage_command(self, text: str) -> str | None:
        lower = text.lower().strip()
        explicit_path = self._extract_storage_path(text)

        folder_match = re.search(r"\b(downloads|documents|pictures|photos)\s+(?:folder\s+)?(?:kholo|open|show|dikhao)\b", lower)
        if folder_match:
            folder = folder_match.group(1)
            if folder == "photos":
                folder = "pictures"
            opened = self.system.open_file(folder)
            if re.search(r"\b(?:batao|list|files|important|kya)\b", lower):
                listing = self.system.list_files_telegram(folder)
                return f"{opened}\n{listing}"
            return opened

        if re.search(r"\b(?:disk|storage|drive|space)\b.*\b(?:cleanup|clean up|report|audit|free)\b", lower) or re.search(r"\b(?:cleanup|clean up)\b.*\b(?:disk|storage|drive|space)\b", lower):
            return self.system.disk_cleanup_report()

        if explicit_path and lower.startswith(("open ", "show ", "launch ")):
            return self.system.open_file(explicit_path)

        if re.search(r"\b(?:open|launch|show)\b", lower) and re.search(r"\b(?:desktop|desk top)\b", lower):
            return self.system.open_desktop_folder()

        if re.search(r"\b(?:open|launch|show)\b", lower) and re.search(r"\b(?:file explorer|explorer|file manager)\b", lower):
            return self.system.open_app("file explorer")

        if re.search(r"\b(?:create|make|new)\s+folder\b", lower):
            folder_path = explicit_path or re.split(r"\b(?:create|make|new)\s+folder\b", text, maxsplit=1, flags=re.I)[-1].strip()
            return self.system.create_folder(folder_path) if folder_path else "Tell me which folder path to create."

        if re.search(r"\b(?:create|make|new)\s+file\b", lower):
            file_path = explicit_path or re.split(r"\b(?:create|make|new)\s+file\b", text, maxsplit=1, flags=re.I)[-1].strip()
            content = ""
            if " with content " in lower:
                parts = re.split(r"\bwith content\b", text, maxsplit=1, flags=re.I)
                file_path = explicit_path or parts[0].split("file", 1)[-1].strip()
                content = parts[1].strip() if len(parts) > 1 else ""
            return self.system.create_file(file_path, content) if file_path else "Tell me which file path to create."

        if re.search(r"\bopen (?:file|folder|document)\b", lower):
            path = explicit_path or re.split(r"\bopen (?:file|folder|document)\b", text, maxsplit=1, flags=re.I)[-1].strip()
            return self.system.open_file(path) if path else "Tell me which path to open."

        if re.search(r"\b(?:find|search|locate|where is)\s+file\b", lower):
            path = explicit_path
            query = re.split(r"\b(?:find|search|look for|locate|where is)\s+file\b", text, maxsplit=1, flags=re.I)[-1].strip()
            if path and query.endswith(path):
                query = query[:-len(path)].strip()
            if " in " in query.lower():
                parts = re.split(r"\bin\b", query, maxsplit=1, flags=re.I)
                query = parts[0].strip()
                path = path or parts[1].strip()
            return self.system.search_files(query, path) if query else "Tell me which filename to search for."

        return None

    def _handle_proactive_command(self, text: str) -> str | None:
        lower = text.lower().strip()
        if "proactive" not in lower:
            return None
        if not self.proactive:
            return "Proactive engine is not available."
        if any(word in lower for word in ("status", "state", "check")):
            return self.proactive.status()
        if any(word in lower for word in ("test", "try", "demo")):
            return self.proactive.test()
        if any(word in lower for word in ("disable", "off", "stop")):
            return self.proactive.set_enabled(False)
        if any(word in lower for word in ("enable", "on", "start")):
            return self.proactive.set_enabled(True)
        return self.proactive.status()

    def _handle_relay_command(self, text: str) -> str | None:
        lower = text.lower().strip()
        if not re.search(r"\b(?:light|bulb|relay|lamp|batti)\b", lower):
            return None
        if not bool(self.cfg.get("relay_enabled", False)):
            return None
        try:
            from jarvis_modules.relay_control import RelayControl
            relay = RelayControl(self.cfg)
            if re.search(r"\b(?:status|state|haal|check)\b", lower):
                status = relay.status()
                return f"Light abhi {status} hai bhai."
            if re.search(r"\b(?:on|chalu|jalaa|jala|start)\b", lower):
                return "Light on kar di bhai." if relay.on() else "NodeMCU ne on confirm nahi kiya."
            if re.search(r"\b(?:off|band|bujha|stop)\b", lower):
                return "Light band kar di bhai." if relay.off() else "NodeMCU ne off confirm nahi kiya."
        except Exception as exc:
            return f"NodeMCU se connect nahi ho paya: {exc}"
        return None

    def _handle_action_plan(self, text: str) -> str | None:
        plan = ActionPlanner.plan(text)
        if not plan or plan.get("intent") != "create_document_and_maybe_send":
            return None

        topic = plan.get("topic")
        recipient = plan.get("recipient")
        missing = set(plan.get("missing_fields") or [])

        if "topic" in missing:
            topic = self._ask_user("Document kis topic par banaun?")
            if not topic:
                return "Document cancelled."
            topic = ActionPlanner._clean_topic(topic)
        if not topic:
            return "Document kis topic par banana hai?"

        steps = plan.get("steps") or []
        should_send = "send_email" in steps
        should_send_telegram = "send_telegram" in steps
        if should_send and "recipient" in missing:
            recipient = self._ask_user("Kis email address par bheju?")
            if not recipient:
                return "Email cancelled."
        if should_send:
            recipient = ActionPlanner.extract_email(recipient or "")
            if not recipient:
                return "Valid email address nahi mila. Please email address clearly bhejo."

        title = self._title_for_topic(topic)
        content = self._draft_document_content(topic)
        doc_path = self.documents.create_docx(title, content, output_name=title)
        if not str(doc_path).lower().endswith(".docx"):
            return str(doc_path)

        if should_send_telegram:
            self._telegram_file_to_send = (doc_path, f"{title} document")
            return f"Document ready hai: {Path(doc_path).name}"

        if not should_send:
            if getattr(self, "_is_telegram_context", False):
                self._telegram_file_to_send = (doc_path, f"{title} document")
                return f"Document ready hai: {Path(doc_path).name}"
            return f"Document ready hai: {doc_path}"

        confirmation = self._ask_confirmation(
            f"Bhai document ready hai: {Path(doc_path).name}. Send kar du {recipient} pe?"
        )
        if not confirmation:
            return f"Document ready hai, send cancel kiya: {doc_path}"

        subject = f"{title} - JARVIS Document"
        body = f"Hi,\n\nAttached is the document on {topic} prepared by JARVIS.\n\nRegards,\nJARVIS"
        return self.email.send_email_with_attachment(recipient, subject, body, doc_path)

    def _ask_user(self, prompt: str) -> str | None:
        if getattr(self, "_is_telegram_context", False) and self.telegram:
            self.telegram.notify(prompt)
            return self.telegram.wait_for_reply()
        self.voice.speak(prompt)
        return self.listen_or_type()

    def _ask_confirmation(self, prompt: str) -> bool:
        answer = (self._ask_user(prompt) or "").strip().lower()
        return any(word in answer for word in ("yes", "haan", "ha", "send", "bhej", "kar do", "ok", "okay"))

    def _extract_email_fields(self, text: str) -> tuple[str | None, str | None, str | None]:
        raw = str(text or "").strip()
        to = EmailModule.normalize_recipient(ActionPlanner.extract_email(raw) or "")
        subject = None
        body = None

        subject_match = re.search(
            r"\bsubject\s+(.+?)(?=\s+(?:aur|and)?\s*body\b|\s+message\b|$)",
            raw,
            re.I,
        )
        if subject_match:
            subject = subject_match.group(1).strip(" .,:;-")

        body_match = re.search(r"\bbody\s+(.+)$", raw, re.I)
        if body_match:
            body = body_match.group(1).strip(" .,:;-")
        elif re.search(r"\bmessage\b", raw, re.I):
            parts = re.split(r"\bmessage\b", raw, maxsplit=1, flags=re.I)
            if len(parts) > 1:
                body = parts[1].strip(" .,:;-")

        return to or None, subject or None, body or None

    def _sensitive_file_response(self, text: str) -> str | None:
        lower = str(text or "").lower()
        sensitive_match = re.search(
            r"\b(?:jarvis_config\.json|\.env|env file|token(?:s)?|api[_ -]?key(?:s)?|password(?:s)?|secret(?:s)?)\b",
            lower,
        )
        if not sensitive_match:
            return None
        wants_file_action = re.search(
            r"\b(?:read|show|print|content|contents|batao|dikhao|dhundo|dhoondo|find|search|locate|where|send|share|bhej|deliver|upload)\b",
            lower,
        )
        if not wants_file_action:
            return None
        query = "jarvis_config.json" if "jarvis_config.json" in lower else sensitive_match.group(0)
        found = self.system.search_files(query)
        return (
            f"{found}\n"
            "Sensitive file hai, isliye main iska content/API keys/passwords Telegram ya chat me reveal nahi karunga."
        )

    def _handle_self_knowledge_command(self, text: str, force_rebuild: bool = False) -> str | None:
        lower = str(text or "").lower()
        wants_self_scan = bool(re.search(
            r"\b(?:read|scan|index|learn|study|update|rebuild)\b.*\b(?:j\.?a\.?r\.?v\.?i\.?s|jarvis|yourself|your code|folder|project)\b",
            lower,
        ))
        asks_self_build = bool(re.search(
            r"\b(?:how were you made|how are you made|what are you made of|what is inside (?:your|the) project|what do you know about yourself|self[- ]knowledge)\b",
            lower,
        ))
        if not (force_rebuild or wants_self_scan or asks_self_build):
            return None

        if not build_self_knowledge or not load_self_knowledge:
            return "Self-knowledge module available nahi hai, Prashant. jarvis_modules/self_knowledge.py import fail ho raha hai."

        should_rebuild = force_rebuild or wants_self_scan or not SELF_KNOWLEDGE_FILE.exists()
        if should_rebuild:
            try:
                data = build_self_knowledge(BASE_DIR, SELF_KNOWLEDGE_FILE)
            except Exception as exc:
                return f"Self-knowledge scan fail ho gaya: {exc}"
            action = "updated"
        else:
            data = load_self_knowledge(SELF_KNOWLEDGE_FILE)
            action = "loaded"

        if not data:
            return "Self-knowledge file abhi empty hai. Scan dobara try karo."

        entries = data.get("entries") or []
        skipped = data.get("skipped") or {}
        summary = str(data.get("summary", "")).strip()
        first_line = summary.splitlines()[0] if summary else "I am JARVIS, Prashant ka personal assistant."
        return (
            f"Self-knowledge {action} ho gaya, Prashant. "
            f"Maine {len(entries)} safe project entries index kiye aur secrets/sessions/binaries skip kiye. "
            f"Skipped: {skipped}. {first_line} "
            f"Ab tum puchh sakte ho: 'JARVIS, how were you made?'"
        )

    def _handle_self_improvement_command(self, text: str, force_queue: bool = False) -> str | None:
        if not bool(self.cfg.get("self_improvement_enabled", True)):
            return None
        if not (looks_like_self_improvement_request and save_self_improvement_request and response_for_request):
            return None
        if not (force_queue or looks_like_self_improvement_request(text)):
            return None

        request = save_self_improvement_request(
            SELF_IMPROVEMENT_REQUESTS_FILE,
            text,
            owner=self.name,
        )
        return response_for_request(request)

    def _handle_screen_document_command(self, text: str) -> str | None:
        lower = str(text or "").lower()
        if not (
            re.search(r"\b(?:current\s+screen|screen)\b", lower)
            and re.search(r"\b(?:summary|summarize|read|text)\b", lower)
            and re.search(r"\b(?:document|docx|doc|save)\b", lower)
        ):
            return None
        screen_text = self.system.read_screen_text()
        if not screen_text or screen_text.lower().startswith(("pyautogui not installed", "screen reading requires", "screen ocr failed")):
            return screen_text
        prompt = (
            "Summarize this current screen text clearly in Hinglish. "
            "Keep it useful and structured for a saved document.\n\n"
            f"{screen_text[:5000]}"
        )
        try:
            summary = self.ai.chat(prompt, context="Summarize current screen into document")
        except Exception:
            summary = screen_text[:3000]
        title = "Current Screen Summary"
        doc_path = self.documents.create_docx(title, str(summary or screen_text), output_name=title)
        if not str(doc_path).lower().endswith(".docx"):
            return str(doc_path)
        if getattr(self, "_is_telegram_context", False):
            self._telegram_file_to_send = (doc_path, "Current screen summary document")
            return f"Screen summary document ready hai: {Path(doc_path).name}"
        return f"Screen summary document saved: {doc_path}"

    def _title_for_topic(self, topic: str) -> str:
        words = re.split(r"\s+", str(topic or "Document").strip())
        return " ".join(word[:1].upper() + word[1:] for word in words if word) or "JARVIS Document"

    def _draft_document_content(self, topic: str) -> str:
        prompt = (
            f"Create a clear student-friendly document about: {topic}.\n"
            "Use a title, short introduction, key concepts, examples, practical uses, and summary. "
            "Keep it concise but complete. Plain text only; headings are allowed."
        )
        try:
            draft = self.ai.chat(prompt, context="Generate DOCX document content")
        except Exception as exc:
            draft = f"{topic}\n\nDocument generation failed through AI: {exc}"
        return str(draft or "").strip()

    def _extract_inline_note(self, text: str) -> str | None:
        raw = str(text or "").strip()
        cleaned = re.sub(r"\b(?:note|notes|bnao|bna|bana|take|write|jot|that|ki|karo)\b", " ", raw, flags=re.I)
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" .,:;-")
        return cleaned or None

    def _extract_inline_reminder(self, text: str) -> tuple[str | None, int | None]:
        raw = str(text or "").strip()
        lower = raw.lower()
        mins = None
        match = re.search(r"(\d+)\s*(?:minute|minutes|min|mins)", lower)
        if match:
            mins = int(match.group(1))
        elif "hour" in lower or "ghante" in lower:
            hour_match = re.search(r"(\d+)\s*(?:hour|hours|ghante|ghanta)", lower)
            mins = int(hour_match.group(1)) * 60 if hour_match else 60
        body = re.sub(r"\b(?:mujhe|remind me|reminder|set reminder|alert me|yaad|dilao|baad|after|in)\b", " ", raw, flags=re.I)
        body = re.sub(r"\d+\s*(?:minute|minutes|min|mins|hour|hours|ghante|ghanta)", " ", body, flags=re.I)
        body = re.sub(r"\s+", " ", body).strip(" .,:;-")
        return (body or None), mins

    def _extract_inline_event(self, text: str) -> tuple[str | None, str | None, str | None]:
        raw = str(text or "").strip()
        lower = raw.lower()
        today = datetime.date.today()
        date = None
        if re.search(r"\b(kal|tomorrow)\b", lower):
            date = (today + datetime.timedelta(days=1)).isoformat()
        elif re.search(r"\b(aaj|today)\b", lower):
            date = today.isoformat()
        else:
            match = re.search(r"\b(20\d{2}-\d{1,2}-\d{1,2})\b", lower)
            if match:
                date = match.group(1)
        time_str = None
        tm = re.search(r"\b(\d{1,2})(?::(\d{2}))?\s*(?:baje|am|pm)?\b", lower)
        if tm:
            hour = int(tm.group(1))
            minute = int(tm.group(2) or 0)
            if "pm" in lower and hour < 12:
                hour += 12
            if "am" in lower and hour == 12:
                hour = 0
            if "baje" in lower and hour <= 7 and not re.search(r"\bsubah|morning|am\b", lower):
                hour += 12
            time_str = f"{hour:02d}:{minute:02d}"
        title = re.sub(r"\b(?:kal|tomorrow|aaj|today|add|event|schedule|book|meeting|create|karo|karna|baje|am|pm)\b", " ", raw, flags=re.I)
        title = re.sub(r"\b20\d{2}-\d{1,2}-\d{1,2}\b", " ", title)
        title = re.sub(r"\b\d{1,2}(?::\d{2})?\b", " ", title)
        title = re.sub(r"\s+", " ", title).strip(" .,:;-")
        return title or None, date, time_str

    def handle(self, text: str) -> str:
        if not text or not text.strip():
            return ""
        if self.proactive:
            try:
                self.proactive.ping()
            except Exception:
                pass

        stripped_lower = text.lower().strip()
        # Explicit slash aliases keep their old behavior.
        # /ai is also a cowork alias — do NOT strip the prefix here so that
        # _aibrain_chat() can detect cowork mode and activate tool calling.
        if stripped_lower.startswith(("/cowork ", "/computer ", "/ai ")):
            intent = "ai_chat"
            lower = text.lower()
        else:
            self_improvement_resp = self._handle_self_improvement_command(text)
            if self_improvement_resp:
                self.voice.speak(self_improvement_resp)
                return self_improvement_resp
            self_knowledge_resp = self._handle_self_knowledge_command(text)
            if self_knowledge_resp:
                self.voice.speak(self_knowledge_resp)
                return self_knowledge_resp
            sensitive_resp = self._sensitive_file_response(text)
            if sensitive_resp:
                self.voice.speak(sensitive_resp)
                return sensitive_resp
            screen_doc_resp = self._handle_screen_document_command(text)
            if screen_doc_resp:
                self.voice.speak(screen_doc_resp)
                return screen_doc_resp
            action_resp = self._handle_action_plan(text)
            if action_resp:
                self.voice.speak(action_resp)
                return action_resp
            agent_resp = self._handle_agent_command(text)
            if agent_resp:
                self.voice.speak(agent_resp)
                return agent_resp
            browser_resp = self._handle_browser_command(text)
            if browser_resp:
                self.voice.speak(browser_resp)
                return browser_resp
            operator_resp = self._handle_operator_command(text)
            if operator_resp:
                self.voice.speak(operator_resp)
                return operator_resp
            proactive_resp = self._handle_proactive_command(text)
            if proactive_resp:
                self.voice.speak(proactive_resp)
                return proactive_resp
            relay_resp = self._handle_relay_command(text)
            if relay_resp:
                self.voice.speak(relay_resp)
                return relay_resp
            storage_resp = self._handle_storage_command(text)
            if storage_resp:
                self.voice.speak(storage_resp)
                return storage_resp
            intent = Intent.classify(text)
            lower  = text.lower()
        resp   = ""

        # ── STOP ─────────────────────────────────────
        if intent == "stop":
            resp = self.greeter.bye()
            self.voice.speak(resp)
            if self.proactive:
                self.proactive.stop()
            self._running = False
            return resp

        # ── GREET ─────────────────────────────────────
        elif intent == "greet":
            resp = self.greeter.greet()

        # ── TIME / DATE ───────────────────────────────
        elif intent == "time":
            resp = f"It's {datetime.datetime.now().strftime('%I:%M %p')}, {self.name}."

        elif intent == "date":
            resp = f"Today is {datetime.datetime.now().strftime('%A, %B %d, %Y')}."

        # ── OPEN APP ──────────────────────────────────
        elif intent == "open_app":
            app = re.sub(r"\b(open|launch|start|run|execute)\b","",lower,flags=re.I).strip()
            resp = self.system.open_app(app or "chrome")

        # ── CLOSE APP ─────────────────────────────────
        elif intent == "close_app":
            app = re.sub(r"\b(close|quit|exit|stop|kill|terminate)\b","",lower,flags=re.I).strip()
            resp = self.system.close_app(app or "")

        # ── LIST FILES ────────────────────────────────
        elif intent == "list_files":
            path = re.sub(r"\b(list|show|files|in|directory|folder|contents|what)\b","",lower,flags=re.I).strip()
            resp = self.system.list_files_telegram(path or None)

        # ── DOWNLOAD FILE ─────────────────────────────
        elif intent == "download_file":
            file_path = re.sub(r"\b(download|send|get|upload|file)\b","",lower,flags=re.I).strip()
            file_path, caption = self.system.download_file_telegram(file_path)
            if file_path:
                # For Telegram context, we'll handle file sending differently
                if hasattr(self, '_is_telegram_context') and self._is_telegram_context:
                    # Store file info for Telegram bot to handle
                    self._telegram_file_to_send = (file_path, caption)
                    resp = f"Preparing to send: {caption}"
                else:
                    resp = f"File ready: {caption}"
            else:
                resp = caption

        # ── HARDWARE ──────────────────────────────────
        elif intent == "hardware":
            resp = self.system.hardware_report()

        elif intent == "thermal_status":
            resp = self.system.thermal_report()

        elif intent == "fan_control":
            resp = self.system.open_hp_thermal_controls()

        elif intent == "network":
            resp = self.system.network_info()

        elif intent == "processes":
            resp = self.system.top_processes()

        # ── EMAIL ─────────────────────────────────────
        elif intent in ("email_check", "email_unread"):
            self.voice.speak("Checking your inbox...")
            resp = self.email.check_inbox(5)

        elif intent == "send_email":
            inline_to, inline_sub, inline_body = self._extract_email_fields(text)
            if getattr(self, "_is_telegram_context", False) and self.telegram:
                to = inline_to
                if not to:
                    self.telegram.notify("To whom?")
                    to = self.telegram.wait_for_reply()
                if not to: return "Email cancelled."
                to = EmailModule.normalize_recipient(to)
                if not to: return "Invalid email address. Please check recipient email and try again."
                sub = inline_sub
                if not sub:
                    self.telegram.notify("Subject?")
                    sub = self.telegram.wait_for_reply()
                if not sub: return "Email cancelled."
                body = inline_body
                if not body:
                    self.telegram.notify("What should I say?")
                    body = self.telegram.wait_for_reply()
                if not body: return "Email cancelled."
            else:
                to = inline_to
                if not to:
                    self.voice.speak("To whom?")
                    to = self.listen_or_type()
                if not to: return "Email cancelled."
                to = EmailModule.normalize_recipient(to)
                if not to: return "Invalid email address. Please check recipient email and try again."
                sub = inline_sub
                if not sub:
                    self.voice.speak("Subject?")
                    sub = self.listen_or_type()
                if not sub: return "Email cancelled."
                body = inline_body
                if not body:
                    self.voice.speak("What should I say?")
                    body = self.listen_or_type()
                if not body: return "Email cancelled."
            if not self._ask_confirmation(f"Send email to {to} with subject '{sub}'?"):
                resp = "Email cancelled."
            else:
                resp = self.email.send_email(to, sub, body) if all([to,sub,body]) else "Email cancelled."

        elif intent in ("send_whatsapp", "send_message"):
            if getattr(self, "_is_telegram_context", False) and self.telegram:
                self.telegram.notify("Contact name or WhatsApp number?")
                target = self.telegram.wait_for_reply()
                if not target: return "WhatsApp message cancelled."
                self.telegram.notify("What is the message?")
                msg = self.telegram.wait_for_reply()
                if not msg: return "WhatsApp message cancelled."
            else:
                self.voice.speak("Contact name or WhatsApp number?")
                target = self.listen_or_type()
                self.voice.speak("What is the message?")
                msg = self.listen_or_type()
            phone = ""
            if target:
                # If user spoke a name, try contacts.csv lookup
                looked = self.contacts.lookup_phone(target)
                phone = looked or target
            if phone and msg:
                if self._ask_confirmation(f"WhatsApp message send karu {target} pe?"):
                    resp = self.whatsapp.send_message(phone, msg)
                else:
                    resp = "WhatsApp message cancelled."
            else:
                resp = "WhatsApp message cancelled."

        elif intent in ("whatsapp_read", "whatsapp_unread"):
            self.voice.speak("Reading WhatsApp screen text now.")
            resp = self.whatsapp.read_incoming()

        elif intent == "media_control":
            if any(k in lower for k in ["next","skip"]):
                resp = self.system.control_media("next")
            elif any(k in lower for k in ["pause","stop"]):
                resp = self.system.control_media("playpause")
            elif any(k in lower for k in ["previous","prev","back"]):
                resp = self.system.control_media("previous")
            else:
                resp = self.system.control_media("playpause")

        elif intent == "screen_read":
            self.voice.speak("Reading screen text now.")
            resp = self.system.read_screen_text()

        # ── CALENDAR ──────────────────────────────────
        elif intent == "calendar_today":
            resp = self.calendar.today()

        elif intent == "calendar_upcoming":
            resp = self.calendar.upcoming()

        elif intent == "add_event":
            if getattr(self, "_is_telegram_context", False) and self.telegram:
                self.telegram.notify("Event title?")
                title = self.telegram.wait_for_reply()
                if not title: return "Event cancelled."
                self.telegram.notify("Date? For example, 2026-05-24.")
                date = self.telegram.wait_for_reply()
                if not date: return "Event cancelled."
                self.telegram.notify("Time?")
                tstr = self.telegram.wait_for_reply() or "00:00"
            else:
                self.voice.speak("Event title?")
                title = self.listen_or_type()
                self.voice.speak("Date? For example, 2026-05-24.")
                date  = self.listen_or_type()
                self.voice.speak("Time?")
                tstr  = self.listen_or_type() or "00:00"
            resp  = self.calendar.add(title, date, tstr) if title and date else "Event cancelled."

        elif intent == "open_calendar":
            webbrowser.open("https://calendar.google.com")
            resp = "Opening Google Calendar."

        # ── WEATHER ───────────────────────────────────
        elif intent == "weather":
            city = None
            ka_match = re.search(r"\b([A-Za-z][A-Za-z\s.-]{1,60}?)\s+ka\s+weather\b", text, re.I)
            if ka_match:
                city = ka_match.group(1).strip()
            for prep in ["in ","for ","at "]:
                if not city and prep in lower:
                    city = lower.split(prep,1)[-1].strip()
                    break
            if city:
                city = re.sub(r"\b(?:weather|batao|dikhao|forecast|ka|kya|hai)\b", " ", city, flags=re.I)
                city = re.sub(r"\s+", " ", city).strip(" .,:;-")
            resp = self.weather.current(city)

        # ── NOTES ─────────────────────────────────────
        elif intent == "note_add":
            content = self._extract_inline_note(text)
            if content:
                resp = self.notes.add(content)
            elif getattr(self, "_is_telegram_context", False) and self.telegram:
                self.telegram.notify("What should I note?")
                content = self.telegram.wait_for_reply()
                resp = self.notes.add(content) if content else "Note cancelled."
            else:
                self.voice.speak("What should I note?")
                content = self.listen_or_type()
                resp = self.notes.add(content) if content else "Note cancelled."

        elif intent == "note_read":
            resp = self.notes.read()

        elif intent == "note_clear":
            resp = self.notes.clear()

        # ── REMINDER ──────────────────────────────────
        elif intent == "reminder":
            rem, mins = self._extract_inline_reminder(text)
            if rem and mins:
                resp = self.notes.set_reminder(rem, mins, lambda m: self.voice.speak(m))
            elif getattr(self, "_is_telegram_context", False) and self.telegram:
                self.telegram.notify("What should I remind you about?")
                rem = self.telegram.wait_for_reply()
                if not rem: return "Reminder cancelled."
                self.telegram.notify("In how many minutes?")
                mins_str = self.telegram.wait_for_reply() or "5"
                try:
                    mins = int(re.search(r"\d+", mins_str).group())
                except:
                    mins = 5
                resp = self.notes.set_reminder(rem, mins, lambda m: self.voice.speak(m))
            else:
                self.voice.speak("What should I remind you about?")
                rem = self.listen_or_type()
                self.voice.speak("In how many minutes?")
                mins_str = self.listen_or_type() or "5"
                try:
                    mins = int(re.search(r"\d+", mins_str).group())
                except:
                    mins = 5
                resp = self.notes.set_reminder(rem, mins, lambda m: self.voice.speak(m))

        # ── NEWS ──────────────────────────────────────
        elif intent == "news":
            query = None
            if re.search(r"\bai\b", lower):
                query = "artificial intelligence"
            for kw in ["about ","on ","regarding "]:
                if not query and kw in lower:
                    query = lower.split(kw,1)[-1].strip()
                    break
            resp = self.news.headlines(query)

        # ── SCREENSHOT ────────────────────────────────
        elif intent == "screenshot":
            resp = self.system.take_screenshot()

        # ── VOLUME ────────────────────────────────────
        elif intent == "volume":
            nums = re.findall(r"\d+", text)
            if nums:
                resp = self.system.set_volume(int(nums[0]))
            elif "mute" in lower:
                resp = self.system.set_volume(0)
            else:
                resp = "Specify a volume level, like 'set volume to 60'."

        # ── FILE ──────────────────────────────────────
        elif intent == "search_files":
            inline_query = re.sub(r"\b(find|search|look for|locate|where is)\s+file\b", "", text, flags=re.I).strip()
            resp = self.system.search_files(inline_query) if inline_query else "Tell me which filename to search for."

        elif intent == "open_file":
            inline_path = self._extract_storage_path(text) or re.sub(r"\bopen (file|folder|document)\b", "", text, flags=re.I).strip()
            resp = self.system.open_file(inline_path) if inline_path else "Tell me which path to open."

        elif intent == "desktop":
            resp = self.system.list_desktop()

        elif intent == "clipboard":
            resp = self.system.get_clipboard()

        # ── CAMERA — PHOTO ────────────────────────────
        elif intent == "camera_photo":
            self.voice.speak("Capturing photo, hold still.")
            ok, result = self.camera.capture_photo()
            if ok and getattr(self, "_is_telegram_context", False):
                self._telegram_file_to_send = (result, "Webcam photo captured by JARVIS")
            resp = f"Photo saved at {result}." if ok else result

        # ── CAMERA — VIDEO ────────────────────────────
        elif intent == "camera_video":
            nums = re.findall(r"\d+", text)
            secs = int(nums[0]) if nums else 10
            self.voice.speak(f"Recording {secs} second video.")
            ok, result = self.camera.record_video(secs)
            if ok and getattr(self, "_is_telegram_context", False):
                self._telegram_file_to_send = (result, f"{secs} second webcam video captured by JARVIS")
            resp = f"Video saved at {result}." if ok else result

        # ── CAMERA — LIVE FEED ────────────────────────
        elif intent == "camera_live":
            self.voice.speak("Opening camera feed. Press Q to close.")
            threading.Thread(target=self.camera.show_live_feed, daemon=True).start()
            resp = "Camera feed opened. Press Q in the camera window to close."

        # ── CAMERA — AI ANALYSIS ──────────────────────
        elif intent == "camera_analyze":
            self.voice.speak("Taking a photo to analyze. One moment.")
            ok, path = self.camera.capture_photo()
            if ok:
                self.voice.speak("Analyzing the image...")
                question = text if len(text) > 15 else "Describe everything you see in detail."
                analysis = self.ai.analyze_image(path, question)
                resp = analysis
            else:
                resp = path  # error message

        # ── CAMERA — LIST ─────────────────────────────
        elif intent == "camera_list":
            resp = self.camera.list_cameras()

        # ── LOCATION ──────────────────────────────────
        elif intent == "location":
            loc = self.location.get_str()
            coords = self.location.get_coords()
            coord_str = f" Coordinates: {coords[0]:.4f}, {coords[1]:.4f}." if coords else ""
            if "map" in lower or "open map" in lower:
                self.location.open_in_browser()
                resp = f"Opening your location on Google Maps.{coord_str}"
            else:
                resp = f"Your current location is {loc}.{coord_str}"

        # ── POWER ─────────────────────────────────────
        elif intent == "shutdown":
            self.voice.speak(f"Are you sure you want to shut down, {self.name}? Say yes to confirm.")
            confirm = self.listen_or_type()
            resp = self.system.shutdown() if confirm and "yes" in confirm.lower() else "Shutdown cancelled."

        elif intent == "restart":
            resp = self.system.restart()

        elif intent == "lock":
            resp = self.system.lock_screen()

        # ── WEB SEARCH ────────────────────────────────
        elif intent == "web_search":
            query = re.sub(r"\b(search for|google|look up|search online|find online|search|karo|karna|pe|par|dikhao|batao)\b"," ",text,flags=re.I).strip()
            query = re.sub(r"\s+", " ", query).strip(" .,:;-")
            if query:
                q = requests.utils.quote(query) if requests else query.replace(" ", "+")
                webbrowser.open(f"https://www.google.com/search?q={q}")
                resp = f"Searching Google for '{query}'."
            else:
                resp = "What should I search for?"

        # ── CLEAR CHAT ────────────────────────────────
        elif intent == "clear_chat":
            self.ai.reset()
            resp = f"Conversation memory cleared, {self.name}. Fresh start."

        elif intent == "self_knowledge":
            resp = self._handle_self_knowledge_command(text, force_rebuild=True)

        elif intent == "self_improvement":
            resp = self._handle_self_improvement_command(text, force_queue=True)

        # ── JOKE ──────────────────────────────────────
        elif intent == "joke":
            jokes = [
                f"I tried to write a joke about artificial intelligence, {self.name}, but I couldn't think of one. I'll outsource it.",
                f"Why do programmers prefer dark mode? Because light attracts bugs, {self.name}.",
                f"I would tell you a UDP joke, {self.name}, but you might not get it.",
                f"Why did the hacker break up with the internet? Too many connections, {self.name}.",
            ]
            resp = random.choice(jokes)

        # ── AI FALLBACK ───────────────────────────────
        else:
            try:
                ctx  = f"Time: {datetime.datetime.now().strftime('%A %I:%M %p')}, OS: {platform.system()}"
                resp = self.ai.chat(text, context=ctx)
            except Exception as e:
                resp = f"AI error: {e}"

        if resp:
            self.voice.speak(resp)
        return resp

    # ── Input helper ──────────────────────────────────────────
    def listen_or_type(self) -> str | None:
        if self.text_mode:
            try:
                return input("  > ").strip() or None
            except:
                return None
        result = self.voice.listen(timeout=8)
        if result is None:
            try:
                result = input("  [Type] > ").strip() or None
            except:
                pass
        return result

    # ── Run modes ─────────────────────────────────────────────
    def run_text(self):
        self.text_mode = True
        self.boot()
        print(f"\n[TEXT MODE] Type commands below. Type 'exit' to quit.\n")
        while self._running:
            try:
                cmd = input(f"[{self.name}] > ").strip()
                if cmd:
                    self.handle(cmd)
            except KeyboardInterrupt:
                break
            except Exception as e:
                print(f"[Error] {e}")

    def run_voice(self):
        self.boot()
        self.voice.speak(f"Listening for '{self.wakes[0]}'.")
        while self._running:
            try:
                if self.voice.wait_wake_word(self.wakes):
                    self.voice.speak("Yes?")
                    cmd = self.voice.listen(timeout=8)
                    if cmd:
                        self.handle(cmd)
            except KeyboardInterrupt:
                break
            except Exception as e:
                if "PyAudio" not in str(e):
                    print(f"[Error] {e}")
                time.sleep(1)
        self.voice.speak(self.greeter.bye())

    def run_hybrid(self):
        self.boot()
        print(f"\n[INFO] Say '{self.wakes[0]}' OR just type below and press Enter.")
        print("[INFO] Press Ctrl+C to exit.\n")

        def text_thread():
            while self._running:
                try:
                    cmd = input("").strip()
                    if cmd:
                        self.handle(cmd)
                except (EOFError, KeyboardInterrupt):
                    break

        threading.Thread(target=text_thread, daemon=True).start()

        while self._running:
            try:
                if self.voice.wait_wake_word(self.wakes):
                    self.voice.speak("Yes?")
                    cmd = self.voice.listen(timeout=8)
                    if cmd:
                        self.handle(cmd)
            except KeyboardInterrupt:
                break
            except Exception as e:
                if "PyAudio" not in str(e):
                    print(f"[Error] {e}")
                time.sleep(0.5)

        self.voice.speak(self.greeter.bye())

# ═══════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    import argparse
    p = argparse.ArgumentParser(description="J.A.R.V.I.S v2.0")
    p.add_argument("--mode", choices=["text","voice","hybrid"], default="hybrid")
    args = p.parse_args()

    j = JARVIS()
    try:
        if   args.mode == "text":   j.run_text()
        elif args.mode == "voice":  j.run_voice()
        else:                       j.run_hybrid()
    except KeyboardInterrupt:
        print("\n[JARVIS] Shutting down gracefully...")
